from pathlib import Path
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "report" / "技术报告_一页纸摘要_图文最终版.docx"
OUT = ROOT / "report" / "技术报告_一页纸摘要_参考扩展版.docx"
ASSET_DIR = ROOT / "results" / "_summary_assets"


BODY_LEFT = [
    (
        "针对问题一：Nyx 数据读取与三维体数据恢复。",
        "Nyx 原始数据以 little-endian float32 的 .dat 二进制形式存储，不能直接作为图片或表格查看，因此本实验首先解决数据如何被正确读出的问题。项目根据文件大小推断体素数量和 n×n×n 网格规模，并按照 z 轴最快、y 轴其次、x 轴最慢的存储顺序，将一维数组恢复为 x-y-z 三维密度体。随后通过 XY、XZ、YZ 三方向中心切片验证读取方向和空间连续性，保证后续统计、体绘制和交互刷选均建立在正确数据结构之上。"
    ),
    (
        "针对问题二：宇宙密度时序统计特征分析。",
        "为刻画 Nyx 密度场在全部时间步中的演化过程，项目逐时计算 mean、std、P99、P99/mean、P99-P01 等指标，并构建 log-density histogram 与 time-density heatmap。图 A 展示全时间步密度分布的变化：早期分布较集中，后期宽度逐渐增大，高密度右尾更明显，说明密度场的不均匀性和极端高密度结构在演化过程中持续增强。"
    ),
    (
        "针对问题三：体数据渲染、传递函数与光照效果设计。",
        "针对二维切片难以表达三维宇宙网整体结构的问题，项目采用 log-density 变换和 P5-P99.7 百分位归一化压缩动态范围，并实现 alpha compositing 体绘制。通过 Balanced、Void、Filament、Node 等传递函数及梯度增强与类光照处理，分别突出整体密度、低密度空洞、中密度丝状结构和高密度节点。图 B 表明，密度场由早期较均匀状态逐步发展出更清晰的丝状连接和团块化节点。"
    ),
]

BODY_RIGHT = [
    (
        "针对问题四：P99 高密度筛选与宇宙网节点验证。",
        "为避免仅依靠体绘制图像进行主观判断，实验进一步使用原始 density 的 P99 作为 top 1% 高密度阈值，生成高密度 mask，并通过最大强度投影和多阈值等值面观察其空间位置。图 C 显示，P99 以上体素并非随机散布，而是集中出现在丝状结构交汇处和局部高密度节点区域，说明直方图右侧高密度尾部具有明确空间结构意义，也为体绘制中的高亮节点提供统计依据。"
    ),
    (
        "针对问题五：相空间交互式刷选与 Web 最终展示系统。",
        "在静态统计图和体绘制结果基础上，项目构建 linked brushing 交互分析机制，并开发 Nyx Density Explorer Web 单页可视化系统。用户可通过 time slider 观察演化，通过 histogram brush 或密度百分位滑条选择区间，并通过 Top 1% 一键筛选、MIP/Mask/Slice 模式切换将统计分布实时映射回空间视图。图 D 展示最终 Web 界面，使统计特征与空间物理结构的双向关联分析能够以交互方式呈现。"
    ),
    (
        "综合结论。",
        "综合来看，本项目形成“数据恢复—时序统计—体绘制—高密度验证—交互展示”的完整科学可视化流程。统计曲线与图 A 揭示密度分布从集中到分化的趋势，图 B 展示丝状结构和高密度节点逐渐显现的空间形态，图 C 与图 D 进一步验证高密度右尾和宇宙网节点之间的对应关系。本文结论基于给定 Nyx 数据和课程实验可视化结果，用于说明科学可视化技术在宇宙学模拟数据分析中的应用价值。"
    ),
]

CAPTIONS = [
    "图 A：时间-密度热力图",
    "图 B：体绘制关键帧",
    "图 C：Top 1% 高密度筛选",
    "图 D：Web 交互展示界面",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=40, start=55, bottom=35, end=55):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_table_width(table, width_twips):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def style_run(run, size=7.8, bold=False, color="111827"):
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def style_para(paragraph, size=7.8, line=9.5, before=0, after=1.2):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line)
    pf.first_line_indent = Pt(0)
    for run in paragraph.runs:
        style_run(run, size=size)


def add_body_block(cell, title, body):
    p = cell.add_paragraph()
    r = p.add_run(title)
    style_run(r, size=7.7, bold=True, color="0F3A5F")
    r2 = p.add_run(body)
    style_run(r2, size=7.7, color="111827")
    style_para(p, size=7.7, line=9.3, after=1.4)


def extract_images():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    images = []
    with zipfile.ZipFile(SRC) as zf:
        for name in zf.namelist():
            if name.startswith("word/media/"):
                out = ASSET_DIR / Path(name).name
                out.write_bytes(zf.read(name))
                images.append(out)
    return images[:4]


def build():
    images = extract_images()
    if len(images) < 4:
        raise RuntimeError(f"Expected 4 images in source DOCX, found {len(images)}")

    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.36)
    section.bottom_margin = Cm(0.34)
    section.left_margin = Cm(0.46)
    section.right_margin = Cm(0.46)

    styles = doc.styles
    styles["Normal"].font.name = "SimSun"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    styles["Normal"].font.size = Pt(7.8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0.6)
    r = p.add_run("摘要")
    style_run(r, size=13, bold=True, color="0B2742")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0.8)
    for idx, text in enumerate(["赛题 II：科学可视化挑战赛", "基于体绘制与相空间刷选的 Nyx 宇宙密度演化可视分析"]):
        if idx:
            p.add_run().add_break(WD_BREAK.LINE)
        r = p.add_run(text)
        style_run(r, size=8.8, bold=True, color="1F2937")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("姓名：          学号：          班级：")
    style_run(r, size=7.6, color="374151")

    usable_twips = 11360
    gap_twips = 120
    col_twips = (usable_twips - gap_twips) // 2
    text_table = doc.add_table(rows=1, cols=2)
    text_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    text_table.autofit = False
    remove_table_borders(text_table)
    set_table_width(text_table, usable_twips)
    for i, cell in enumerate(text_table.rows[0].cells):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_width(cell, col_twips)
        set_cell_margins(cell, 20, 50, 20, 50)
        set_cell_shading(cell, "F8FAFC")
        # Remove the empty default paragraph after adding content.
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    for title, body in BODY_LEFT:
        add_body_block(text_table.cell(0, 0), title, body)
    for title, body in BODY_RIGHT:
        add_body_block(text_table.cell(0, 1), title, body)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    spacer.paragraph_format.line_spacing = Pt(2)

    fig_table = doc.add_table(rows=2, cols=2)
    fig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fig_table.autofit = False
    remove_table_borders(fig_table)
    set_table_width(fig_table, usable_twips)
    for idx, cell in enumerate([c for row in fig_table.rows for c in row.cells]):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_width(cell, col_twips)
        set_cell_margins(cell, 25, 50, 10, 50)
        p_img = cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(0)
        run = p_img.add_run()
        run.add_picture(str(images[idx]), width=Cm(9.15))
        p_cap = cell.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(1)
        p_cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p_cap.paragraph_format.line_spacing = Pt(8)
        r = p_cap.add_run(CAPTIONS[idx])
        style_run(r, size=6.8, bold=True, color="334155")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
