from pathlib import Path
import re
import zipfile

from PIL import Image, ImageChops
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "report" / "技术报告_一页纸摘要_参考扩展版.docx"
OUT = ROOT / "report" / "技术报告_一页纸摘要_格式优化版.docx"
ASSET_DIR = ROOT / "results" / "_format_assets"

TITLE = "摘要"
SUBTITLE = "赛题 II：科学可视化挑战赛"
WORK_TITLE = "基于体绘制与相空间刷选的 Nyx 宇宙密度演化可视分析"
MEMBERS = "1191230327 邵奕博 1191230328 周凡升 1191230330 沈世轩"

CAPTIONS = [
    "图 A：时间-密度热力图，展示全时间步密度分布演化",
    "图 B：体绘制关键帧，展示宇宙网结构形成过程",
    "图 C：Top 1% 高密度筛选，验证高密度尾部空间位置",
    "图 D：Web 交互界面，实现直方图刷选与空间联动",
]


def set_run_font(run, size_pt, bold=False, color="000000", east_asia="SimSun", latin="Times New Roman"):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def para_compact(paragraph, line_pt=13, before=0, after=0, align=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_pt)
    pf.first_line_indent = None
    if align is not None:
        paragraph.alignment = align


def set_cell_margins(cell, top=30, start=55, bottom=25, end=55):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width))
    tbl_w.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def extract_body_blocks():
    doc = Document(SRC)
    text = "\n".join(
        p.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        for p in cell.paragraphs
        if p.text.strip()
    )
    text = re.sub(r"\n+", "\n", text)
    labels = ["针对问题一：", "针对问题二：", "针对问题三：", "针对问题四：", "针对问题五：", "综合结论。"]
    blocks = []
    for i, label in enumerate(labels):
        start = text.find(label)
        if start < 0:
            raise RuntimeError(f"Missing block: {label}")
        end_candidates = [text.find(next_label, start + len(label)) for next_label in labels[i + 1 :]]
        end_candidates = [x for x in end_candidates if x >= 0]
        end = min(end_candidates) if end_candidates else len(text)
        block = text[start:end].strip()
        block = re.sub(r"\s*图 A：.*$", "", block, flags=re.S)
        blocks.append((label, block[len(label) :].strip()))
    return blocks


def crop_image(src, dst):
    image = Image.open(src).convert("RGB")
    bg = Image.new("RGB", image.size, (255, 255, 255))
    diff = ImageChops.difference(image, bg)
    bbox = diff.getbbox()
    if bbox:
        pad = 8
        left = max(bbox[0] - pad, 0)
        top = max(bbox[1] - pad, 0)
        right = min(bbox[2] + pad, image.size[0])
        bottom = min(bbox[3] + pad, image.size[1])
        image = image.crop((left, top, right, bottom))
    image.save(dst)
    return dst


def extract_images():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    images = []
    with zipfile.ZipFile(SRC) as zf:
        for name in zf.namelist():
            if name.startswith("word/media/"):
                raw = ASSET_DIR / Path(name).name
                raw.write_bytes(zf.read(name))
                cropped = ASSET_DIR / f"cropped_{Path(name).name}"
                images.append(crop_image(raw, cropped))
    if len(images) < 4:
        raise RuntimeError("Expected four images in the source DOCX")
    return images[:4]


def add_body_paragraph(cell, label, body):
    p = cell.add_paragraph()
    para_compact(p, line_pt=12.4, after=0.5)
    r1 = p.add_run(label)
    set_run_font(r1, 8.5, bold=True, east_asia="SimHei", color="111827")
    r2 = p.add_run(body)
    set_run_font(r2, 8.5, color="111827")


def build():
    blocks = extract_body_blocks()
    images = extract_images()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(8.5)

    p = doc.add_paragraph()
    para_compact(p, line_pt=18, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(TITLE)
    set_run_font(r, 16, bold=True, east_asia="SimHei", latin="Times New Roman", color="0B1F33")

    p = doc.add_paragraph()
    para_compact(p, line_pt=14, after=0.6, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(SUBTITLE)
    set_run_font(r, 12, bold=True, east_asia="SimHei", color="111827")

    p = doc.add_paragraph()
    para_compact(p, line_pt=14, after=0.6, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(WORK_TITLE)
    set_run_font(r, 12, bold=True, east_asia="SimHei", color="111827")

    p = doc.add_paragraph()
    para_compact(p, line_pt=11.5, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(MEMBERS)
    set_run_font(r, 10.5, east_asia="SimSun", color="374151")

    usable_twips = 10205
    col_gap = 200
    col_twips = (usable_twips - col_gap) // 2

    body_table = doc.add_table(rows=1, cols=2)
    body_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    body_table.autofit = False
    set_table_width(body_table, usable_twips)
    remove_table_borders(body_table)
    left, right = body_table.rows[0].cells
    for cell in (left, right):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_width(cell, col_twips)
        set_cell_margins(cell, top=20, start=35, bottom=20, end=35)
        set_cell_shading(cell, "F7F9FC")
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    for label, body in blocks[:3]:
        add_body_paragraph(left, label, body)
    for label, body in blocks[3:]:
        add_body_paragraph(right, label, body)

    p = doc.add_paragraph()
    para_compact(p, line_pt=3, after=3)

    fig_table = doc.add_table(rows=2, cols=2)
    fig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fig_table.autofit = False
    set_table_width(fig_table, usable_twips)
    remove_table_borders(fig_table)

    for idx, cell in enumerate([c for row in fig_table.rows for c in row.cells]):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_width(cell, col_twips)
        set_cell_margins(cell, top=10, start=45, bottom=5, end=45)
        p_img = cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(0)
        p_img.paragraph_format.space_after = Pt(0)
        p_img.add_run().add_picture(str(images[idx]), width=Cm(8.28))
        p_cap = cell.add_paragraph()
        para_compact(p_cap, line_pt=8.2, after=1.6, align=WD_ALIGN_PARAGRAPH.CENTER)
        r = p_cap.add_run(CAPTIONS[idx])
        set_run_font(r, 6.8, bold=True, east_asia="SimSun", color="1F2937")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
