from pathlib import Path
import csv

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "report"
RESULTS_DIR = ROOT / "results"
STORY_DIR = RESULTS_DIR / "07_visual_story"
PROCESSED_DIR = ROOT / "data" / "processed"
STATS_PATH = PROCESSED_DIR / "density_stats.csv"
METRICS_PATH = PROCESSED_DIR / "structure_metrics.csv"


def read_csv(path):
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    for name in ["Normal", "Heading 1", "Heading 2", "List Bullet"]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)

    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(15.5)
    h1.font.color.rgb = RGBColor(26, 68, 112)
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(5)

    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(12.5)
    h2.font.color.rgb = RGBColor(46, 116, 181)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(3)


def add_title(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("数据可视化课程大作业实验报告")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(23)
    r.font.bold = True

    subtitles = [
        "赛题 II：科学可视化挑战赛",
        "Nyx 宇宙学模拟密度场可视分析",
        "作品主题：从暗雾到宇宙网骨架",
        "联合制作软件：MATLAB R2024b + Python + FFmpeg",
        "报告结构参考往期优秀模板：封面、目录、环境、方法、结果、创新与总结",
    ]
    for text in subtitles:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(12)
    doc.add_page_break()


def add_p(doc, text):
    doc.add_paragraph(text)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = text
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = "Microsoft YaHei"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    r.font.size = Pt(9)
                    if i == 0:
                        r.bold = True
    return table


def add_image(doc, path, caption, width=6.15):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(85, 85, 85)


def build_report():
    stats = read_csv(STATS_PATH)
    metrics = read_csv(METRICS_PATH)
    first, mid, last = stats[0], stats[49], stats[-1]
    m_first, m_mid, m_last = metrics[0], metrics[49], metrics[-1]
    std_growth = (float(last["std_density"]) / float(first["std_density"]) - 1) * 100
    max_growth = (float(last["max_density"]) / float(first["max_density"]) - 1) * 100
    spread_growth = (float(m_last["spread_p99_p01"]) / float(m_first["spread_p99_p01"]) - 1) * 100
    void_delta = (float(m_last["void_fraction_vs_t0000_p05"]) - float(m_first["void_fraction_vs_t0000_p05"])) * 100
    entropy_growth = (float(m_last["log_density_entropy"]) / float(m_first["log_density_entropy"]) - 1) * 100

    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("目录", level=1)
    for item in [
        "1 赛题理解与作品定位",
        "2 多软件联合制作流程",
        "3 数据读取、质量检查与预处理",
        "4 体绘制方案：传递函数、光照与视觉编码",
        "5 宇宙网视觉叙事：从暗雾到骨架",
        "6 时序统计：100 步密度分布如何偏移",
        "7 相空间交互刷选：从统计尾部回到空间结构",
        "8 演示视频、作品创新点与提交说明",
        "9 总结、局限与进一步改进",
    ]:
        doc.add_paragraph(item)
    doc.add_page_break()

    doc.add_heading("1 赛题理解与作品定位", level=1)
    add_p(doc, "本作品选择赛题 II：科学可视化挑战赛。Nyx 数据描述宇宙学模拟中气体密度随时间的演化，核心问题不是简单画出一张漂亮图片，而是把三维标量场、时间序列统计和可交互阈值筛选连接起来。")
    add_p(doc, "我将作品定位为一个“宇宙网形成观测台”：第一层通过体绘制观察空间形态，第二层通过 100 个时间步统计曲线观察分布迁移，第三层通过相空间刷选验证高密度尾部是否对应真实的三维节点结构。")
    add_p(doc, "视觉叙事上，报告把宇宙密度演化概括为四个阶段：暗雾、丝束、节点、空洞。暗雾代表早期近均匀背景；丝束代表物质被引力牵引形成连续结构；节点代表高密度峰值抬升；空洞代表低密度区域在结构分化中扩张。")

    doc.add_heading("2 多软件联合制作流程", level=1)
    add_table(doc, [
        ["软件/工具", "承担任务", "形成成果"],
        ["MATLAB R2024b", "体数据读取、统计计算、体绘制、直方图、相空间交互仪表盘", "PNG 图像、CSV/MAT 统计文件、interactive_dashboard.m"],
        ["Python + Pillow", "基于 MATLAB 输出二次视觉设计，生成图谱、热力图、指标仪表盘、视频帧", "07_visual_story 作品图谱"],
        ["Python + python-docx / ReportLab", "自动生成 DOCX 与 PDF 报告，保证报告和数据结果同步", "技术报告、Answer Sheet"],
        ["FFmpeg", "将 100 帧数据驱动分镜合成为演示视频", "video/demo.mp4"],
    ])
    add_p(doc, "与单一软件截图不同，本项目让 MATLAB 负责科学计算与交互，Python 负责视觉叙事包装与文档生成，FFmpeg 负责视频化展示。三者形成从数据、图像、报告到视频的制作链路。")
    add_image(doc, RESULTS_DIR / "06_python_summary" / "software_workflow_summary.png", "图 1  MATLAB + Python 联合制作流程与结果汇总", width=6.4)

    doc.add_heading("3 数据读取、质量检查与预处理", level=1)
    add_p(doc, "每个 Nyx 文件以 little-endian float32 保存，大小为 8,388,608 字节，对应 128×128×128 个体素。题目说明数据线性顺序为先 z、再 y、最后 x，因此读取时先 reshape 为 z-y-x，再转换为 x-y-z 空间。")
    add_p(doc, "数据预处理坚持“只用给定数据”。没有引入任何外部天文数据或纹理。所有图像、曲线和视频帧都来自 data/raw 中 100 个时间步。为了让高密度峰值与主体分布同时可见，对密度采用 log10 动态范围压缩。")
    add_image(doc, RESULTS_DIR / "01_data_check" / "slice_0000.png", "图 2  t0000 中央切片质量检查：密度场连续，无明显读取错位")
    doc.add_page_break()

    doc.add_heading("4 体绘制方案：传递函数、光照与视觉编码", level=1)
    add_p(doc, "体绘制采用前向 alpha 合成。传递函数的关键思想是：低密度空洞不完全隐藏，而是以深色暗雾保留背景；中等密度结构用青蓝色显示连续丝束；极高密度节点用金色到暖白色突出。这样既能看到宇宙网骨架，也能看到空洞与丝束之间的空间关系。")
    add_p(doc, "光照并不追求真实天文照明，而是使用梯度幅值近似“结构边界强度”。密度变化越剧烈的区域越亮，视觉上更容易识别空洞边界、丝状结构和团块节点。")
    add_table(doc, [
        ["环节", "实现方式", "分析目的"],
        ["动态范围压缩", "对密度取 log10，并用分位数裁剪到稳定显示区间", "保留主体结构，同时避免极端峰值压暗整体图像"],
        ["透明度设计", "低密度给极低 alpha，中高密度随 smoothstep 增强", "让空洞作为背景存在，让丝状结构和节点成为视觉焦点"],
        ["梯度光照", "用三维梯度幅值增强边界亮度", "突出密度突变边界，帮助识别空洞壁和宇宙网骨架"],
        ["时间对比", "代表帧采用同一视觉语法和一致阈值思想", "避免逐帧单独调色造成演化强弱的误判"],
    ])
    add_image(doc, STORY_DIR / "transfer_function_design.png", "图 3  传递函数设计图：颜色、透明度与梯度光照的联合编码", width=6.4)
    add_image(doc, RESULTS_DIR / "02_volume_render" / "volume_t0000.png", "图 4  t0000 体绘制：密度背景较均匀，结构边界较弱")
    add_image(doc, RESULTS_DIR / "02_volume_render" / "volume_t0099.png", "图 5  t0099 体绘制：高密度节点和丝状连接更加突出")

    doc.add_heading("5 宇宙网视觉叙事：从暗雾到骨架", level=1)
    add_p(doc, "优秀的科学可视化不只是把数据画出来，还要让读者看到数据背后的过程。因此我额外制作了“宇宙网演化图谱”，用同一视觉语法比较 t0000、t0030、t0060、t0099 四个代表时间步。")
    add_p(doc, "图谱显示早期物质分布像弥散暗雾，之后沿引力势阱被拉成丝束，最终高密度节点逐渐发亮，低密度区域退成大尺度空洞。这个过程与宇宙大尺度结构形成的基本图景一致。")
    add_image(doc, STORY_DIR / "cosmic_web_atlas.png", "图 6  Nyx 宇宙网演化图谱：从暗雾到骨架", width=6.4)
    doc.add_page_break()

    doc.add_heading("6 时序统计：100 步密度分布如何偏移", level=1)
    add_table(doc, [
        ["指标", "t0000", "t0049", "t0099", "变化含义"],
        ["均值", f"{float(first['mean_density']):.4f}", f"{float(mid['mean_density']):.4f}", f"{float(last['mean_density']):.4f}", "全域平均密度缓慢变化"],
        ["标准差", f"{float(first['std_density']):.4f}", f"{float(mid['std_density']):.4f}", f"{float(last['std_density']):.4f}", "密度波动增强"],
        ["最大值", f"{float(first['max_density']):.4f}", f"{float(mid['max_density']):.4f}", f"{float(last['max_density']):.4f}", "高密度峰值抬升"],
        ["P99-P01", f"{float(m_first['spread_p99_p01']):.4f}", f"{float(m_mid['spread_p99_p01']):.4f}", f"{float(m_last['spread_p99_p01']):.4f}", "分布跨度增大"],
        ["低密度空洞占比", f"{float(m_first['void_fraction_vs_t0000_p05'])*100:.2f}%", f"{float(m_mid['void_fraction_vs_t0000_p05'])*100:.2f}%", f"{float(m_last['void_fraction_vs_t0000_p05'])*100:.2f}%", "相对早期阈值的空洞扩张"],
    ])
    add_p(doc, f"从统计量看，标准差由 {float(first['std_density']):.4f} 增加到 {float(last['std_density']):.4f}，最大密度由 {float(first['max_density']):.4f} 增加到 {float(last['max_density']):.4f}。这说明演化后期物质分布更不均匀，极端高密度区域更强。")
    add_table(doc, [
        ["派生量", "变化幅度", "解释"],
        ["标准差增长率", f"{std_growth:+.1f}%", "全域密度波动增强，均匀背景被拉开"],
        ["最大密度增长率", f"{max_growth:+.1f}%", "局部高密度峰值持续被引力聚集放大"],
        ["P99-P01 跨度增长率", f"{spread_growth:+.1f}%", "低密度端与高密度端同步分化"],
        ["空洞占比增量", f"{void_delta:+.1f} 个百分点", "相对 t0000 的低密度阈值，后期空洞区域显著扩张"],
        ["log 密度熵增长率", f"{entropy_growth:+.1f}%", "分布形态更复杂，单峰集中性减弱"],
    ])
    add_image(doc, STORY_DIR / "structure_metrics_dashboard.png", "图 7  结构形成统计指纹：波动、峰值、空洞占比和分布复杂度")
    add_image(doc, STORY_DIR / "histogram_temporal_heatmap.png", "图 8  100 个时间步密度对数直方图热力图")
    add_p(doc, "热力图把 100 张直方图压缩到一张图里。横向看是时间推进，纵向看是密度区间。高密度尾部在后期保持并略微增强，低密度端也逐步展开，形成两极分化。")

    doc.add_heading("7 相空间交互刷选：从统计尾部回到空间结构", level=1)
    add_p(doc, "交互仪表盘 interactive_dashboard.m 的目标是把统计空间和物理空间打通。用户在直方图中选择密度百分位区间，例如 99% 到 100% 的前 1% 高密度尾部，右侧三维空间视图会实时显示满足阈值的体素。")
    add_p(doc, "这个设计回答了一个关键问题：直方图中的尾部数据到底是什么？结果表明，高密度尾部不是随机散点，而是集中出现在宇宙网节点与致密丝束附近。统计异常因此获得了空间物理解释。")
    add_table(doc, [
        ["交互环节", "数据处理", "联动结果"],
        ["直方图刷选", "滑块把百分位转换为 log10(density) 上下界", "红色参考线实时标出被选密度区间"],
        ["空间掩膜", "对三维体素计算 low <= logDensity <= high", "只保留满足阈值的物理单元格"],
        ["三维投影", "沿视线方向对掩膜后的 log 密度取最大投影", "节点和致密丝束在右侧空间视图中显影"],
        ["解释闭环", "统计尾部与空间位置共享同一阈值", "验证高密度尾部对应宇宙网骨架而非随机噪声"],
    ])
    add_image(doc, STORY_DIR / "interaction_storyboard.png", "图 9  相空间刷选故事板：从直方图尾部定位三维节点结构", width=6.4)
    add_image(doc, RESULTS_DIR / "05_interaction_selection" / "top1_percent_t0099.png", "图 10  t0099 前 1% 高密度区域刷选结果")

    doc.add_heading("8 演示视频、作品创新点与提交说明", level=1)
    add_p(doc, "为了让宇宙演化过程更直观，项目额外生成 video/demo.mp4。视频不是图片轮播，而是逐帧读取 100 个原始时间步，在每一帧同时展示空间投影、当前统计指标和直方图热力图中的时间位置。")
    add_bullets(doc, [
        "创新点 1：同一套传递函数贯穿所有时间步，避免每张图单独调色导致的对比误差。",
        "创新点 2：用 100 步直方图热力图展示全域分布迁移，比只放 4 张直方图更完整。",
        "创新点 3：把前 1% 高密度尾部与三维空间结构联动，完成统计特征到物理结构的解释闭环。",
        "创新点 4：MATLAB、Python、FFmpeg 多软件协作，覆盖计算、交互、视觉叙事、报告和视频生成。",
    ])
    add_p(doc, "最终提交包包含核心代码、处理数据、结果图、DOCX/PDF 报告、Answer Sheet 与演示视频。原始 data/raw 数据约 800MB，默认不重复放入 zip，避免超过课程电子材料大小限制。")

    doc.add_heading("9 总结、局限与进一步改进", level=1)
    add_p(doc, "本项目用体绘制、统计曲线、时序热力图和交互刷选共同说明：Nyx 密度场在 100 个时间步内由相对均匀的分布逐步发展出更明显的高密度节点、丝状连接和低密度空洞。可视化技术的价值在于让抽象的三维标量场同时具备空间形态、统计证据和交互验证。")
    add_p(doc, "局限性在于当前体绘制采用 CPU 前向合成与投影近似，未实现真正的 GPU 射线投射；交互仪表盘侧重密度阈值筛选，尚未加入温度、速度或暗物质粒子的多变量联动。后续可以加入 WebGL/Volume Viewer、等值面提取、连通域跟踪和节点生命周期分析。")

    out = REPORT_DIR / "技术报告.docx"
    doc.save(out)
    return out


def build_answer_sheet():
    doc = Document()
    style_document(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Answer Sheet")
    r.font.size = Pt(18)
    r.bold = True
    doc.add_heading("作品信息", level=1)
    add_table(doc, [
        ["项目", "内容"],
        ["赛题", "赛题 II：科学可视化挑战赛"],
        ["作品", "Nyx 宇宙学模拟密度场可视分析"],
        ["主题", "从暗雾到宇宙网骨架"],
        ["联合制作软件", "MATLAB R2024b + Python + FFmpeg"],
        ["核心成果", "体绘制、时序热力图、结构指标仪表盘、相空间交互刷选、代表图、演示视频"],
    ])
    doc.add_heading("任务覆盖说明", level=1)
    add_table(doc, [
        ["任务要求", "对应成果", "核心文件"],
        ["体数据渲染、传递函数与光照", "报告第 4 节；体绘制图、传递函数设计图", "code/step4_volume_render.m；results/02_volume_render；results/07_visual_story/transfer_function_design.png"],
        ["宇宙密度演化特征归纳", "报告第 5、9 节；宇宙网演化图谱", "results/07_visual_story/cosmic_web_atlas.png；video/demo.mp4"],
        ["100 步密度统计与对数直方图", "报告第 6 节；统计曲线、直方图热力图", "code/step2_density_statistics.m；code/step3_draw_histograms.m；data/processed/density_stats.csv"],
        ["相空间交互刷选联动分析", "报告第 7 节；前 1% 高密度刷选结果和交互故事板", "code/interactive_dashboard.m；code/step5_high_density_selection.m；results/05_interaction_selection"],
    ])
    doc.add_heading("提交清单", level=1)
    add_bullets(doc, [
        "code：MATLAB 主程序、交互仪表盘、Python 视觉故事/报告/视频生成脚本。",
        "results：体绘制、直方图、统计曲线、高密度刷选、联合流程图、视觉故事图谱。",
        "data/processed：density_stats.csv、high_density_threshold.csv、structure_metrics.csv、statistics.mat。",
        "report：技术报告 DOCX/PDF 与 Answer Sheet DOCX/PDF。",
        "results/representative_image.jpg：官方要求的 JPG 代表图。",
        "video：demo.mp4 演示视频。",
    ])
    out = REPORT_DIR / "Answer_Sheet.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    print(build_report())
    print(build_answer_sheet())
