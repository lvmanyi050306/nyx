# -*- coding: utf-8 -*-
"""生成 Nyx 项目的“实验迭代过程版”报告、代码说明和路线图。

运行方式：
    python code/python/build_iteration_process_documents.py

输出：
    report/大作业技术报告草稿_迭代过程版.md
    report/代码说明_迭代过程版.md
    report/Nyx宇宙密度演化可视分析_技术报告_迭代过程版.docx
    report/Nyx宇宙密度演化可视分析_代码说明_迭代过程版.docx
    results/report_figures/iteration_roadmap.png
"""

from __future__ import annotations

from pathlib import Path
import re
import textwrap


PROJECT_ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = PROJECT_ROOT / "report"
FIG_DIR = PROJECT_ROOT / "results" / "report_figures"


def create_iteration_roadmap(output_path: Path) -> None:
    """用代码自动生成“从无到有”的实验迭代路线图。"""

    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
    from matplotlib.font_manager import FontProperties

    font_path = Path("C:/Windows/Fonts/msyh.ttc")
    if not font_path.exists():
        font_path = Path("C:/Windows/Fonts/simhei.ttf")
    cn_font = FontProperties(fname=str(font_path))
    cn_font_bold = FontProperties(fname=str(font_path), weight="bold")

    stages = [
        ("1 数据理解", ".dat 二进制文件\nlittle-endian float32\n体素数量推断"),
        ("2 数据恢复", "读取数据\n推断 n×n×n\nz-y-x 转 x-y-z"),
        ("3 读取验证", "文件大小检查\nXY/XZ/YZ 中心切片\n确认方向正确"),
        ("4 基础统计", "单帧统计\n全时间步统计\ndensity_stats.csv"),
        ("5 显示增强", "原始 density\nlog-density\nP5-P99.7 归一化"),
        ("6 空间可视化", "中心切片\nMIP 与体绘制\n传递函数设计"),
        ("7 结构验证", "P99 筛选\n多阈值等值面\n结构指标"),
        ("8 交互分析", "直方图刷选\n空间投影联动\nlinked brushing"),
        ("9 进阶分析", "density-gradient\nHessian 分类\n时间相似性矩阵"),
        ("10 最终系统", "统计分布 + 三维结构\n交互探索 + 结构识别"),
    ]

    output_path.parent.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(14, 7.6), dpi=300)
    fig.patch.set_facecolor("white")
    ax.set_xlim(0, 5)
    ax.set_ylim(0, 2.25)
    ax.axis("off")

    colors = ["#E8F1FA", "#EAF7EF", "#FFF4D9", "#FDEDEE", "#F0EDF8"]
    edge = "#D0D7DE"
    title_color = "#17324D"

    positions = []
    for i in range(10):
        row = 1 if i < 5 else 0
        col = i if i < 5 else 9 - i
        x = col + 0.08
        y = 1.24 if row == 1 else 0.18
        positions.append((x, y))

    for i, ((title, desc), (x, y)) in enumerate(zip(stages, positions)):
        box = FancyBboxPatch(
            (x, y),
            0.84,
            0.72,
            boxstyle="round,pad=0.025,rounding_size=0.045",
            linewidth=1.4,
            edgecolor=edge,
            facecolor=colors[i % len(colors)],
        )
        ax.add_patch(box)
        ax.text(x + 0.42, y + 0.54, title, ha="center", va="center",
                fontsize=11, color=title_color, fontproperties=cn_font_bold)
        ax.text(x + 0.42, y + 0.28, desc, ha="center", va="center",
                fontsize=8.3, color="#2B3440", linespacing=1.25, fontproperties=cn_font)

    for i in range(9):
        x0, y0 = positions[i]
        x1, y1 = positions[i + 1]
        if i == 4:
            start = (x0 + 0.42, y0 - 0.02)
            end = (x1 + 0.42, y1 + 0.74)
            conn = "arc3,rad=0.25"
        else:
            start = (x0 + (0.86 if i < 4 else -0.02), y0 + 0.36)
            end = (x1 + (-0.02 if i < 4 else 0.86), y1 + 0.36)
            conn = "arc3,rad=0.0"
        arrow = FancyArrowPatch(
            start,
            end,
            arrowstyle="-|>",
            mutation_scale=12,
            linewidth=1.4,
            color="#7A869A",
            connectionstyle=conn,
        )
        ax.add_patch(arrow)

    ax.text(
        2.5,
        2.12,
        "Nyx Visualization Project: Iterative Development Roadmap",
        ha="center",
        va="center",
        fontsize=17,
        color=title_color,
        fontproperties=cn_font_bold,
    )
    ax.text(
        2.5,
        2.00,
        "Problem-driven workflow: data recovery → validation → statistics → rendering → verification → interaction → advanced analysis",
        ha="center",
        va="center",
        fontsize=9.5,
        color="#58606A",
        fontproperties=cn_font,
    )
    fig.savefig(output_path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def create_web_system_overview(output_path: Path) -> None:
    """生成 Web 交互展示系统概览图。"""

    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
    from matplotlib.font_manager import FontProperties

    font_path = Path("C:/Windows/Fonts/msyh.ttc")
    if not font_path.exists():
        font_path = Path("C:/Windows/Fonts/simhei.ttf")
    cn_font = FontProperties(fname=str(font_path))
    cn_font_bold = FontProperties(fname=str(font_path), weight="bold")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(13.8, 6.8), dpi=300)
    fig.patch.set_facecolor("white")
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6)
    ax.axis("off")

    boxes = [
        ("原始 Nyx 数据", ".dat little-endian float32\n128×128×128 时序体", 0.45, 4.1, "#E8F1FA"),
        ("Web 预处理", "log-density\nP5-P99.7 归一化\n64³ 降采样", 3.0, 4.1, "#EAF7EF"),
        ("Web 数据资产", "metadata / stats / histograms\nvol_0000.bin ... vol_0099.bin", 5.55, 4.1, "#FFF4D9"),
        ("HTML 单页系统", "index.html + Canvas\n懒加载当前时间步", 8.1, 4.1, "#F0EDF8"),
        ("交互控制", "time slider / Play\npercentile brush\nprojection / transfer function", 1.4, 1.35, "#FDEDEE"),
        ("联动视图", "Spatial View\nHistogram & Brushing\nTime-density Heatmap\nMetric Curves", 4.35, 1.35, "#E9F7FA"),
        ("答辩演示", "top 1% 高密度体素\n统计右尾与空间节点\n课程实验结论", 7.3, 1.35, "#FFF0E5"),
    ]

    for title, desc, x, y, color in boxes:
        patch = FancyBboxPatch(
            (x, y),
            2.08,
            1.06,
            boxstyle="round,pad=0.04,rounding_size=0.06",
            linewidth=1.35,
            edgecolor="#D0D7DE",
            facecolor=color,
        )
        ax.add_patch(patch)
        ax.text(x + 1.04, y + 0.76, title, ha="center", va="center",
                fontsize=12, color="#17324D", fontproperties=cn_font_bold)
        ax.text(x + 1.04, y + 0.39, desc, ha="center", va="center",
                fontsize=8.7, color="#2B3440", linespacing=1.25, fontproperties=cn_font)

    arrows = [
        ((2.53, 4.63), (3.0, 4.63)),
        ((5.08, 4.63), (5.55, 4.63)),
        ((7.63, 4.63), (8.1, 4.63)),
        ((9.14, 4.1), (8.34, 2.41)),
        ((2.44, 1.88), (4.35, 1.88)),
        ((6.43, 1.88), (7.3, 1.88)),
        ((6.55, 4.1), (5.4, 2.41)),
    ]
    for start, end in arrows:
        ax.add_patch(FancyArrowPatch(start, end, arrowstyle="-|>", mutation_scale=13,
                                     linewidth=1.5, color="#7A869A"))

    ax.text(6, 5.65, "Nyx Web Interactive Visualization System",
            ha="center", va="center", fontsize=17, color="#17324D", fontproperties=cn_font_bold)
    ax.text(6, 5.42, "Offline preprocessing + browser-side lazy loading + linked brushing dashboard",
            ha="center", va="center", fontsize=9.8, color="#58606A", fontproperties=cn_font)
    fig.savefig(output_path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def report_markdown() -> str:
    return r"""# 基于体绘制与相空间刷选的 Nyx 宇宙密度演化可视分析

课程名称：数据可视化  
作品类型：科学可视化 / 体数据可视化 / 时序数据可视分析  
数据来源：Nyx 宇宙学模拟密度数据  
说明：本文以 ChinaVis 2026 赛题 II 科学可视化挑战赛中的 Nyx 数据为课程大作业数据来源，目标是完成面向课程实验的可视分析系统设计与技术报告撰写，并非正式参赛答卷或天体物理研究论文。

## 摘要

Nyx 宇宙学模拟数据以连续二进制 `.dat` 文件存储多个时间步的三维密度场，文件本身无法像图像或表格一样直接查看。因此，本作业不是一次性搭建完整可视化系统，而是围绕“数据能否读出、读出后如何验证、如何从局部观察扩展到全时序分析、如何把统计分布与三维结构联系起来”这一系列问题逐步迭代。项目首先从 little-endian float32 二进制读取开始，通过体素数量自动推断网格尺寸，并将原始 z-y-x 线性顺序恢复为 MATLAB 中更直观的 x-y-z 三维体数据；随后利用 XY、XZ、YZ 中心切片和文件大小检查验证数据读取是否正确。在此基础上，分析从单帧切片扩展到全时间步统计，计算 mean、std、max、百分位数、偏度、峰度等指标，并引入 log-density 变换和百分位归一化解决原始 density 动态范围过大的显示问题。进一步地，项目加入自定义 alpha compositing 体绘制、LUT、smoothstep 透明度函数和梯度增强，以展示宇宙网的三维空间形态。为了避免只停留在“看起来像”的视觉判断，后续又加入 P99 top 1% 高密度筛选、多阈值等值面、结构指标和首末时间步差分 MIP，将直方图右尾与三维空间节点联系起来。静态图像仍然缺少探索性，因此项目继续实现 linked brushing 仪表盘，使用户能够通过密度百分位区间联动查看统计分布和空间投影。最后，项目在原有流程上扩展 density-gradient 二维相空间、Hessian 特征值宇宙网结构分类和时间步相似性矩阵，并进一步整理为 HTML 单页 Web 交互展示系统。Web 版本通过 Python 预处理将体数据降采样为浏览器可懒加载的 Float32Array，并用原生 Canvas 实现时间步播放、直方图刷选、top 1% 高密度区域筛选、MIP/Mask/Slice/简化体合成和 time-density heatmap 联动。实验结果表明，该迭代流程形成了“数据恢复—统计分析—三维渲染—结构验证—交互探索—进阶识别—Web 展示”的完整链条，能够较系统地展示 Nyx 密度场从相对均匀分布到空洞、丝状结构和高密度节点逐渐显现的可视化特征。本文结论主要服务于数据可视化课程实验，不作为正式天体物理科学结论。

## 目录

1 问题描述  
1.1 研究背景  
1.2 数据特点  
1.3 本作业的可视化目标  
1.4 实验迭代思路  

2 数据读取与初步验证：从不可见二进制到三维密度体  
2.1 原始 .dat 数据格式分析  
2.2 little-endian float32 读取  
2.3 网格尺寸自动推断  
2.4 z-y-x 到 x-y-z 的轴顺序恢复  
2.5 中心切片验证与文件大小检查  
2.6 本阶段小结：解决“数据能否正确读出”的问题  

3 基础统计分析：从单帧观察到全时间步量化  
3.1 为什么不能只看单个时间步  
3.2 全时间步统计量设计  
3.3 mean、std、max 和百分位数  
3.4 log-density 变换的引入  
3.5 对数密度直方图  
3.6 时间-密度热力图  
3.7 本阶段小结：从图像观察扩展到时序统计  

4 三维空间可视化：从二维切片到体绘制  
4.1 二维切片的局限  
4.2 百分位归一化与显示范围控制  
4.3 自定义 alpha compositing 体绘制  
4.4 自定义 LUT 与 smoothstep 透明度函数  
4.5 梯度增强与类光照  
4.6 三种传递函数设计  
4.7 本阶段小结：从局部截面升级到三维整体结构表达  

5 结构验证分析：从“看起来像”到“统计阈值验证”  
5.1 体绘制结果的不足  
5.2 P99 top 1% 高密度区域筛选  
5.3 X/Y/Z 最大强度投影  
5.4 P90/P95/P99 多阈值等值面  
5.5 结构指标：P99-P01、P99/mean、熵、Gini、空洞比例和致密比例  
5.6 首末时间步差分 MIP  
5.7 本阶段小结：将统计高尾部与三维空间节点对应起来  

6 交互式可视分析：从静态图片到 linked brushing  
6.1 静态图像分析的局限  
6.2 linked brushing 的设计思路  
6.3 直方图刷选与空间投影联动  
6.4 时间步滑条、百分位区间和投影方向切换  
6.5 交互分析案例  
6.6 本阶段小结：从固定结果展示升级到可探索分析  

7 进阶可视分析：从空间观察到结构识别和阶段划分  
7.1 为什么需要进阶分析  
7.2 density-gradient 二维相空间  
7.3 基于 Hessian 特征值的 Void / Sheet / Filament / Node 分类  
7.4 时间步相似性矩阵与演化阶段划分  
7.5 进阶模块与原有方法的关系  
7.6 本阶段小结：从可视化展示升级到结构识别  
7.7 HTML Web 交互展示系统  

8 最终结果汇总与综合分析  
8.1 数据读取检查结果  
8.2 体绘制关键帧结果  
8.3 密度直方图与时间-密度热力图结果  
8.4 时间序列统计曲线结果  
8.5 P99 高密度筛选结果  
8.6 多阈值等值面与差分 MIP 结果  
8.7 交互仪表盘结果  
8.8 density-gradient 二维相空间结果  
8.9 Hessian 宇宙网分类结果  
8.10 时间步相似性与阶段划分结果  
8.11 Web 交互展示作品结果  
8.12 综合规律总结  

9 实验过程反思与方法演进总结  
9.1 从无到有的开发路线  
9.2 每一次改进解决的问题  
9.3 当前系统的优势  
9.4 当前系统的局限  
9.5 后续可扩展方向  

10 结论  
11 附录：核心代码

# 1 问题描述

## 1.1 研究背景

宇宙大尺度结构描述了物质在宇宙空间中的整体分布形态。现代宇宙学认为，早期宇宙中的微小密度扰动会在引力作用下逐渐放大，形成空洞、片状结构、丝状结构以及高密度团块和节点。这些结构共同构成通常所说的“宇宙网”。对于此类数据，仅凭单张二维图像很难理解其空间结构和时间演化，因此科学可视化方法在数据理解中具有重要意义。

Nyx 宇宙学模拟输出的三维密度场为课程大作业提供了典型的科学体数据场景。它既有三维空间维度，又有时间序列维度，同时存在较大的数值动态范围。与普通二维图像不同，Nyx 数据需要先从二进制文件中恢复出三维体结构，再设计合适的统计图、体绘制、阈值筛选和交互刷选方法。

## 1.2 数据特点

本作业使用的 Nyx 数据由 `0000.dat`、`0001.dat`、`0002.dat` 直到 `0099.dat` 等文件组成。每个文件代表一个时间步，每个时间步内部存储一个 n×n×n 三维密度体。文件格式为 little-endian float32，原始线性顺序为 z 轴最快、y 轴其次、x 轴最慢。每个体素表示宇宙空间中对应位置的 density 值。

这类数据具有三个直接挑战。第一，`.dat` 文件本身不可直接查看，必须通过二进制解析恢复数据结构。第二，density 的动态范围较大，直接线性显示容易让高密度区域主导图像，从而掩盖中低密度结构。第三，数据包含多个时间步，单帧观察无法体现演化过程，因此需要将空间结构和时序统计结合起来。

## 1.3 本作业的可视化目标

本作业的目标不是进行严格天体物理参数推断，而是构建一套适合数据可视化课程展示的 Nyx 密度场分析流程。具体目标包括：正确读取并恢复三维密度体；通过切片和文件大小检查验证读取结果；通过统计曲线和直方图观察全时间步密度分布变化；通过体绘制展示三维宇宙网结构；通过 P99 高密度筛选、多阈值等值面和结构指标验证高密度尾部与空间节点的关系；通过 linked brushing 建立统计分布和空间投影之间的联动；最后通过 density-gradient 相空间、Hessian 分类和时间步相似性矩阵形成进阶结构分析。

## 1.4 实验迭代思路

在项目初始阶段，`.dat` 文件本身无法直接查看，因此首要任务不是绘图，而是恢复数据结构。整个项目按照“问题驱动—方法改进—结果验证—继续优化”的方式推进：先解决数据能否读出，再验证读出的数据是否可信；先观察单帧，再扩展到全时间步；先做二维切片，再进入三维体绘制；先做视觉展示，再用统计阈值和结构指标验证；先生成静态图，再加入交互探索；最后引入进阶相空间和结构识别方法。

![图 1：Nyx 可视化项目从无到有的实验迭代路线图](../results/report_figures/iteration_roadmap.png)

图 1 展示了本作业从数据理解到最终系统形成的迭代路线。该路线图强调：每个模块都不是孤立添加的功能，而是针对前一阶段暴露出的不足所做的改进。

![图 2：Nyx 宇宙密度演化可视分析总体流程图](../results/report_figures/overall_workflow.png)

图 2 从系统角度概括了项目最终形成的完整流程。与图 1 的“实验迭代路线”不同，图 2 更强调最终系统中数据读取、统计分析、体绘制、高密度筛选、结构指标和交互分析之间的功能关系。

![图 3：项目关键词词云](../results/report_figures/nyx_keyword_wordcloud.png)

![图 4：方法亮点气泡标签图](../results/report_figures/nyx_method_bubble_tags.png)

图 3 和图 4 作为项目概览图，用于展示本作业涉及的核心概念和方法体系。它们说明后续章节并非孤立技术罗列，而是围绕 Nyx density、log-density、volume rendering、P99、linked brushing 和进阶结构识别逐步展开。

# 2 数据读取与初步验证：从不可见二进制到三维密度体

## 2.1 原始 .dat 数据格式分析

**问题。** 在最初阶段，Nyx 的 `.dat` 文件无法直接打开查看，也没有表头、字段名或图像元数据。如果不先理解数据格式，后续任何图像都可能建立在错误的维度或错误的轴顺序上。

**方法。** 本作业首先根据文件大小分析数据格式。由于每个 density 值为 float32，占 4 字节，因此可通过文件总字节数除以 4 得到体素数量，再判断该数量是否能组成 n³ 的立方体网格。例如 8,388,608 字节对应 2,097,152 个 float32 数值，即 128×128×128 体素。

**结果。** 文件大小分析说明，当前数据可以被解释为规则三维体数据，而不是任意长度的一维数组。这为后续自动推断网格尺寸提供了依据。

**不足。** 文件大小只能说明数据数量正确，不能说明读取字节序、reshape 顺序和坐标轴方向正确。

**下一步。** 因此需要进一步使用 little-endian float32 方式读取数据，并恢复 z-y-x 到 x-y-z 的轴顺序。

## 2.2 little-endian float32 读取

**问题。** 如果字节序或数据类型设置错误，读出的数值会出现异常，例如极端大值、NaN 或不合理分布。

**方法。** `read_nyx_dat.m` 使用 `fopen(filename,'r','ieee-le')` 指定 little-endian 字节序，并使用 `fread(fid, inf, 'single=>single')` 读取 float32 数据。读取后保留 single 类型，既符合原始数据精度，也能降低内存占用。

**结果。** 程序可以稳定读取单个 `.dat` 文件，并获得一维体素数组。

**不足。** 一维数组仍不能直接作为三维密度场使用。

**下一步。** 需要根据体素数量自动推断 n，并 reshape 为三维体数据。

## 2.3 网格尺寸自动推断

**问题。** 数据文件可能是 128³，也可能是其他 n³ 尺寸。如果在代码中写死 128，会降低可复现性和鲁棒性。

**方法。** `infer_grid_size.m` 根据 `numel(data)` 计算立方根并取整，检查 `n^3` 是否等于体素总数。若不相等，则说明文件尺寸异常或数据格式不符合假设，程序会报错提示。

**结果。** 读取函数不依赖硬编码网格尺寸，可自动适配符合 n×n×n 的 Nyx 体数据。

**不足。** 自动推断只能确认体素数量能构成立方体，仍不能确认原始线性顺序如何映射到三维坐标。

**下一步。** 需要按照题目给出的 z-y-x 存储顺序进行轴恢复。

## 2.4 z-y-x 到 x-y-z 的轴顺序恢复

**问题。** 原始线性存储顺序为 z 轴最快、y 轴其次、x 轴最慢。如果直接按 MATLAB 默认理解，很可能导致空间方向错乱。

**方法。** 程序先将一维数组 `reshape(data,[nz,ny,nx])`，得到 `[z,y,x]` 顺序的数组，再通过 `permute(Vzyx,[3,2,1])` 转换为 `[x,y,z]`。这样后续切片、MIP、体绘制和交互投影都建立在一致的坐标解释之上。

**结果。** 单个 `.dat` 文件被恢复为 MATLAB 中更直观的三维 density 体数据。

**不足。** 代码逻辑正确并不等于数据方向已经实际验证正确。

**下一步。** 需要通过中心切片和文件大小检查进行读取验证。

## 2.5 中心切片验证与文件大小检查

**问题。** 虽然成功读取了二进制数据，但仍无法确认 reshape 和 permute 顺序是否正确。

**方法。** `step1_check_data.m` 优先读取 `0000.dat`，输出尺寸、最小值、最大值、均值、标准差、NaN 数量、Inf 数量和负值数量。同时提取 XY、XZ、YZ 三个中心切片，并使用 `log10(max(V,eps))` 进行显示。所有 `.dat` 文件还会进行文件大小一致性检查。

**结果。** 三个方向的中心切片呈现连续纹理，没有明显条纹或错位时，可说明数据读取和轴顺序恢复基本正确。文件大小一致性图则用于检查 100 个时间步是否具有相同网格规模。

![图 5：数据读取检查中心切片图](../results/01_data_check/center_slices_0000.png)

![图 6：文件大小一致性检查图](../results/01_data_check/file_size_check.png)

**不足。** 中心切片验证解决了数据读取是否正确的问题，但仍然只能观察局部截面。

**下一步。** 因此需要进一步从单帧观察扩展到全时间步统计，并在后续引入 MIP 和体绘制观察整体空间结构。

## 2.6 本阶段小结：解决“数据能否正确读出”的问题

本阶段完成了从不可见二进制文件到三维密度体的恢复。通过 little-endian float32 读取、网格尺寸自动推断、轴顺序重排、中心切片和文件大小检查，项目解决了“数据能否正确读出”的基础问题。此时仍存在两个不足：一是只能观察局部切片，二是尚未分析时间演化。因此下一阶段转向全时间步统计。

# 3 基础统计分析：从单帧观察到全时间步量化

## 3.1 为什么不能只看单个时间步

**问题。** 单个时间步只能说明某一时刻的密度分布状态，无法回答密度场是否随时间变得更不均匀、高密度尾部是否增强、低密度区域是否扩张等演化问题。

**方法。** 本作业将分析对象从单帧扩展到全部 100 个时间步。`step2_statistics_all_frames.m` 遍历 `data/raw/` 下所有非空 `.dat` 文件，并按文件名排序，逐帧读取三维体数据。

**结果。** 全时间步统计使分析对象从单个时间步扩展到完整演化过程。

**不足。** 仅遍历数据还不够，需要选择能够描述密度分布变化的统计量。

**下一步。** 因此设计 mean、std、max、百分位数和 log-density 统计指标。

## 3.2 全时间步统计量设计

**问题。** Nyx density 分布具有长尾特征，单纯使用均值很难反映局部高密度结构增强。

**方法。** 对每个时间步计算 mean、std、min、max、P01、P05、P50、P95、P99、P99.7 和 P99.9。同时对 log-density 计算均值、标准差、偏度和峰度。

**结果。** 程序输出 `density_stats.csv`、`statistics.mat` 和 `high_density_threshold.csv`，为后续曲线绘制和 P99 筛选提供数据基础。

**不足。** 统计表本身不直观，需要通过曲线和热力图展示趋势。

**下一步。** 因此绘制统计曲线、直方图和 time-density heatmap。

## 3.3 mean、std、max 和百分位数

**问题。** mean density 可能变化不明显，不能单独作为判断结构演化的依据。

**方法。** 本作业同时绘制 mean±std、max density、P01/P05/P50/P95/P99 等百分位曲线。std、max、P95、P99、P99/mean 和 P99-P01 更适合描述密度波动增强和高密度尾部变化。

**结果。** 统计曲线可展示密度分布是否逐渐扩散、极端高密度是否增强。

**不足。** 原始 density 动态范围大，直接显示和直接比较容易被极端值影响。

**下一步。** 因此引入 log-density 变换。

## 3.4 log-density 变换的引入

**问题。** log-density 变换的引入，是因为原始 density 动态范围较大，直接显示会掩盖中低密度结构。

**方法。** 对 density 使用 `log10(max(density, eps))`。其中 `eps` 用于避免对 0 或极小值取对数导致数值异常。

**结果。** log-density 压缩了动态范围，使低密度空洞、中密度丝状结构和高密度节点能够在同一幅图中更均衡地显示。

**不足。** log-density 只改变数值尺度，仍不能单独表达空间位置。

**下一步。** 因此需要将 log-density 用于直方图、热力图、MIP 和体绘制。

## 3.5 对数密度直方图

**问题。** 统计曲线描述的是若干摘要指标，不能展示完整分布形态。

**方法。** `step3_histogram_analysis.m` 选取 t=0000、0030、0060、0099 四个代表时间步，使用 120 个 bins 绘制 log-density probability histogram，并生成四帧叠加对比图。

**结果。** 直方图可以观察分布中心、宽度和右尾变化。早期分布通常更集中，后期高密度尾部更明显。

![图 7：代表时间步对数密度直方图对比图](../results/02_histograms/histogram_compare_keyframes.png)

**不足。** 直方图只表达数值分布，不包含空间位置。

**下一步。** 因此需要 time-density heatmap 描述全时序分布变化，并在后续用 P99 mask 和 linked brushing 建立空间对应。

## 3.6 时间-密度热力图

**问题。** 四个代表时间步仍然是抽样观察，不能完整展示 100 个时间步的分布演化。

**方法。** 对全部时间步使用统一 bins 计算 log-density histogram，将所有 histogram 堆叠为二维热力图。横轴为 time step，纵轴为 log10(density)，颜色为 probability。

**结果。** time-density heatmap 可以展示密度分布随时间整体迁移、扩散和高密度尾部增强。

![图 8：时间-密度热力图](../results/02_histograms/time_density_heatmap.png)

**不足。** 热力图仍然是统计空间，无法说明某个密度区间在三维空间中的位置。

**下一步。** 因此需要进入三维空间可视化。

## 3.7 本阶段小结：从图像观察扩展到时序统计

本阶段完成了从单帧切片观察到全时间步统计分析的转变。统计曲线、直方图和 time-density heatmap 揭示了密度分布随时间变化的趋势，但它们仍然缺少三维空间形态。因此下一阶段需要用 MIP 和体绘制表达空间结构。

# 4 三维空间可视化：从二维切片到体绘制

## 4.1 二维切片的局限

**问题。** 中心切片只能观察穿过体数据中心的一张截面，可能错过空间中的丝状结构、团块和节点。二维切片无法完整表达三维宇宙网的整体形态。

**方法。** 在切片之外，本作业加入最大强度投影和自定义体绘制。MIP 用于快速观察高密度结构的大致分布，体绘制用于表现三维密度场的整体空间感。

**结果。** 体绘制解决了二维切片空间信息不足的问题，但体绘制本身仍然需要统计结果支撑。

**不足。** 体绘制对显示范围和传递函数非常敏感。

**下一步。** 因此需要百分位归一化和可控传递函数。

## 4.2 百分位归一化与显示范围控制

**问题。** 即使采用 log-density，极端高值仍可能影响显示范围，导致图像对比度不稳定。

**方法。** `normalize_percentile.m` 根据给定 lowerPct 和 upperPct 计算阈值，将数据裁剪到该范围后归一化到 [0,1]。体绘制默认使用 log-density 的 P5 到 P99.7。

**结果。** 百分位归一化提高了不同时间步之间的显示一致性，也避免少数极端值压缩主要结构的可见范围。

**不足。** 归一化只解决显示范围，还需要决定不同密度值对应什么颜色和透明度。

**下一步。** 因此设计自定义 LUT 和透明度传递函数。

## 4.3 自定义 alpha compositing 体绘制

**问题。** 直接依赖 `volshow` 不利于展示算法细节，也不便于控制传递函数。

**方法。** `volume_render_alpha_composite.m` 实现沿 z 方向从后向前的 alpha compositing。对每一层，根据归一化密度查找颜色和透明度，再与已累积图像进行前向合成。

**结果。** 项目获得了不依赖 `volshow` 的自定义体绘制流程，能够在报告中解释体绘制的核心机制。

**不足。** 固定 z 方向合成缺少自由旋转，且视觉效果依赖传递函数。

**下一步。** 因此需要更明确地设计 LUT、透明度函数和梯度增强。

## 4.4 自定义 LUT 与 smoothstep 透明度函数

**问题。** 线性透明度函数往往无法同时展示低密度背景、中密度丝状结构和高密度节点。

**方法。** 本作业设计从深蓝/黑色到紫色、橙色、黄色和白色的自定义 LUT，并使用 `smoothstep` 控制透明度过渡。低密度区域接近透明，中密度区域半透明，高密度区域更不透明。

**结果。** 自定义 LUT 与 smoothstep 使不同密度层次的视觉表达更加可控。

**不足。** 仅依据密度值仍可能让边界不够清晰。

**下一步。** 因此加入梯度增强与类光照。

## 4.5 梯度增强与类光照

**问题。** 宇宙网的丝状结构和团块边界通常体现在密度变化较快的位置，单纯密度映射可能使边界显得平。

**方法。** 对归一化 log-density 计算三维梯度幅值，并用梯度增强透明度或亮度。该处理不是物理真实光照，而是用于突出密度边界、丝状结构和团块边缘的类光照增强。

**结果。** 梯度增强使结构边界更清晰，有助于报告展示。

**不足。** 视觉增强可能改变观察重点，因此需要通过统计阈值验证结构意义。

**下一步。** 在体绘制基础上进一步设计不同传递函数，并与 P99 筛选结合验证。

## 4.6 三种传递函数设计

**问题。** 不同分析任务关注的密度范围不同，单一传递函数无法同时适合观察空洞、丝状结构和高密度节点。

**方法。** 项目设计三类传递函数：空洞观察型强调低密度区域，丝状结构型增强中密度连续结构，高密度节点型突出高密度尾部。传递函数不是单纯美化图像，而是根据不同分析任务设计视觉映射。

![图 9：三种体绘制传递函数设计对比](../results/report_figures/transfer_function_design.png)

**结果。** 三种传递函数使同一体数据能够从不同角度观察结构特征。

**不足。** 传递函数选择具有主观性，仍需要统计图和阈值结果支撑。

**下一步。** 因此进入结构验证分析。

## 4.7 本阶段小结：从局部截面升级到三维整体结构表达

本阶段从二维切片升级到三维体绘制，实现了 Nyx 密度场整体空间结构的表达。体绘制能够直观展示类似宇宙网的丝状结构和节点，但它仍然属于视觉展示，需要通过 P99 筛选、等值面和结构指标进一步验证。

# 5 结构验证分析：从“看起来像”到“统计阈值验证”

## 5.1 体绘制结果的不足

**问题。** 体绘制能够让结构“看起来像”宇宙网，但如果缺少统计依据，就难以说明高亮区域是否真的对应密度分布中的高尾部。

**方法。** 本作业引入 P99 top 1% 高密度筛选、多阈值等值面和结构指标，将视觉观察与统计阈值联系起来。

**结果。** 结构验证分析使体绘制结果从视觉展示进一步转向可解释分析。

**不足。** 阈值选择仍具有经验性。

**下一步。** 因此需要同时观察 P99、P90/P95/P99 层级和连续结构指标。

## 5.2 P99 top 1% 高密度区域筛选

**问题。** 直方图右尾表示高密度体素，但单凭直方图无法知道这些体素位于三维空间哪里。

**方法。** 对每个时间步使用原始 density 的 P99 作为阈值，生成 `density >= P99` 的 top 1% mask。P99 是阈值，top 1% 是超过该阈值的空间体素集合。

**结果。** P99 高密度筛选将直方图右尾与三维空间节点联系起来，使可视化结果具有统计依据。

**不足。** P99 是人为阈值，不同阈值会改变选区范围。

**下一步。** 因此进一步绘制 MIP 和多阈值等值面。

## 5.3 X/Y/Z 最大强度投影

**问题。** 三维 mask 不方便直接放入报告，需要一种二维表达方式展示高密度区域的空间分布。

**方法。** 对 P99 mask 区域中的 log-density 分别沿 X、Y、Z 三个方向做最大强度投影。

**结果。** MIP 显示 P99 以上体素并非随机分布，而是更集中于丝状结构交汇处和高密度节点附近。

**不足。** MIP 会丢失深度信息，不能完全替代真实三维交互。

**下一步。** 因此引入等值面和嵌套等值面观察层级结构。

## 5.4 P90/P95/P99 多阈值等值面

**问题。** 单一 P99 阈值只能展示最极端高密度区域，无法表现中高密度外壳到高密度核心的层级。

**方法。** 对 t0099 或最后一个时间步绘制 P90、P95、P99 三个阈值的嵌套等值面。P90 表示中高密度丝状区域，P95 表示高密度聚集区域，P99 表示极高密度节点。

**结果。** 多阈值等值面体现了密度结构从大尺度丝状外壳向高密度核心收缩的层级关系。

**不足。** 等值面仍受阈值选择影响，且对复杂重叠结构的深度关系表达有限。

**下一步。** 因此需要补充结构指标进行量化分析。

## 5.5 结构指标：P99-P01、P99/mean、熵、Gini、空洞比例和致密比例

**问题。** 图像展示容易受到视觉映射影响，需要更客观的数值指标描述结构演化。

**方法。** `step6_structure_metrics.m` 计算 P99-P01、P99/mean、void fraction、dense fraction、log-density entropy、Gini 系数，以及可选的 P99 mask 连通域指标。

**结果。** 这些指标从不同角度量化密度分布的扩散、高尾部增强和不均匀程度。

![图 10：密度统计曲线](../results/report_figures/combined_statistics_curves.png)

**不足。** 指标能够描述总体趋势，但仍不能直接展示空间中哪些区域增强或减弱。

**下一步。** 因此加入首末时间步差分 MIP。

## 5.6 首末时间步差分 MIP

**问题。** 如果只看单个时间步，难以观察哪些空间区域在演化中密度增强或降低。

**方法。** 计算最后一个时间步与第一个时间步的 log-density 差分，并使用红蓝发散色图做三方向差分 MIP。

**结果。** 红色区域表示相对增强，蓝色区域表示相对降低。该图为“低密度空洞扩展”和“高密度结构增强”提供空间线索。

**不足。** 差分 MIP 仍然压缩了深度信息。

**下一步。** 因此需要交互式分析，让用户主动探索不同密度区间和空间投影。

## 5.7 本阶段小结：将统计高尾部与三维空间节点对应起来

本阶段通过 P99 筛选、多阈值等值面、结构指标和差分 MIP，将统计分布中的高密度尾部与三维空间中的节点、团块和丝状交汇处联系起来。它解决了体绘制缺少统计依据的问题，但静态图像仍缺少灵活探索能力。

# 6 交互式可视分析：从静态图片到 linked brushing

## 6.1 静态图像分析的局限

**问题。** 静态图只能展示预先选择的时间步、阈值和视角。用户如果想查看其他密度区间或其他时间步，需要重新运行脚本。

**方法。** 本作业实现 `step7_interactive_dashboard.m`，使用时间步滑条、密度百分位滑条、投影方向选择和显示模式选择构建交互仪表盘。

**结果。** linked brushing 的加入，使分析从预设结果展示转向用户可探索的交互式分析。

**不足。** 初始仪表盘主要支持一维密度区间刷选。

**下一步。** 因此在进阶模块中扩展到 density-gradient 二维相空间。

## 6.2 linked brushing 的设计思路

**问题。** 统计图和空间图如果彼此独立，用户很难判断某一段直方图对应空间中的哪个区域。

**方法。** 左侧显示当前时间步的 log-density histogram，右侧显示空间投影视图。用户通过百分位滑条选择密度区间后，系统生成 mask 并实时更新空间投影。

**结果。** 该设计建立了统计分布与三维空间结构之间的双向关联。

**不足。** 当前空间视图以投影为主，深度关系仍有压缩。

**下一步。** 可进一步结合三维交互体绘制。

## 6.3 直方图刷选与空间投影联动

**问题。** 单独直方图无法表达空间位置，单独投影无法说明选区的统计意义。

**方法。** 在 histogram 中标出当前百分位区间，并将同一选区映射到空间投影中。

**结果。** 用户可以观察低密度、中密度、高密度区间在三维空间投影中的不同分布。

**不足。** 仅以 density 作为刷选变量，无法区分边界和核心。

**下一步。** density-gradient 相空间将解决这一不足。

## 6.4 时间步滑条、百分位区间和投影方向切换

**问题。** 科学可视分析需要在时间、数值区间和空间方向之间灵活切换。

**方法。** 仪表盘提供 time slider、density lower/upper percentile slider、X/Y/Z 投影方向和 MIP/Mask/LogDensity 显示模式。

**结果。** 用户可以快速比较不同时间步和不同密度区间的空间分布。

**不足。** 交互结果依赖当前投影方式。

**下一步。** 后续可扩展为三维旋转或体绘制交互。

## 6.5 交互分析案例

**问题。** 需要展示交互分析如何帮助解释数据，而不是只展示界面。

**方法。** 默认状态下显示某一时间步的 histogram 和空间投影，调整百分位区间观察低密度背景、中密度丝状结构和高密度节点的分布差异。

![图 11：交互式刷选仪表盘截图](../results/07_dashboard/dashboard_preview.png)

**结果。** 仪表盘能够回答“某一密度区间在哪里”这一问题。

**不足。** 一维密度刷选仍无法表达局部变化程度。

**下一步。** 因此引入 density-gradient 二维相空间。

## 6.6 本阶段小结：从固定结果展示升级到可探索分析

本阶段将静态可视化结果扩展为 linked brushing 交互仪表盘。它强化了统计分布和空间投影之间的联系，但仍主要基于一维 density 特征。为了进一步识别边界、核心和结构类型，下一阶段加入进阶分析模块。

# 7 进阶可视分析：从空间观察到结构识别和阶段划分

## 7.1 为什么需要进阶分析

**问题。** 基础流程已经能够读取数据、统计趋势、体绘制和交互刷选，但仍有三个不足：一维密度刷选不能区分结构边界和核心；视觉观察难以明确区分 Void、Sheet、Filament 和 Node；代表时间步选择缺少数据依据。

**方法。** 在原有流程上新增 density-gradient 二维相空间、Hessian 特征值结构分类和时间步相似性矩阵。

**结果。** 进阶模块使分析从空间观察进一步走向结构识别和阶段划分。

**不足。** 这些方法仍依赖阈值、平滑尺度和统计特征设计。

**下一步。** 在综合分析中将进阶结果与基础图像和统计结果互相印证。

## 7.2 density-gradient 二维相空间

**问题。** 一维密度刷选只能回答“密度高不高”，不能区分“高密度边界”和“高密度核心”。

**方法。** 对 normalized log-density 计算 gradient magnitude，将每个体素表示为二维点 `(density, gradient)`。横轴为 normalized log-density，纵轴为 normalized gradient magnitude，颜色为 `log10(voxel count + 1)`。

**结果。** density-gradient 二维相空间进一步把一维密度刷选扩展为二维特征刷选。低密度低梯度对应 void background，中密度高梯度对应 filament boundary，高密度高梯度对应 node boundary，高密度低梯度对应 dense core。

**不足。** 当前二维相空间选择主要输出 MIP，仍不能完全替代三维交互探索。

**下一步。** 可将二维相空间刷选并入 GUI，实现更完整的交互式二维传递函数。

## 7.3 基于 Hessian 特征值的 Void / Sheet / Filament / Node 分类

**问题。** 仅靠体绘制和 MIP 观察，很难稳定地区分不同局部结构类型。

**方法。** 对平滑后的 log-density 计算 Hessian 特征值。设阈值 `tau`，统计小于 `-tau` 的特征值数量。0、1、2、3 个负曲率方向分别对应 Void、Sheet、Filament 和 Node。

**结果。** Hessian 分类使宇宙网结构从视觉观察进一步扩展为局部结构识别。

**不足。** 分类依赖平滑尺度、降采样尺度和阈值 `tau`，因此只能作为课程可视化实验中的近似结构识别。

**下一步。** 可与更多物理变量或多尺度 Hessian 分析结合。

## 7.4 时间步相似性矩阵与演化阶段划分

**问题。** t=0000、0030、0060、0099 作为代表时间步虽然直观，但需要数据依据支撑。

**方法。** 将每个时间步的 log-density histogram 看作概率向量，计算时间步之间的 cosine similarity。再根据 distance from start 的分位数划分 Early、Transition、Structure enhancement 和 Late 四个阶段。

**结果。** 时间步相似性矩阵为代表帧选择和演化阶段划分提供了数据依据。

**不足。** 相似性矩阵基于统计分布，不包含完整空间对应关系。

**下一步。** 可进一步结合空间特征向量或结构指标进行阶段划分。

## 7.5 进阶模块与原有方法的关系

density-gradient 相空间是对 linked brushing 的扩展，Hessian 分类是对体绘制和等值面观察的结构化补充，时间步相似性矩阵是对 time-density heatmap 的时序组织补充。三者并不是替代原有方法，而是针对原有方法的不足继续深化。

## 7.6 本阶段小结：从可视化展示升级到结构识别

本阶段将项目从“展示 Nyx 密度场”推进到“识别和组织 Nyx 密度结构”。虽然这些进阶分析仍具有课程实验性质，但它们使最终系统在方法链条上更完整。

## 7.7 HTML Web 交互展示系统

**问题。** MATLAB 交互仪表盘适合本地分析，但在课程答辩和作品展示时仍存在两个不足：一是需要 MATLAB 环境支持，展示迁移性较弱；二是静态报告图虽然完整，却不能让观众现场切换时间步、刷选密度区间和观察空间联动。因此，在完成数据读取、统计分析、体绘制、结构验证和进阶分析之后，项目进一步增加 HTML 单页 Web 交互展示系统，将前面生成的数据和结果整理为可浏览器演示的作品。

**方法。** Web 系统采用“离线预处理 + 浏览器懒加载”的策略，而不是直接把 100 个完整 128×128×128 体数据一次性加载进浏览器。预处理脚本 `preprocess_for_web.py` 从 `data/raw/` 读取 `.dat` 文件，完成 little-endian float32 解析、z-y-x 到 x-y-z 轴顺序恢复、log-density 变换、P5-P99.7 百分位归一化，并将每个时间步降采样为默认 64×64×64 的 Float32Array 二进制文件 `vol_0000.bin` 至 `vol_0099.bin`。同时，脚本生成 `metadata.json`、`density_stats.json`、`histograms.json`、`time_density_heatmap.json` 和阶段划分 JSON，供浏览器端按需读取。

**结果。** Web 页面 `index.html` 使用原生 HTML/CSS/JavaScript 和 Canvas 实现，不依赖外部可视化库。页面包含时间步滑条、播放按钮、密度百分位区间刷选、投影方向切换、MIP/Mask/Slice/简化 Volume Composite 视图、log-density histogram、time-density heatmap、统计指标曲线和结论卡片。浏览器端只加载当前时间步体数据，并将已经访问过的时间步缓存在内存中，从而保证演示时能够流畅切换。

**与原有方法的关系。** Web 版本不是替代 MATLAB/Python 分析流程，而是最终展示层：MATLAB/Python 负责数据恢复、统计计算、体绘制和结构分析，Web 负责将这些结果组织成可演示、可交互、易传播的可视化作品。它把 linked brushing 从 MATLAB GUI 延伸到浏览器环境，使答辩现场可以直接展示“统计直方图刷选—空间投影响应—时间演化播放”的完整交互链条。

![图 12：Nyx Web 交互展示系统概览图](../results/report_figures/web_system_overview.png)

**不足。** 当前 Web 端体绘制为 CPU Canvas 上的简化 alpha compositing，交互效率优先于真实三维体渲染；空间视图以 MIP、Mask、Slice 和固定方向合成为主，尚未实现 WebGL 自由旋转。后续可以引入 Three.js 或 WebGL shader，实现真正的 3D volume rendering 和二维传递函数刷选。

# 8 最终结果汇总与综合分析

## 8.1 数据读取检查结果

数据读取检查结果如图 5 和图 6 所示。中心切片呈现连续纹理，文件大小检查显示时间步文件规模一致，说明 little-endian 读取、网格尺寸推断和轴顺序恢复基本可靠。该结果为后续全部可视化分析提供了前提。

## 8.2 体绘制关键帧结果

图 13 展示 t=0000、0030、0060、0099 的体绘制对比。t=0000 中密度结构相对均匀；t=0030 开始出现局部聚集；t=0060 丝状结构增强；t=0099 高密度节点和宇宙网形态更加明显。体绘制适合观察空间整体形态，但具体密度变化仍需统计曲线和阈值结果支撑。

![图 13：代表时间步体绘制对比图](../results/04_volume_render/volume_keyframes_compare.png)

## 8.3 密度直方图与时间-密度热力图结果

图 7 和图 8 表明，log-density 分布随时间推进出现扩散和高密度尾部增强。直方图有助于观察单帧分布形态，time-density heatmap 则展示完整时序趋势。它们共同说明，密度场从较集中的数值分布逐渐发展出更明显的不均匀性。

## 8.4 时间序列统计曲线结果

统计曲线如图 10 所示。mean density 可能变化不明显，因此不应作为唯一依据；std、max、P95、P99、P99/mean 和 P99-P01 更能反映密度波动增强和高密度尾部拉长。偏度和峰度可补充分布不对称性和极端尾部的信息。

## 8.5 P99 高密度筛选结果

图 14 展示 P99 top 1% 高密度筛选结果。P99 以上体素并非随机散布，而是集中于丝状结构交汇处和高密度节点附近。这说明直方图右尾具有明确空间结构意义。

![图 14：Top 1% 高密度筛选图](../results/report_figures/combined_top1_percent_mips.png)

## 8.6 多阈值等值面与差分 MIP 结果

图 15 展示 P90、P95、P99 多阈值等值面。P90 通常覆盖更大范围的中高密度丝状区域，P95 更集中，P99 只保留最致密节点。图 16 的首末时间步差分 MIP 则显示密度增强和降低区域，为结构演化提供空间对照。

![图 15：多阈值等值面图](../results/05_high_density/nested_isosurfaces_t0099.png)

![图 16：首末时间步差分 MIP](../results/06_structure_metrics/difference_mip_first_last.png)

## 8.7 交互仪表盘结果

图 11 展示 linked brushing dashboard。用户可以通过滑条选择时间步和密度百分位区间，并查看对应空间投影。该界面把“统计区间是什么”和“空间位置在哪里”联系起来，使分析从固定结果展示转向探索式分析。

## 8.8 density-gradient 二维相空间结果

图 17 展示四个代表时间步的 density-gradient 二维相空间。高梯度区域对应结构边界，高密度低梯度区域更可能对应致密核心。图 18 进一步展示 t0099 中四类典型相空间区域的空间选择结果，说明二维相空间能够比单变量 density 更细致地区分空间结构类型。

![图 17：density-gradient 二维相空间对比图](../results/report_figures/density_gradient_phase_compare.png)

![图 18：density-gradient 典型区域空间选择图](../results/report_figures/phase_space_selection_t0099.png)

## 8.9 Hessian 宇宙网分类结果

图 19 展示 Hessian 分类切片，图 20 展示 Void、Sheet、Filament 和 Node 占比曲线。该方法不只是显示高密度区域，而是尝试根据局部二阶结构识别不同宇宙网形态。Filament+Node MIP 进一步展示中高密度连续结构和节点区域的空间分布。

![图 19：Hessian 宇宙网分类切片对比图](../results/report_figures/hessian_class_slice_compare.png)

![图 20：Hessian 结构占比曲线](../results/report_figures/hessian_class_fraction_curve.png)

## 8.10 时间步相似性与阶段划分结果

图 21 展示时间步相似性矩阵。矩阵对角线附近相似度最高，说明相邻时间步连续变化；远离对角线后相似度降低，说明早期和后期分布差异增加。图 22 根据 distance from start 给出阶段划分，为代表时间步选择提供数据依据。

![图 21：时间步相似性矩阵](../results/report_figures/time_step_similarity_matrix.png)

![图 22：演化阶段划分图](../results/report_figures/evolution_stage_segmentation.png)

## 8.11 Web 交互展示作品结果

图 12 展示了最终 HTML Web 交互展示系统的组织方式。与前面各阶段输出的静态图不同，Web 作品将降采样体数据、统计曲线、直方图、time-density heatmap 和结论卡片组织到同一个浏览器页面中。用户可以拖动时间步滑条或点击播放按钮观察演化过程，也可以将密度百分位区间设置为 99%-100%，实时查看 top 1% 高密度体素在 X/Y/Z 方向投影中的空间分布。该结果对应前面“交互分析”阶段的进一步产品化：它不仅证明 linked brushing 方法能够工作，也使课程答辩中可以现场演示“选择密度右尾—观察空间节点—切换时间步验证演化”的完整分析路径。

从结果解释角度看，Web 作品强化了本项目的展示闭环。统计图回答“密度分布如何变化”，空间视图回答“这些密度区间在哪里”，time-density heatmap 和指标曲线回答“变化是否具有时序趋势”，而浏览器交互把这些回答放在同一操作界面中。其局限在于 Web 端为了保证流畅性使用了 64×64×64 降采样体数据和简化体合成，因此更适合作为课程展示和探索入口，而不是替代原始数据上的完整离线分析。

## 8.12 综合规律总结

综合统计图、体绘制、P99 筛选、交互刷选、进阶分析和 Web 展示系统可以得到以下课程实验层面的观察：早期密度分布相对集中，后期密度波动逐渐增强；高密度尾部拉长，P99、max 和结构指标更能体现团块化趋势；低密度区域和高密度区域逐渐分化，形成空洞、丝状结构和节点共存的空间形态；体绘制适合观察空间结构，直方图和统计曲线适合观察数值分布，linked brushing 与 density-gradient 相空间适合建立二者之间的联系；HTML Web 页面则使这些分析结果能够在浏览器中以可播放、可刷选、可切换的方式完成最终展示。

# 9 实验过程反思与方法演进总结

## 9.1 从无到有的开发路线

本作业不是一次性完成完整系统，而是围绕数据理解和可视化问题逐步改进。最初面对的是不可直接查看的 `.dat` 文件，因此首先恢复数据结构；随后通过中心切片验证读取结果；之后扩展到全时间步统计；再加入 log-density、体绘制、P99 筛选、交互仪表盘和进阶结构识别；最后将离线分析结果整理为 HTML 单页交互展示作品。最终形成“数据恢复—统计分析—三维渲染—结构验证—交互探索—进阶识别—Web 展示”的完整链条。

## 9.2 每一次改进解决的问题

表 1 总结了从基础版本到最终系统的方法迭代关系。

表 1：从基础版本到最终系统的方法迭代对照表

| 阶段 | 最初问题 | 采用方法 | 得到结果 | 暴露不足 | 下一步改进 |
|---|---|---|---|---|---|
| 阶段 1 数据读取 | `.dat` 文件无法直接查看 | little-endian float32 读取、网格推断 | 恢复 n×n×n 三维密度体 | 无法确认方向是否正确 | 中心切片验证 |
| 阶段 2 中心切片 | 无法确认轴顺序 | XY/XZ/YZ 中心切片 | 验证读取方向基本正确 | 只能观察局部截面 | MIP 和体绘制 |
| 阶段 3 全时间步统计 | 单帧无法体现演化 | 遍历 100 个时间步统计指标 | 生成 density_stats.csv 和统计曲线 | 统计图缺少空间位置 | 体绘制和 P99 空间筛选 |
| 阶段 4 log-density | 原始密度动态范围大 | log10 变换和百分位归一化 | 中低密度结构更清晰 | 仍需要三维表达 | 体绘制 |
| 阶段 5 体绘制 | 二维切片空间信息不足 | alpha compositing、LUT、smoothstep | 展示宇宙网整体结构 | 视觉结果缺少统计验证 | P99 高密度筛选 |
| 阶段 6 P99 筛选 | 高亮区域缺少统计依据 | P99 mask、MIP、多阈值等值面 | 右尾对应空间节点 | 阈值固定、交互不足 | linked brushing |
| 阶段 7 linked brushing | 静态图无法探索 | 直方图滑条与空间投影联动 | 可交互选择密度区间 | 只支持一维密度刷选 | density-gradient 二维相空间 |
| 阶段 8 进阶结构识别 | 仅靠观察难以识别结构类型 | density-gradient、Hessian、time similarity | 实现结构边界分析、Void/Sheet/Filament/Node 分类和阶段划分 | 分类依赖阈值和平滑尺度 | 多变量联合分析和更强交互系统 |
| 阶段 9 Web 交互展示 | MATLAB GUI 和静态报告不便于答辩现场演示 | Python 预处理、降采样体数据、HTML/CSS/JavaScript Canvas 仪表盘 | 实现时间播放、直方图刷选、top 1% 高密度联动展示 | Web 端仍是固定方向投影和简化体合成 | WebGL 体绘制和二维传递函数 |

## 9.3 当前系统的优势

当前系统的优势在于流程完整、层次清晰、结果可复现。数据读取、统计、渲染、阈值验证、交互刷选、进阶分析和 Web 展示构成连续链条；自定义体绘制使传递函数设计可控；log-density 与百分位归一化解决了高动态范围显示问题；P99 筛选将统计尾部与空间节点联系起来；linked brushing 和 density-gradient 相空间增强了探索性分析能力；HTML Web 页面进一步提高了作品的展示性和可传播性，使课程答辩时可以直接演示时间演化和密度刷选联动。

## 9.4 当前系统的局限

当前体绘制主要沿固定 z 方向合成，缺少自由三维旋转；MIP 会丢失深度信息；Hessian 分类依赖平滑尺度和阈值 `tau`；时间步相似性矩阵基于统计分布，不包含完整空间对应关系；Web 端为了保证浏览器流畅性使用降采样体数据和 CPU Canvas 简化渲染，不能完全替代原始分辨率离线分析；当前分析主要基于 density 单变量，尚未结合 velocity、temperature 等多变量数据。因此，本文结论应理解为课程可视化实验分析结果，而不是正式天体物理科学结论。

## 9.5 后续可扩展方向

后续可以从四个方向扩展。第一，加入真正的三维交互视角和二维传递函数刷选，使用户能够同时控制 density 和 gradient。第二，使用多尺度 Hessian 或更稳健的拓扑方法识别宇宙网结构。第三，将 density 与速度、温度等变量联合分析，构建多变量科学可视化系统。第四，在 Web 端引入 WebGL 或 Three.js，将当前 Canvas 投影升级为可旋转的三维体绘制展示。

# 10 结论

本作业围绕 Nyx 宇宙密度演化数据，完成了一套从二进制数据恢复到 HTML 交互展示的可视分析流程。整个项目不是一次性完成完整系统，而是在每一阶段根据前一阶段的不足逐步改进：数据读取解决“看不见”的问题，中心切片解决“读得对不对”的问题，全时间步统计解决“是否存在演化趋势”的问题，log-density 和百分位归一化解决“动态范围过大”的问题，体绘制解决“二维切片空间信息不足”的问题，P99 筛选和结构指标解决“视觉结果缺少统计依据”的问题，linked brushing 解决“静态图无法探索”的问题，density-gradient、Hessian 和时间相似性分析进一步解决“边界、结构类型和阶段划分”问题，Web 交互展示则解决“如何将完整分析流程用于课程答辩现场演示”的问题。最终系统形成了“数据恢复—统计分析—三维渲染—结构验证—交互探索—进阶识别—Web 展示”的完整链条。实验表明，该流程能够较系统地展示 Nyx 密度场从相对均匀到结构增强、从统计尾部到空间节点、从静态展示到交互探索的可视化分析过程。需要强调的是，本作业的结论主要用于数据可视化课程实验，不作为正式天体物理科学结论。

# 11 附录：核心代码

Code Listing 1：Nyx 二进制体数据读取核心代码

对应模块：`read_nyx_dat.m` / `infer_grid_size.m`

```matlab
fid = fopen(filename, 'r', 'ieee-le');
data = fread(fid, inf, 'single=>single');
fclose(fid);
n = infer_grid_size(numel(data));
Vzyx = reshape(data, [n, n, n]);
V = permute(Vzyx, [3, 2, 1]);
```

Code Listing 2：log-density 与百分位归一化核心代码

对应模块：`normalize_percentile.m`

```matlab
logV = log10(max(V, eps('single')));
Vn = normalize_percentile(logV, 5, 99.7);
```

Code Listing 3：P99 高密度筛选核心代码

对应模块：`step5_high_density_selection.m`

```matlab
threshold = prctile(double(V(:)), 99);
mask = V >= threshold;
selectedRatio = nnz(mask) / numel(mask);
```

Code Listing 4：density-gradient 二维相空间核心代码

对应模块：`step8_density_gradient_phase_space.m` / `compute_gradient_magnitude.m`

```matlab
Vn = normalize_percentile(logV, 5, 99.7);
gradMag = compute_gradient_magnitude(Vn);
gradN = normalize_percentile(gradMag, 1, 99);
[counts, xEdges, yEdges] = histcounts2(Vn(:), gradN(:), ...
    linspace(0,1,121), linspace(0,1,121));
```

Code Listing 5：Hessian 分类核心代码

对应模块：`step9_hessian_cosmic_web_classification.m` / `compute_hessian_eigenvalues_3d.m` / `classify_cosmic_web_hessian.m`

```matlab
[l1, l2, l3] = compute_hessian_eigenvalues_3d(logVs);
tau = max(prctile(abs([l1(:); l2(:); l3(:)]), 70), 1e-4);
classVol = classify_cosmic_web_hessian(l1, l2, l3, tau);
```

Code Listing 6：时间步相似性矩阵核心代码

对应模块：`step10_time_similarity_stage_analysis.m` / `cosine_similarity_matrix.m` / `simple_stage_segmentation.m`

```matlab
H(i,:) = histcounts(logV(:), edges, 'Normalization', 'probability');
S = cosine_similarity_matrix(H);
D = 1 - S;
[stageLabel, splitIdx] = simple_stage_segmentation(D(:,1));
```

Code Listing 7：Web 体数据预处理核心代码

对应模块：`Nyx_Web_Visualization/scripts/preprocess_for_web.py`

```python
density = read_nyx_dat(path)
log_v = np.log10(np.maximum(density, EPS)).astype(np.float32)
norm_v, low, high = normalize_percentile(log_v, 5, 99.7)
small = downsample_volume(norm_v, target_size)
small.astype("<f4", copy=False).tofile(volume_dir / f"vol_{t:04d}.bin")
```

Code Listing 8：Web 端懒加载和百分位刷选核心代码

对应模块：`Nyx_Web_Visualization/web/app.js`

```javascript
async function loadVolume(timeIndex) {
  if (state.volumeCache.has(timeIndex)) return state.volumeCache.get(timeIndex);
  const step = state.metadata.steps[timeIndex];
  const res = await fetch(DATA_BASE + step.volume_file);
  const volume = new Float32Array(await res.arrayBuffer());
  state.volumeCache.set(timeIndex, volume);
  return volume;
}

function makeMask(volume, lowPct, highPct) {
  const low = computePercentile(volume, lowPct);
  const high = computePercentile(volume, highPct);
  return volume.map(v => (v >= low && v <= high ? 1 : 0));
}
```
"""


def code_explanation_markdown() -> str:
    return r"""# Nyx 宇宙密度演化可视分析代码说明：迭代过程版

## 1 代码总体结构

本项目采用“主入口 + 分步骤分析模块 + 工具函数 + Web 展示系统”的组织方式。`code/main.m` 负责自动识别项目根目录、添加路径、检查数据文件、创建输出目录，并依次运行基础分析模块和进阶分析模块。每个 `step*.m` 文件对应报告中的一个实验阶段，`code/utils/` 保存通用工具函数。最终展示层位于 `Nyx_Web_Visualization/`，其中 Python 脚本负责生成浏览器可加载的轻量数据，HTML/CSS/JavaScript 负责实现单页交互展示。

## 2 主入口

`code/main.m` 是一键运行入口。基础流程包括：

1. `step1_check_data.m`
2. `step2_statistics_all_frames.m`
3. `step3_histogram_analysis.m`
4. `step4_volume_render_keyframes.m`
5. `step5_high_density_selection.m`
6. `step6_structure_metrics.m`

进阶模块以 `try/catch` 方式追加运行：

1. `step8_density_gradient_phase_space.m`
2. `step9_hessian_cosmic_web_classification.m`
3. `step10_time_similarity_stage_analysis.m`

交互仪表盘 `step7_interactive_dashboard.m` 不默认阻塞运行，可由用户手动调用。

Web 展示系统不由 `main.m` 直接启动，而是通过 `start_nyx_web_demo.bat` 或 `start_nyx_web_demo.py` 一键运行。启动器会检查 `web/assets/data/` 是否存在，必要时自动调用 `scripts/preprocess_for_web.py` 生成 Web 数据，然后启动本地 HTTP 服务器并打开浏览器。

## 3 从无到有的代码演进关系

| 阶段 | 代码文件 | 解决的问题 | 输出结果 | 对应报告章节 |
|---|---|---|---|---|
| 数据读取 | `read_nyx_dat.m` / `infer_grid_size.m` | `.dat` 二进制文件无法直接查看 | 恢复 n×n×n 三维密度体 | 第 2 章 |
| 读取验证 | `step1_check_data.m` | 无法确认读取方向和文件一致性 | 中心切片、文件大小检查 | 第 2 章 |
| 全时间步量化 | `step2_statistics_all_frames.m` | 单帧无法体现演化过程 | `density_stats.csv`、统计曲线 | 第 3 章 |
| 分布展示 | `step3_histogram_analysis.m` | 统计指标不能展示完整分布 | log-density histogram、time-density heatmap | 第 3 章 |
| 三维结构展示 | `step4_volume_render_keyframes.m` | 二维切片空间信息不足 | 体绘制关键帧、传递函数对比 | 第 4 章 |
| 高密度定位 | `step5_high_density_selection.m` | 高密度尾部缺少空间定位 | P99 mask、MIP、等值面 | 第 5 章 |
| 结构演化量化 | `step6_structure_metrics.m` | 视觉结果缺少量化验证 | 结构指标、差分 MIP | 第 5 章 |
| 交互探索 | `step7_interactive_dashboard.m` | 静态图无法探索其他密度区间 | linked brushing dashboard | 第 6 章 |
| 二维相空间 | `step8_density_gradient_phase_space.m` | 单变量密度刷选不足 | density-gradient 相空间和区域 MIP | 第 7 章 |
| 结构识别 | `step9_hessian_cosmic_web_classification.m` | 仅靠视觉观察难以分类结构 | Void/Sheet/Filament/Node 分类 | 第 7 章 |
| 阶段划分 | `step10_time_similarity_stage_analysis.m` | 代表时间步选择缺少依据 | 相似性矩阵和阶段划分 | 第 7 章 |
| Web 展示 | `Nyx_Web_Visualization/` / `start_nyx_web_demo.py` | MATLAB GUI 和静态报告不便于答辩现场演示 | HTML 单页交互系统、时间播放、直方图刷选、top 1% 联动 | 第 7.7 节、第 8.11 节 |

## 4 数据读取相关代码

### `code/read_nyx_dat.m`

功能：读取单个 Nyx `.dat` 文件，并恢复为三维体数据。

输入：`.dat` 文件路径。

输出：`single` 类型三维数组 `V`。

关键逻辑：使用 little-endian float32 读取；自动推断网格尺寸；先 reshape 为 `[z,y,x]`，再 permute 为 `[x,y,z]`。

### `code/infer_grid_size.m`

功能：根据体素总数推断 n，使 `numel(data)=n^3`。若无法构成规则立方体，则报错提示文件尺寸异常。

## 5 分步骤分析代码

### `code/step1_check_data.m`

功能：完成数据完整性检查，生成中心切片和文件大小一致性检查图。

对应问题：读取出的数据是否合理，轴顺序是否基本正确。

### `code/step2_statistics_all_frames.m`

功能：遍历所有时间步并计算 density 和 log-density 统计量。

对应问题：如何从单帧观察扩展到全时间步量化分析。

### `code/step3_histogram_analysis.m`

功能：绘制代表时间步 log-density histogram 和 time-density heatmap。

对应问题：如何展示密度分布形态及其随时间的整体变化。

### `code/step4_volume_render_keyframes.m`

功能：对代表时间步进行自定义 alpha compositing 体绘制，并比较三种传递函数。

对应问题：二维切片无法完整表达三维空间结构。

### `code/step5_high_density_selection.m`

功能：使用原始 density 的 P99 作为 top 1% 阈值，生成高密度 mask、MIP 和多阈值等值面。

对应问题：如何验证直方图右尾是否对应空间中的高密度节点。

### `code/step6_structure_metrics.m`

功能：计算结构指标、连通域指标和首末时间步差分 MIP。

对应问题：如何对结构演化进行量化补充。

### `code/step7_interactive_dashboard.m`

功能：实现 linked brushing 交互仪表盘。

对应问题：静态图像无法灵活探索不同时间步和密度区间。

### `code/step8_density_gradient_phase_space.m`

功能：构建 normalized log-density 与 gradient magnitude 的二维相空间热力图，并输出典型区域空间 MIP。

对应问题：一维 density 刷选不能区分结构边界和致密核心。

### `code/step9_hessian_cosmic_web_classification.m`

功能：基于 Hessian 特征值进行 Void、Sheet、Filament、Node 粗分类，输出分类切片、结构占比 CSV 和占比曲线。

对应问题：仅靠视觉观察难以识别局部结构类型。

### `code/step10_time_similarity_stage_analysis.m`

功能：基于全时间步 log-density histogram 计算时间步相似性矩阵，输出演化阶段划分。

对应问题：代表时间步选择和阶段划分缺少数据依据。

## 6 Web 交互展示代码

### `Nyx_Web_Visualization/scripts/preprocess_for_web.py`

功能：将原始 Nyx `.dat` 文件转换为浏览器可加载的轻量 Web 数据。

输入：`data/raw/` 下的原始 `.dat` 文件，以及可选的 `data/processed/` 统计结果。

输出：`metadata.json`、`density_stats.json`、`histograms.json`、`time_density_heatmap.json`、`time_similarity_stage.json` 和 `vol_0000.bin` 至 `vol_0099.bin` 等降采样体数据。

关键逻辑：读取 little-endian float32；恢复 x-y-z 三维体；计算 log-density；做 P5-P99.7 百分位归一化；降采样到 64×64×64；保存为 Float32Array 二进制文件。

### `Nyx_Web_Visualization/web/index.html`

功能：定义最终 HTML 单页交互展示系统的页面结构。

主要区域：顶部项目说明、左侧控制面板、主空间视图、右侧直方图刷选视图、底部 time-density heatmap、统计指标曲线、路线图、方法图和结论卡片。

### `Nyx_Web_Visualization/web/style.css`

功能：定义深色宇宙主题仪表盘风格。

设计要点：深蓝紫背景、半透明面板、高亮指标卡、Canvas 主视图、适合 1920×1080 课堂录屏展示。

### `Nyx_Web_Visualization/web/app.js`

功能：实现浏览器端数据加载、体数据懒加载、时间步播放、百分位刷选、MIP/Mask/Slice/简化体合成、histogram brush、time-density heatmap 和统计曲线联动。

对应问题：如何把 MATLAB/Python 离线分析结果转化为可以现场演示的 Web 交互作品。

### `start_nyx_web_demo.py` / `start_nyx_web_demo.bat`

功能：一键启动 Web 演示。脚本会检查 Web 数据是否存在，必要时自动运行预处理，然后寻找可用端口、启动本地 HTTP 服务器并打开浏览器。

使用方式：双击 `start_nyx_web_demo.bat`，或在命令行运行 `python start_nyx_web_demo.py`。

## 7 工具函数

| 工具函数 | 功能 |
|---|---|
| `ensure_folder.m` | 若目标目录不存在，则自动创建 |
| `save_figure.m` | 统一保存 MATLAB figure，兼容新旧 MATLAB |
| `normalize_percentile.m` | 百分位裁剪并归一化到 [0,1] |
| `smoothstep.m` | 平滑阶跃函数，用于透明度传递函数 |
| `custom_nyx_colormap.m` | 自定义 Nyx 颜色映射 |
| `volume_render_alpha_composite.m` | 自定义 alpha compositing 体绘制 |
| `gini_coefficient.m` | 计算 density 的 Gini 系数 |
| `compute_gradient_magnitude.m` | 计算三维梯度幅值 |
| `compute_hessian_eigenvalues_3d.m` | 计算三维 Hessian 特征值 |
| `classify_cosmic_web_hessian.m` | 基于 Hessian 特征值进行结构分类 |
| `cosine_similarity_matrix.m` | 计算 histogram 向量之间的余弦相似性 |
| `simple_stage_segmentation.m` | 基于 distance from start 划分演化阶段 |

## 8 答辩讲解顺序

1. 说明 `.dat` 文件不能直接查看，因此首先恢复三维体数据。
2. 展示中心切片，说明读取和轴顺序恢复基本正确。
3. 展示全时间步统计曲线和 time-density heatmap，说明从单帧扩展到演化过程。
4. 说明 log-density 和百分位归一化解决高动态范围显示问题。
5. 展示体绘制结果，说明三维宇宙网结构的整体表达。
6. 展示 P99 筛选和多阈值等值面，说明统计右尾与空间节点对应。
7. 展示 linked brushing dashboard，说明交互探索能力。
8. 展示 density-gradient、Hessian 和 time similarity，说明进阶结构识别和阶段划分。
9. 打开 Web 页面，现场演示 time slider、Play、histogram brush 和 top 1% 高密度空间联动。
10. 最后总结系统优势、局限和后续扩展方向。

## 9 可选功能和降级方案

| 功能 | 首选实现 | 降级方案 |
|---|---|---|
| P99 连通域分析 | `bwconncomp` | 工具箱不可用时跳过并 warning |
| 3D Gaussian smoothing | `imgaussfilt3` | `smooth3` 或 separable convolution |
| 体数据降采样 | `imresize3` | 使用步长抽样 |
| 体绘制 | 自定义 alpha compositing | 如果较慢，可使用 MIP 帧 |
| 视频生成 | FFmpeg | 只保留 `results/frames/` |
| 阶段划分 | distance_from_start 分位划分 | 分位不稳定时按时间均分四段 |
| Web 启动 | `start_nyx_web_demo.bat` | 命令行运行 `python start_nyx_web_demo.py` |
| Web 体绘制 | Canvas 简化 alpha compositing | 切换为 MIP、Mask 或 Slice 模式 |
"""


def write_outputs() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    create_iteration_roadmap(FIG_DIR / "iteration_roadmap.png")
    create_web_system_overview(FIG_DIR / "web_system_overview.png")
    (REPORT_DIR / "大作业技术报告草稿_迭代过程版.md").write_text(report_markdown(), encoding="utf-8")
    (REPORT_DIR / "代码说明_迭代过程版.md").write_text(code_explanation_markdown(), encoding="utf-8")
    (REPORT_DIR / "大作业技术报告草稿_Web展示版.md").write_text(report_markdown(), encoding="utf-8")
    (REPORT_DIR / "代码说明_Web展示版.md").write_text(code_explanation_markdown(), encoding="utf-8")


def add_markdown_to_docx(md_path: Path, docx_path: Path) -> bool:
    """将本任务生成的 Markdown 转为简洁可编辑 DOCX。"""

    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT, WD_SECTION
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        print("未检测到 python-docx，跳过 Word 版本生成。")
        return False

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    def apply_page_setup(sec, landscape: bool = False) -> None:
        """统一设置页面方向、页边距和页眉页脚。"""
        if landscape:
            sec.orientation = WD_ORIENT.LANDSCAPE
            sec.page_width, sec.page_height = Inches(11.69), Inches(8.27)
        else:
            sec.orientation = WD_ORIENT.PORTRAIT
            sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)
        sec.top_margin = Inches(0.9)
        sec.bottom_margin = Inches(0.9)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)

    def add_field(paragraph, instruction: str) -> None:
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)

    def decorate_section(sec) -> None:
        header = sec.header.paragraphs[0]
        header.text = "Nyx 宇宙密度演化可视分析"
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header.runs:
            run.font.size = Pt(9)
            run.font.name = "宋体"
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("第 ")
        add_field(footer, "PAGE")
        footer.add_run(" 页 / 共 ")
        add_field(footer, "NUMPAGES")
        footer.add_run(" 页")
        for run in footer.runs:
            run.font.size = Pt(9)
            run.font.name = "宋体"

    apply_page_setup(section)
    decorate_section(section)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True

    lines = md_path.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_buffer: list[str] = []

    def flush_code() -> None:
        nonlocal code_buffer
        if not code_buffer:
            return
        p = doc.add_paragraph()
        for run_text in ["\n".join(code_buffer)]:
            run = p.add_run(run_text)
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
        code_buffer = []

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue
        if not line:
            i += 1
            continue

        img_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if img_match:
            caption = img_match.group(1)
            rel = img_match.group(2)
            img_path = (md_path.parent / rel).resolve()
            if not img_path.exists():
                img_path = (PROJECT_ROOT / rel.replace("../", "")).resolve()
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                try:
                    run.add_picture(str(img_path), width=Inches(5.9))
                except Exception:
                    pass
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for tl in table_lines:
                if re.match(r"^\|\s*[-:]+", tl):
                    continue
                rows.append([c.strip() for c in tl.strip("|").split("|")])
            if rows:
                wide_table = len(rows[0]) >= 6
                if wide_table:
                    landscape_section = doc.add_section(WD_SECTION.NEW_PAGE)
                    apply_page_setup(landscape_section, landscape=True)
                    decorate_section(landscape_section)
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                table.autofit = True
                for r, cells in enumerate(rows):
                    for c, cell_text in enumerate(cells):
                        table.cell(r, c).text = cell_text
                        para = table.cell(r, c).paragraphs[0]
                        for run in para.runs:
                            run.font.name = "宋体"
                            run.font.size = Pt(8 if wide_table else 9)
                            if r == 0:
                                run.bold = True
                        if r == 0:
                            tc_pr = table.cell(r, c)._tc.get_or_add_tcPr()
                            shd = OxmlElement("w:shd")
                            shd.set(qn("w:fill"), "EDEDED")
                            tc_pr.append(shd)
                if wide_table:
                    portrait_section = doc.add_section(WD_SECTION.NEW_PAGE)
                    apply_page_setup(portrait_section, landscape=False)
                    decorate_section(portrait_section)
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = "黑体"
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            i += 1
            continue
        if re.match(r"^\d+\. ", line):
            doc.add_paragraph(line, style=None)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
        i += 1

    doc.save(docx_path)
    return True


def build_docx_outputs() -> None:
    add_markdown_to_docx(
        REPORT_DIR / "大作业技术报告草稿_迭代过程版.md",
        REPORT_DIR / "Nyx宇宙密度演化可视分析_技术报告_迭代过程版.docx",
    )
    add_markdown_to_docx(
        REPORT_DIR / "代码说明_迭代过程版.md",
        REPORT_DIR / "Nyx宇宙密度演化可视分析_代码说明_迭代过程版.docx",
    )
    add_markdown_to_docx(
        REPORT_DIR / "大作业技术报告草稿_迭代过程版.md",
        REPORT_DIR / "Nyx宇宙密度演化可视分析_技术报告_最终提交版.docx",
    )
    add_markdown_to_docx(
        REPORT_DIR / "代码说明_迭代过程版.md",
        REPORT_DIR / "Nyx宇宙密度演化可视分析_代码说明_最终提交版.docx",
    )
    add_markdown_to_docx(
        REPORT_DIR / "大作业技术报告草稿_Web展示版.md",
        REPORT_DIR / "Nyx宇宙密度演化可视分析_技术报告_Web展示版.docx",
    )
    add_markdown_to_docx(
        REPORT_DIR / "代码说明_Web展示版.md",
        REPORT_DIR / "Nyx宇宙密度演化可视分析_代码说明_Web展示版.docx",
    )


def main() -> None:
    write_outputs()
    build_docx_outputs()
    print("已生成迭代过程版报告、代码说明和路线图。")
    print((FIG_DIR / "iteration_roadmap.png").relative_to(PROJECT_ROOT))
    print((REPORT_DIR / "大作业技术报告草稿_迭代过程版.md").relative_to(PROJECT_ROOT))
    print((REPORT_DIR / "代码说明_迭代过程版.md").relative_to(PROJECT_ROOT))
    print((REPORT_DIR / "Nyx宇宙密度演化可视分析_技术报告_最终提交版.docx").relative_to(PROJECT_ROOT))
    print((REPORT_DIR / "Nyx宇宙密度演化可视分析_代码说明_最终提交版.docx").relative_to(PROJECT_ROOT))
    print((REPORT_DIR / "Nyx宇宙密度演化可视分析_技术报告_Web展示版.docx").relative_to(PROJECT_ROOT))
    print((REPORT_DIR / "Nyx宇宙密度演化可视分析_代码说明_Web展示版.docx").relative_to(PROJECT_ROOT))


if __name__ == "__main__":
    main()
