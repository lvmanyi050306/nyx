from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    Image as RLImage,
    KeepTogether,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont


ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = ROOT / "report"
OUT_MD = REPORT_DIR / "Nyx宇宙密度演化可视分析_Answer_Sheet_答题卡格式版.md"
OUT_DOCX = REPORT_DIR / "Nyx宇宙密度演化可视分析_Answer_Sheet_答题卡格式版.docx"
OUT_PDF = REPORT_DIR / "Nyx宇宙密度演化可视分析_Answer_Sheet_答题卡格式版.pdf"


def rel(path: Path) -> str:
    return "../" + str(path.relative_to(ROOT).as_posix())


IMG = {
    "slices": ROOT / "results" / "01_data_check" / "center_slices_0000.png",
    "volume": ROOT / "results" / "04_volume_render" / "volume_keyframes_compare.png",
    "transfer": ROOT / "results" / "report_figures" / "transfer_function_design.png",
    "hist": ROOT / "results" / "02_histograms" / "histogram_compare_keyframes.png",
    "heatmap": ROOT / "results" / "02_histograms" / "time_density_heatmap.png",
    "stats": ROOT / "results" / "report_figures" / "combined_statistics_curves.png",
    "top1": ROOT / "results" / "report_figures" / "combined_top1_percent_mips.png",
    "nested": ROOT / "results" / "05_high_density" / "nested_isosurfaces_t0099.png",
    "web": ROOT / "results" / "report_figures" / "web_dashboard_actual_screenshot.png",
    "roadmap": ROOT / "results" / "report_figures" / "iteration_roadmap.png",
    "similarity": ROOT / "results" / "10_time_similarity" / "time_step_similarity_matrix.png",
}


QUESTIONS = [
    {
        "title": "【第一题】数据读取与体绘制：如何展示宇宙演化中密度信息的变化？",
        "answer": "Nyx 原始数据为 little-endian float32 格式的 `.dat` 二进制体数据，不能直接作为图片查看。本作品首先根据文件大小推断三维网格尺寸，将 z-y-x 存储顺序恢复为 x-y-z 体数据，再通过中心切片检查读取方向。随后使用 log-density 与百分位归一化压缩动态范围，并采用自定义 alpha compositing、传递函数和梯度增强进行体绘制，从空间形态上展示宇宙密度结构随时间演化的变化。",
        "process": [
            "读取 `.dat` 文件中的 little-endian float32 数据，并根据体素数量推断 `n×n×n` 网格。",
            "按照 z 轴最快、y 轴其次、x 轴最慢的存储顺序恢复三维体数据。",
            "提取 XY、XZ、YZ 三个中心切片，验证读取方向和轴顺序是否合理。",
            "对 density 进行 `log10(max(density, eps))` 变换，并使用 P5-P99.7 百分位归一化。",
            "使用自定义 alpha compositing、LUT、smoothstep 透明度函数和梯度增强生成关键帧体绘制图。"
        ],
        "findings": [
            "中心切片呈现连续纹理，说明数据读取和三维恢复基本正确。",
            "t=0000 时密度结构相对连续，高密度节点不明显；t=0060 和 t=0099 中丝状连接与局部团块更加清晰。",
            "传递函数改变了可观察重点：低密度空洞、中密度丝状结构和高密度节点可以分别被突出显示。"
        ],
        "summary": "体绘制能够直观展示 Nyx 密度场的三维空间结构，是回答“宇宙演化中密度信息如何变化”的主要可视化手段；统计图和筛选结果则用于补充验证体绘制观察。",
        "images": [("图 1：数据读取检查中心切片图", IMG["slices"]), ("图 2：代表时间步体绘制结果图", IMG["volume"])]
    },
    {
        "title": "【第二题】宇宙密度时序统计特征分析：密度分布随时间如何变化？",
        "answer": "为避免只依赖单帧图像，本作品遍历全部时间步，计算 mean、std、max、P01、P05、P50、P95、P99、P99/mean 和 P99-P01 等指标，并绘制代表时间步 log-density histogram 与 time-density heatmap。统计结果用于从数值分布角度观察密度场由相对集中到逐渐分化的变化过程。",
        "process": [
            "对 100 个时间步逐帧读取密度体数据，并展开为体素数组。",
            "计算全局统计量和百分位指标，保存为 `density_stats.csv`。",
            "对密度进行 log-density 变换，构建代表时间步的对数密度直方图。",
            "使用统一 bins 叠加全部时间步的 histogram，生成 time-density heatmap。",
            "通过统计曲线追踪 std、P99、P99/mean 和 P99-P01 的变化。"
        ],
        "findings": [
            "早期时间步密度分布更集中，后期分布逐渐变宽。",
            "高密度右尾增强，说明局部极高密度体素逐渐突出。",
            "mean 变化不一定最明显，std、P99/mean 和 P99-P01 更能体现密度不均匀性增强。",
            "统计趋势与体绘制中高密度节点逐渐突出的现象相互印证。"
        ],
        "summary": "本题从统计分布角度验证了 Nyx 密度场演化过程中由相对均匀到高低密度分化的变化趋势。",
        "images": [("图 3：代表时间步 log-density histogram 对比图", IMG["hist"]), ("图 4：time-density heatmap 与统计曲线", IMG["stats"])]
    },
    {
        "title": "【第三题】Top 1% 高密度区域验证：高密度尾部是否具有空间结构意义？",
        "answer": "本作品使用每个时间步原始 density 的 P99 作为 top 1% 高密度阈值，生成高密度 mask，并通过 X/Y/Z 三方向 MIP 和 P90/P95/P99 多阈值等值面观察空间分布。该分析直接回答直方图高密度尾部是否对应空间中的宇宙网致密节点。",
        "process": [
            "选取每个时间步密度分布的 P99 作为极高密度阈值。",
            "筛选 `density >= P99` 的体素，构建 top 1% 高密度 mask。",
            "分别沿 X、Y、Z 三个方向进行最大强度投影，观察高密度体素的空间位置。",
            "使用 P90、P95、P99 多阈值等值面观察结构层级。",
            "将筛选结果与直方图右尾和体绘制结果进行对照。"
        ],
        "findings": [
            "P99 以上体素并非随机分散，而是集中在若干高密度节点和丝状结构交汇区域。",
            "P90、P95、P99 等值面呈现由外围结构到核心结构逐渐收缩的层级关系。",
            "高密度右尾在空间中对应宇宙网的致密节点区域。",
            "P99 筛选结果为体绘制中高亮结构提供了统计依据。"
        ],
        "summary": "P99 top 1% 高密度筛选建立了“统计分布右尾”和“三维空间节点结构”之间的对应关系，是本作品回答赛题核心问题的重要证据。",
        "images": [("图 5：Top 1% 高密度区域三方向 MIP", IMG["top1"]), ("图 6：P90/P95/P99 多阈值等值面图", IMG["nested"])]
    },
    {
        "title": "【第四题】相空间交互式刷选分析：如何实现统计特征与空间结构的双向关联？",
        "answer": "本作品构建了 linked brushing 仪表盘和 HTML Web 展示系统。用户可以在 log-density histogram 或密度百分位滑条中选择密度区间，系统实时生成 mask 并映射回空间投影视图。当选择 99%-100% 区间时，页面会突出显示 top 1% 高密度体素，用于交互验证高密度尾部与空间节点结构之间的关系。",
        "process": [
            "构建 log-density histogram 作为统计分布视图。",
            "设置 density percentile range，用于选择指定密度区间。",
            "根据用户选择实时生成 mask，并映射回三维体数据空间。",
            "在空间投影视图中显示被选体素，支持 MIP、Mask、Masked MIP 和 Slice。",
            "在 Web 页面中加入 time slider、Play/Pause、投影方向切换和 Top 1% 一键筛选。"
        ],
        "findings": [
            "普通密度区间可用于观察不同密度层级的空间分布范围。",
            "选择 99%-100% 时，空间中突出显示的区域集中于高密度节点和丝状交汇处。",
            "直方图刷选与空间联动使用户可以从统计分布反查空间结构。",
            "Web 展示系统比静态图更适合答辩演示和交互探索。"
        ],
        "summary": "相空间交互式刷选使本作品从静态结果展示升级为可探索的可视分析系统，能够直观验证密度分布特征与空间结构之间的对应关系。",
        "images": [("图 7：Nyx Density Explorer Web 交互展示界面", IMG["web"])]
    },
    {
        "title": "【第五题】综合结论：宇宙密度演化规律、可视化技术价值与实验挑战是什么？",
        "answer": "综合体绘制、统计曲线、time-density heatmap、P99 高密度筛选和 linked brushing 结果，Nyx 密度场在课程实验可视化层面呈现由相对均匀到逐渐分化的趋势。高密度尾部增强，低密度区域与高密度节点逐渐分离，丝状结构和团块化特征更加明显。",
        "process": [
            "从数据读取、中心切片、统计分析逐步扩展到体绘制和结构验证。",
            "利用 P99 筛选把密度分布右尾映射回空间节点。",
            "利用 Web 系统实现时间播放、密度刷选和空间联动展示。",
            "使用 density-gradient、Hessian 分类和时间步相似性矩阵作为进阶补充分析。"
        ],
        "findings": [
            "早期密度分布相对集中，后期密度波动增强。",
            "高密度尾部增强，P99、P99/mean 与 P99-P01 可作为团块化观察指标。",
            "P99 高密度体素集中在丝状结构交汇和节点区域。",
            "体绘制、统计曲线、time-density heatmap 和 P99 筛选结果相互印证。",
            "Web 交互系统进一步支持用户从统计分布主动探索空间结构。"
        ],
        "summary": "本作品主要服务于数据可视化课程实验，结论基于给定 Nyx 数据和可视化结果，用于说明科学可视化方法在宇宙密度场分析中的应用价值，不作为正式天体物理结论。",
        "images": [("图 8：从无到有的实验迭代路线图", IMG["roadmap"]), ("图 9：时间步相似性矩阵", IMG["similarity"])]
    },
]


def make_md() -> str:
    lines = [
        "# 答题卡",
        "",
        "**赛题 II：科学可视化挑战赛**  ",
        "**基于体绘制与相空间刷选的 Nyx 宇宙密度演化可视分析**",
        "",
        "课程名称：数据可视化  ",
        "作品名称：Nyx Density Explorer：宇宙密度演化交互式可视分析系统  ",
        "姓名：  ",
        "学号：  ",
        "班级：  ",
        "完成日期：  ",
        "",
    ]
    for q in QUESTIONS:
        lines.append(f"## {q['title']}")
        lines.append("")
        lines.append(q["answer"])
        lines.append("")
        lines.append("●过程：")
        for i, item in enumerate(q["process"], 1):
            lines.append(f"（{i}）{item}")
        lines.append("")
        lines.append("●发现：")
        for i, item in enumerate(q["findings"], 1):
            lines.append(f"（{i}）{item}")
        lines.append("")
        for caption, path in q["images"]:
            if path.exists():
                lines.append(f"![{caption}]({rel(path)})")
                lines.append("")
                lines.append(caption)
            else:
                lines.append(f"[此处插入：{path}]")
            lines.append("")
        lines.append("●总结：")
        lines.append(q["summary"])
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bullet_block(doc, title, items):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(f"（{i}）{item}")
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.05)


def add_image(doc, path: Path, caption: str, width_cm=13.5):
    if not path.exists():
        p = doc.add_paragraph(f"[此处插入：{path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(str(path), width=Cm(width_cm))
    except Exception:
        p.add_run(f"[此处插入：{path}]")
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def make_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.line_spacing = 1.05
    styles["Normal"].paragraph_format.space_after = Pt(3)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("答题卡")
    r.bold = True
    r.font.size = Pt(24)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("赛题 II：科学可视化挑战赛\n基于体绘制与相空间刷选的 Nyx 宇宙密度演化可视分析")
    r.bold = True
    r.font.size = Pt(13)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    info = [
        ("课程名称", "数据可视化"),
        ("作品名称", "Nyx Density Explorer：宇宙密度演化交互式可视分析系统"),
        ("姓名 / 学号 / 班级 / 完成日期", " /  /  / "),
    ]
    for row, (k, v) in zip(table.rows, info):
        row.cells[0].text = k
        row.cells[1].text = v
        set_cell_shading(row.cells[0], "EAF1F8")

    for qi, q in enumerate(QUESTIONS):
        if qi:
            doc.add_paragraph()
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(8)
        hp.paragraph_format.space_after = Pt(4)
        hr = hp.add_run(q["title"])
        hr.bold = True
        hr.font.size = Pt(13)
        hr.font.name = "黑体"
        hr._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

        p = doc.add_paragraph(q["answer"])
        p.paragraph_format.first_line_indent = Cm(0.74)
        add_bullet_block(doc, "●过程：", q["process"])
        add_bullet_block(doc, "●发现：", q["findings"])

        for caption, path in q["images"]:
            width = 12.4 if qi in (0, 3) else 11.8
            if path == IMG["web"]:
                width = 14.0
            add_image(doc, path, caption, width_cm=width)

        p = doc.add_paragraph()
        r = p.add_run("●总结：")
        r.bold = True
        p.add_run(q["summary"])

    doc.save(OUT_DOCX)


def scaled_rl_image(path: Path, max_w_cm=8.0, max_h_cm=4.8):
    if not path.exists():
        return Paragraph(f"[此处插入：{path}]", get_pdf_styles()["body"])
    with PILImage.open(path) as im:
        w, h = im.size
    max_w = max_w_cm * cm
    max_h = max_h_cm * cm
    scale = min(max_w / w, max_h / h)
    return RLImage(str(path), width=w * scale, height=h * scale)


def get_pdf_styles():
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "CNTitle", parent=styles["Title"], fontName="STSong-Light",
            fontSize=24, leading=30, alignment=TA_CENTER, spaceAfter=8
        ),
        "subtitle": ParagraphStyle(
            "CNSubtitle", parent=styles["Normal"], fontName="STSong-Light",
            fontSize=12, leading=17, alignment=TA_CENTER, spaceAfter=10
        ),
        "qtitle": ParagraphStyle(
            "CNQTitle", parent=styles["Heading2"], fontName="STSong-Light",
            fontSize=12.5, leading=17, textColor=colors.HexColor("#111827"),
            spaceBefore=8, spaceAfter=5
        ),
        "body": ParagraphStyle(
            "CNBody", parent=styles["BodyText"], fontName="STSong-Light",
            fontSize=9.2, leading=13.3, alignment=TA_LEFT, spaceAfter=4
        ),
        "label": ParagraphStyle(
            "CNLabel", parent=styles["BodyText"], fontName="STSong-Light",
            fontSize=9.5, leading=13.5, textColor=colors.HexColor("#111827"),
            spaceBefore=2, spaceAfter=2
        ),
        "caption": ParagraphStyle(
            "CNCaption", parent=styles["BodyText"], fontName="STSong-Light",
            fontSize=8.2, leading=10.5, alignment=TA_CENTER, textColor=colors.HexColor("#374151"),
        ),
    }


def add_pdf_image_group(story, image_items, styles):
    cells = []
    for caption, path in image_items:
        cell_story = [scaled_rl_image(path), Spacer(1, 0.08 * cm), Paragraph(caption, styles["caption"])]
        cells.append(cell_story)
    if len(cells) == 1:
        tbl = Table([[cells[0]]], colWidths=[16.2 * cm])
    else:
        tbl = Table([cells[:2]], colWidths=[8.05 * cm, 8.05 * cm])
    tbl.setStyle(TableStyle([
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 0.08 * cm))


def make_pdf():
    styles = get_pdf_styles()
    doc = SimpleDocTemplate(
        str(OUT_PDF),
        pagesize=A4,
        rightMargin=1.35 * cm,
        leftMargin=1.35 * cm,
        topMargin=1.25 * cm,
        bottomMargin=1.2 * cm,
    )
    story = []
    story.append(Paragraph("答题卡", styles["title"]))
    story.append(Paragraph("赛题 II：科学可视化挑战赛<br/>基于体绘制与相空间刷选的 Nyx 宇宙密度演化可视分析", styles["subtitle"]))

    info = [
        ["课程名称", "数据可视化"],
        ["作品名称", "Nyx Density Explorer：宇宙密度演化交互式可视分析系统"],
        ["姓名 / 学号 / 班级 / 完成日期", " /  /  / "],
    ]
    info_tbl = Table(info, colWidths=[4.1 * cm, 12.1 * cm])
    info_tbl.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
        ("FONTSIZE", (0, 0), (-1, -1), 9.2),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EAF1F8")),
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#9CA3AF")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(info_tbl)
    story.append(Spacer(1, 0.18 * cm))

    for q in QUESTIONS:
        block = [Paragraph(q["title"], styles["qtitle"])]
        block.append(Paragraph(q["answer"], styles["body"]))
        block.append(Paragraph("●过程：", styles["label"]))
        for i, item in enumerate(q["process"], 1):
            block.append(Paragraph(f"（{i}）{item}", styles["body"]))
        block.append(Paragraph("●发现：", styles["label"]))
        for i, item in enumerate(q["findings"], 1):
            block.append(Paragraph(f"（{i}）{item}", styles["body"]))
        story.append(KeepTogether(block[:3]))
        story.extend(block[3:])
        add_pdf_image_group(story, q["images"], styles)
        story.append(Paragraph("●总结：" + q["summary"], styles["body"]))
        story.append(Spacer(1, 0.12 * cm))

    doc.build(story)


def main():
    REPORT_DIR.mkdir(exist_ok=True)
    OUT_MD.write_text(make_md(), encoding="utf-8")
    make_docx()
    make_pdf()
    print(OUT_MD)
    print(OUT_DOCX)
    print(OUT_PDF)


if __name__ == "__main__":
    main()
