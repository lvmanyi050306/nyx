from pathlib import Path

from PIL import Image, ImageOps
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
REPORT = ROOT / "report"
OUT_DOCX = REPORT / "技术报告_一页纸摘要_内容丰富版.docx"
THUMB_DIR = ROOT / "results" / "report_figures" / "abstract_thumbs"


IMG_SOURCES = [
    ("图 A：时间-密度热力图，展示全时间步密度分布演化", ROOT / "results" / "02_histograms" / "time_density_heatmap.png"),
    ("图 B：代表时间步体绘制结果，展示宇宙网结构形成过程", ROOT / "results" / "04_volume_render" / "volume_keyframes_compare.png"),
    ("图 C：Top 1% 高密度筛选结果，验证高密度尾部空间位置", ROOT / "results" / "05_high_density" / "top1_percent_t0099.png"),
    ("图 D：Web 交互展示系统界面，支持直方图刷选与空间联动", ROOT / "results" / "report_figures" / "web_dashboard_actual_screenshot.png"),
]


PROBLEMS = [
    (
        "针对问题一：Nyx 数据读取与三维体数据恢复",
        "Nyx 原始数据为 little-endian float32 的 `.dat` 二进制体，不能直接作为图像或表格查看。本项目根据文件大小推断 n×n×n 网格，并按 z-y-x 到 x-y-z 恢复三维密度体；中心切片验证空间方向和数据连续性，解决后续体绘制与统计分析的数据基础问题。"
    ),
    (
        "针对问题二：宇宙密度时序统计特征分析",
        "为刻画 100 个时间步的整体演化，本项目计算 mean、std、P99、P99/mean、P99-P01，并构建 log-density histogram 与 time-density heatmap。图 A 显示密度分布随时间扩散，高密度右尾增强，说明后期不均匀性和极端高密度结构更加明显。"
    ),
    (
        "针对问题三：体数据渲染、传递函数与光照效果设计",
        "针对二维切片难以表达三维宇宙网的问题，项目使用 log-density、P5-P99.7 百分位归一化、alpha compositing、LUT 传递函数和梯度增强生成体绘制关键帧。图 B 展示密度结构由较均匀状态逐渐发展为丝状连接和高密度节点。"
    ),
    (
        "针对问题四：P99 高密度筛选与宇宙网节点验证",
        "为避免仅凭体绘制主观判断，项目以原始 density 的 P99 作为 top 1% 高密度阈值，生成 mask，并通过 MIP 与多阈值等值面分析空间分布。图 C 表明 P99 以上体素集中在丝状交汇处和局部节点区域，验证直方图右尾的空间意义。"
    ),
    (
        "针对问题五：相空间交互式刷选与 Web 最终展示系统",
        "在静态结果基础上，项目构建 linked brushing 与 Nyx Density Explorer Web 页面，支持 time slider、Play/Pause、histogram brush、Top 1% 一键筛选和 MIP/Mask/Slice 模式切换。图 D 展示用户可从统计密度区间实时反查空间结构。"
    ),
]

CONCLUSION = (
    "综合结论",
    "本作品形成“数据恢复—时序统计—体绘制—高密度验证—交互展示”的完整科学可视化流程。结果表明 Nyx 密度场在课程实验层面由相对均匀走向高低密度分化，高密度尾部与空间节点结构相互印证；相关结论基于给定数据和可视化结果，不作为正式天体物理结论。"
)


def set_run_font(run, name="宋体", size=8.0, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def compact(paragraph, before=0, after=0, line=0.92):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_borderless(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tc_pr.append(borders)


def crop_for_abstract(img: Image.Image, name: str) -> Image.Image:
    """为摘要页裁剪重点区域，减少白边和长截图压缩。"""
    w, h = img.size
    if name == "abstract_3":
        # Top 1% MIP 图通常是横向三联图，只去除少量外边缘。
        return img.crop((int(w * 0.015), int(h * 0.05), int(w * 0.985), int(h * 0.93)))
    if name == "abstract_4":
        # Web 长截图只保留首屏主要交互区域，裁成接近 16:9，避免左右留白。
        return img.crop((0, int(h * 0.112), w, int(h * 0.537)))
    return img


def make_thumb(src: Path, name: str, size=(1120, 650)) -> Path:
    THUMB_DIR.mkdir(parents=True, exist_ok=True)
    dst = THUMB_DIR / f"{name}.png"
    if not src.exists():
        return src

    img = Image.open(src).convert("RGB")
    img = crop_for_abstract(img, name)
    if name == "abstract_3":
        # 横向三联 MIP 按原比例缩放，避免上下留白挤占图像可读性。
        target_w = 1120
        target_h = max(1, round(img.height * target_w / img.width))
        img = img.resize((target_w, target_h), Image.Resampling.LANCZOS)
        img = ImageOps.expand(img, border=2, fill=(210, 216, 225))
        img.save(dst, quality=92)
        return dst

    if name == "abstract_4":
        # 首屏 Web 截图裁剪后直接缩放到统一横向比例，突出控制、空间视图和直方图。
        img = img.resize(size, Image.Resampling.LANCZOS)
        img = ImageOps.expand(img, border=2, fill=(210, 216, 225))
        img.save(dst, quality=92)
        return dst

    bg = Image.new("RGB", size, (248, 250, 252))
    img.thumbnail((size[0], size[1]), Image.Resampling.LANCZOS)
    x = (size[0] - img.width) // 2
    y = (size[1] - img.height) // 2
    bg.paste(img, (x, y))
    bg = ImageOps.expand(bg, border=2, fill=(210, 216, 225))
    bg.save(dst, quality=92)
    return dst


def add_problem(cell, title, body):
    p = cell.add_paragraph()
    compact(p, after=0)
    r = p.add_run(title)
    set_run_font(r, name="黑体", size=7.25, bold=True, color="1F4E79")

    p = cell.add_paragraph()
    compact(p, after=0.8, line=0.86)
    r = p.add_run(body)
    set_run_font(r, size=6.55)


def add_image_cell(cell, img_path: Path, caption: str):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for p in cell.paragraphs:
        compact(p, after=0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if img_path.exists():
        p.add_run().add_picture(str(img_path), width=Cm(8.95))
    else:
        r = p.add_run("[待插入：对应结果图]")
        set_run_font(r, size=7.2, color="8A4B00")

    cp = cell.add_paragraph()
    compact(cp, after=1)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(caption)
    set_run_font(r, size=6.15, color="4B5563")


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.65)
    section.bottom_margin = Cm(0.55)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(7.2)

    title = doc.add_paragraph()
    compact(title, after=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("摘要")
    set_run_font(r, name="黑体", size=16, bold=True)

    subtitle = doc.add_paragraph()
    compact(subtitle, after=0, line=0.95)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("赛题 II：科学可视化挑战赛\n基于体绘制与相空间刷选的 Nyx 宇宙密度演化可视分析")
    set_run_font(r, name="宋体", size=8.5, bold=True)

    info = doc.add_paragraph()
    compact(info, after=2)
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info.add_run("姓名：          学号：          班级：")
    set_run_font(r, size=7.4)

    text_tbl = doc.add_table(rows=1, cols=2)
    text_tbl.autofit = False
    for cell in text_tbl.rows[0].cells:
        set_borderless(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_width(text_tbl.rows[0].cells[0], 9.55)
    set_cell_width(text_tbl.rows[0].cells[1], 9.55)

    left = text_tbl.rows[0].cells[0]
    right = text_tbl.rows[0].cells[1]
    for title_text, body in PROBLEMS[:3]:
        add_problem(left, title_text, body)
    for title_text, body in PROBLEMS[3:]:
        add_problem(right, title_text, body)
    add_problem(right, CONCLUSION[0], CONCLUSION[1])

    img_tbl = doc.add_table(rows=2, cols=2)
    img_tbl.autofit = False
    for row in img_tbl.rows:
        for cell in row.cells:
            set_borderless(cell)
            set_cell_width(cell, 9.55)

    thumbs = [
        (caption, make_thumb(path, f"abstract_{idx}"))
        for idx, (caption, path) in enumerate(IMG_SOURCES, start=1)
    ]
    cells = [img_tbl.cell(0, 0), img_tbl.cell(0, 1), img_tbl.cell(1, 0), img_tbl.cell(1, 1)]
    for cell, (caption, path) in zip(cells, thumbs):
        add_image_cell(cell, path, caption)

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    build_docx()
    print(OUT_DOCX)
