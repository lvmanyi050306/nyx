"""生成 Nyx 技术报告/PPT 使用的辅助概览型图像。

运行命令：
    python code/python/generate_report_summary_figures.py

输出：
    results/report_figures/nyx_keyword_wordcloud.png
    results/report_figures/nyx_method_bubble_tags.png
    results/report_figures/nyx_structure_wordcloud.png
"""

from __future__ import annotations

import math
import random
from pathlib import Path
from typing import Dict, Iterable, List, Sequence, Tuple

import matplotlib

# 使用无 GUI 后端，保证在命令行、服务器或 MATLAB 调用环境中都能运行。
matplotlib.use("Agg")

import matplotlib.pyplot as plt
from matplotlib.colors import to_hex
from matplotlib.patches import Circle, FancyBboxPatch
from matplotlib import font_manager
from PIL import Image, ImageDraw, ImageFont


try:
    from wordcloud import WordCloud

    WORDCLOUD_AVAILABLE = True
except Exception:
    WordCloud = None
    WORDCLOUD_AVAILABLE = False


# code/python/generate_report_summary_figures.py -> code/python -> code -> project root
PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = PROJECT_ROOT / "results" / "report_figures"


KEYWORD_FREQ = {
    "Nyx": 110,
    "Density": 105,
    "Volume Rendering": 100,
    "Linked Brushing": 95,
    "Log-Density": 90,
    "P99": 86,
    "Isosurface": 82,
    "MIP": 78,
    "Time-Density Heatmap": 76,
    "Histogram": 74,
    "Transfer Function": 72,
    "Alpha Compositing": 70,
    "Gradient Enhancement": 68,
    "Structure Metrics": 65,
    "Entropy": 62,
    "Gini": 60,
    "Void": 58,
    "Filament": 56,
    "Node": 54,
    "Cosmic Web": 52,
    "Time Series": 50,
    "Phase-space Selection": 48,
    "Interactive Dashboard": 46,
    "Percentile Normalization": 44,
    "Top 1%": 42,
    "Binary Parsing": 40,
    "Axis Reordering": 38,
}


STRUCTURE_FREQ = {
    "Cosmic Web": 110,
    "Filament": 100,
    "Node": 95,
    "Void": 90,
    "Cluster": 85,
    "Density Growth": 80,
    "Aggregation": 76,
    "Non-uniformity": 72,
    "High Density": 68,
    "Low Density": 64,
    "Tail Amplification": 60,
    "Hierarchy": 56,
    "Spatial Structure": 52,
    "Evolution": 48,
    "Connectivity": 44,
    "Collapse": 40,
    "Void Expansion": 38,
    "Dense Core": 36,
    "Structure Formation": 34,
    "Matter Clustering": 32,
}


PALETTE = [
    "#1F4E79",  # 深蓝
    "#5B3F8C",  # 紫色
    "#B35C1E",  # 低饱和橙
    "#9E2F3F",  # 暗红
    "#236B5E",  # 墨绿
    "#5A6E8C",  # 蓝灰
]


def palette_color_func(*args, **kwargs) -> str:
    """wordcloud 的颜色函数：从低饱和专业配色中选择颜色。"""

    return random.choice(PALETTE)


def create_wordcloud_or_fallback(
    freq_dict: Dict[str, int],
    title: str,
    output_path: Path,
    colormap_name: str = "viridis",
) -> None:
    """优先使用 wordcloud 生成词云；若不可用，则自动降级为 matplotlib 文本散布。

    Parameters
    ----------
    freq_dict:
        关键词及权重。
    title:
        图像内部英文标题，避免中文字体缺失。
    output_path:
        输出图片路径。
    colormap_name:
        fallback 文本散布使用的 matplotlib colormap。
    """

    if WORDCLOUD_AVAILABLE:
        random.seed(42)
        wc = WordCloud(
            width=1800,
            height=1000,
            background_color="white",
            prefer_horizontal=0.9,
            collocations=False,
            max_words=120,
            random_state=42,
            margin=6,
            color_func=palette_color_func,
            font_path=None,
        ).generate_from_frequencies(freq_dict)

        fig, ax = plt.subplots(figsize=(10, 6), dpi=300)
        ax.imshow(wc, interpolation="bilinear")
        ax.set_axis_off()
        ax.set_title(title, fontsize=18, fontweight="bold", color="#17324D", pad=14)
        fig.patch.set_facecolor("white")
        fig.savefig(output_path, dpi=300, bbox_inches="tight", facecolor="white")
        plt.close(fig)
    else:
        print("未检测到 wordcloud 库，已使用 matplotlib 文本散布方式生成替代词云。")
        create_fallback_text_cloud(freq_dict, title, output_path, colormap_name)


def create_fallback_text_cloud(
    freq_dict: Dict[str, int],
    title: str,
    output_path: Path,
    colormap_name: str = "viridis",
) -> None:
    """?? PIL + matplotlib ???????????

    wordcloud ????????????????????????
    ?????????????????????????????
    ???????????????????????
    """

    rng = random.Random(20260627)
    items = sorted(freq_dict.items(), key=lambda x: x[1], reverse=True)
    min_w = min(freq_dict.values())
    max_w = max(freq_dict.values())
    cmap = plt.get_cmap(colormap_name)

    width, height = 1800, 1000
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    # ?? matplotlib ?? DejaVu ??????????????
    regular_font = font_manager.findfont("DejaVu Sans")
    bold_font = font_manager.findfont("DejaVu Sans:bold")

    title_font = ImageFont.truetype(bold_font, 60)
    title_bbox = draw.textbbox((0, 0), title, font=title_font)
    title_w = title_bbox[2] - title_bbox[0]
    draw.text(((width - title_w) / 2, 28), title, fill="#17324D", font=title_font)

    cloud_box = (90, 170, width - 90, height - 90)
    placed: List[Tuple[int, int, int, int]] = []

    def overlaps(box: Tuple[int, int, int, int], pad: int = 8) -> bool:
        x0, y0, x1, y1 = box
        x0 -= pad
        y0 -= pad
        x1 += pad
        y1 += pad
        for a0, b0, a1, b1 in placed:
            if not (x1 < a0 or x0 > a1 or y1 < b0 or y0 > b1):
                return True
        return False

    def inside(box: Tuple[int, int, int, int]) -> bool:
        x0, y0, x1, y1 = box
        return x0 >= cloud_box[0] and y0 >= cloud_box[1] and x1 <= cloud_box[2] and y1 <= cloud_box[3]

    def candidates(rank: int):
        fixed = [
            (0.50, 0.52),
            (0.39, 0.56),
            (0.62, 0.56),
            (0.50, 0.43),
            (0.35, 0.45),
            (0.66, 0.45),
        ]
        if rank < len(fixed):
            yield int(width * fixed[rank][0]), int(height * fixed[rank][1])
        for k in range(900):
            angle = 0.50 * k + rng.uniform(-0.18, 0.18)
            radius = 8 + 2.35 * k
            x = width * 0.50 + radius * math.cos(angle) * 1.32
            y = height * 0.52 + radius * math.sin(angle) * 0.78
            yield int(x), int(y)

    for rank, (word, weight) in enumerate(items):
        norm = (weight - min_w) / max(max_w - min_w, 1)
        base_size = int(30 + 96 * math.sqrt(norm))
        rgba = cmap(0.18 + 0.72 * norm)
        color = tuple(int(255 * c) for c in rgba[:3])

        placed_this = False
        for scale in [1.00, 0.92, 0.84, 0.76, 0.68, 0.60, 0.52]:
            size = max(18, int(base_size * scale))
            font_path = bold_font if rank < 9 else regular_font
            font = ImageFont.truetype(font_path, size)
            bbox0 = draw.textbbox((0, 0), word, font=font)
            tw = bbox0[2] - bbox0[0]
            th = bbox0[3] - bbox0[1]
            for cx, cy in candidates(rank):
                box = (cx - tw // 2, cy - th // 2, cx + tw // 2, cy + th // 2)
                if inside(box) and not overlaps(box, pad=7):
                    draw.text((box[0], box[1]), word, fill=color, font=font)
                    placed.append(box)
                    placed_this = True
                    break
            if placed_this:
                break

        if not placed_this:
            # ???????????????????????????
            continue

    image.save(output_path, dpi=(300, 300))

def bubble_size(weight: int) -> float:
    """??????????????"""

    return 0.043 + (weight - 40) / 70 * 0.017


def draw_bubble(
    ax: plt.Axes,
    x: float,
    y: float,
    radius: float,
    label: str,
    color: str,
    fontsize: float,
) -> None:
    """??????????????"""

    circle = Circle((x, y), radius, facecolor=color, edgecolor="white", linewidth=2.0, alpha=0.92)
    ax.add_patch(circle)
    ax.text(
        x,
        y,
        label,
        ha="center",
        va="center",
        fontsize=fontsize,
        color="white",
        fontweight="bold",
        wrap=True,
    )


def create_method_bubble_tags(output_path: Path) -> None:
    """???????????????????"""

    groups = [
        {
            "title": "Data Parsing",
            "rect": (0.05, 0.64, 0.46, 0.25),
            "color": "#2F6F9F",
            "items": [
                ("Binary\nParsing", 78, 0.27, 0.62),
                ("Little-\nendian\nfloat32", 70, 0.67, 0.62),
                ("Grid\nInference", 66, 0.38, 0.28),
                ("Axis\nReordering", 62, 0.76, 0.28),
            ],
        },
        {
            "title": "Statistical Analysis",
            "rect": (0.69, 0.64, 0.46, 0.25),
            "color": "#6B5AA6",
            "items": [
                ("Histogram", 78, 0.27, 0.62),
                ("Time-\nDensity\nHeatmap", 76, 0.69, 0.62),
                ("Gini", 60, 0.20, 0.27),
                ("Percentiles", 70, 0.50, 0.27),
                ("Entropy", 62, 0.80, 0.27),
            ],
        },
        {
            "title": "Volume Visualization",
            "rect": (0.31, 0.36, 0.58, 0.22),
            "color": "#C46A2B",
            "items": [
                ("Log\nTransform", 74, 0.18, 0.60),
                ("Percentile\nNorm.", 72, 0.50, 0.63),
                ("Transfer\nFunction", 78, 0.82, 0.60),
                ("Alpha\nComposite", 76, 0.34, 0.26),
                ("Gradient\nEnhance.", 70, 0.66, 0.25),
            ],
        },
        {
            "title": "Spatial Structure",
            "rect": (0.05, 0.07, 0.46, 0.24),
            "color": "#2F7A68",
            "items": [
                ("P99\nSelection", 76, 0.28, 0.62),
                ("Top 1%\nMask", 72, 0.68, 0.62),
                ("Nested\nThres.", 70, 0.20, 0.25),
                ("MIP", 68, 0.50, 0.24),
                ("Iso-\nsurface", 74, 0.80, 0.25),
            ],
        },
        {
            "title": "Interaction",
            "rect": (0.69, 0.07, 0.46, 0.24),
            "color": "#A53F55",
            "items": [
                ("Linked\nBrushing", 80, 0.28, 0.62),
                ("Phase-\nspace\nSelection", 74, 0.68, 0.62),
                ("Density\nRange\nSlider", 66, 0.20, 0.25),
                ("Dashboard", 70, 0.50, 0.24),
                ("Time\nSlider", 64, 0.80, 0.25),
            ],
        },
    ]

    fig, ax = plt.subplots(figsize=(11.2, 6.4), dpi=300)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    ax.set_xlim(0, 1.2)
    ax.set_ylim(0, 1)
    ax.set_aspect("equal", adjustable="box")
    ax.set_axis_off()
    ax.set_title(
        "Visualization Method System for Nyx Density Analysis",
        fontsize=17,
        fontweight="bold",
        color="#17324D",
        pad=14,
    )

    for group in groups:
        x0, y0, panel_w, panel_h = group["rect"]
        panel = FancyBboxPatch(
            (x0, y0),
            panel_w,
            panel_h,
            boxstyle="round,pad=0.012,rounding_size=0.025",
            linewidth=1.2,
            edgecolor="#D7DCE3",
            facecolor="#FAFBFC",
            alpha=1.0,
        )
        ax.add_patch(panel)
        ax.text(
            x0 + panel_w / 2,
            y0 + panel_h - 0.028,
            group["title"],
            ha="center",
            va="center",
            fontsize=12.5,
            fontweight="bold",
            color="#2B3440",
        )

        for label, weight, rx, ry in group["items"]:
            radius = bubble_size(weight)
            max_line_len = max(len(line) for line in label.split("\n"))
            fontsize = 4.9 if "\n" in label and len(label) > 13 else 6.0
            if max_line_len > 8:
                fontsize = min(fontsize, 5.2)
            x = x0 + panel_w * rx
            y = y0 + panel_h * ry
            draw_bubble(ax, x, y, radius, label, group["color"], fontsize)

    fig.savefig(output_path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close(fig)

def insert_summary_figures_to_docx(docx_path: Path, output_docx_path: Path) -> None:
    """可选：将两张报告正文概览图插入技术报告。

    该函数默认不在 main() 中执行。若需要使用，可在 Python 中手动调用。
    如果无法定位章节位置，函数不会破坏原文档，只在末尾给出插入提示。
    """

    try:
        from docx import Document
        from docx.shared import Inches
    except Exception:
        print("未检测到 python-docx，已跳过 Word 自动插图功能。")
        return

    doc = Document(docx_path)
    target_index = None
    for i, paragraph in enumerate(doc.paragraphs):
        if "3 关键技术描述" in paragraph.text or "关键技术描述" == paragraph.text.strip():
            target_index = i
            break

    keyword_img = OUTPUT_DIR / "nyx_keyword_wordcloud.png"
    bubble_img = OUTPUT_DIR / "nyx_method_bubble_tags.png"

    if target_index is None:
        print("未定位到第 3 章位置，未自动插入图片。请手动插入生成的 PNG。")
        return

    # python-docx 不方便在任意段落前插入复杂图片，这里采用稳妥方式：追加到文档末尾并提示用户移动。
    doc.add_page_break()
    doc.add_heading("报告概览辅助图", level=1)
    doc.add_paragraph("以下两张图片可移动到第 3 章“关键技术描述”开头或方法对比表之前。")
    doc.add_picture(str(keyword_img), width=Inches(6.2))
    doc.add_paragraph("图 X：Nyx 宇宙密度演化可视分析关键词词云")
    doc.add_picture(str(bubble_img), width=Inches(6.2))
    doc.add_paragraph("图 X：Nyx 宇宙密度演化可视分析方法亮点气泡标签图")
    doc.save(output_docx_path)
    print(f"已生成可选插图版 Word：{output_docx_path}")


def print_insert_suggestions() -> None:
    """打印报告和 PPT 插入建议。"""

    print("\n技术报告正文建议插入：")
    print("1. nyx_keyword_wordcloud.png")
    print("   插入位置：第 3 章“关键技术描述”开头，图题为“图 X：Nyx 宇宙密度演化可视分析关键词词云”。")
    print("\n2. nyx_method_bubble_tags.png")
    print("   插入位置：第 3 章“关键技术描述”方法对比表之前，图题为“图 X：Nyx 宇宙密度演化可视分析方法亮点气泡标签图”。")
    print("\nPPT 或附录建议插入：")
    print("3. nyx_structure_wordcloud.png")
    print("   插入位置：技术报告附录或答辩 PPT 结论页，图题为“图 X：Nyx 密度演化结构特征词云”。")


def main() -> None:
    """主函数：依次生成三张报告/PPT 辅助概览图。"""

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    keyword_path = OUTPUT_DIR / "nyx_keyword_wordcloud.png"
    bubble_path = OUTPUT_DIR / "nyx_method_bubble_tags.png"
    structure_path = OUTPUT_DIR / "nyx_structure_wordcloud.png"

    create_wordcloud_or_fallback(
        KEYWORD_FREQ,
        "Key Terms for Nyx Visualization",
        keyword_path,
        colormap_name="plasma",
    )
    create_method_bubble_tags(bubble_path)
    create_wordcloud_or_fallback(
        STRUCTURE_FREQ,
        "Structure Features of Nyx Density Evolution",
        structure_path,
        colormap_name="magma",
    )

    print("已生成图片：")
    print(f"1. {keyword_path.relative_to(PROJECT_ROOT).as_posix()}")
    print(f"2. {bubble_path.relative_to(PROJECT_ROOT).as_posix()}")
    print(f"3. {structure_path.relative_to(PROJECT_ROOT).as_posix()}")
    print_insert_suggestions()


if __name__ == "__main__":
    main()
