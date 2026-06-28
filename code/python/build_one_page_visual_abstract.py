from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image


ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = ROOT / "report"
OUT_MD = REPORT_DIR / "技术报告_一页纸摘要_图文版.md"
OUT_DOCX = REPORT_DIR / "技术报告_一页纸摘要_图文版.docx"


IMAGES = [
    ("图 A：时间-密度热力图", ROOT / "results" / "02_histograms" / "time_density_heatmap.png"),
    ("图 B：代表时间步体绘制结果", ROOT / "results" / "04_volume_render" / "volume_keyframes_compare.png"),
    ("图 C：Top 1% 高密度筛选结果", ROOT / "results" / "report_figures" / "combined_top1_percent_mips.png"),
    ("图 D：Web 交互展示系统界面", ROOT / "results" / "report_figures" / "web_dashboard_actual_screenshot.png"),
]


PROBLEMS = [
    (
        "针对问题一：Nyx 数据读取与三维体数据恢复",
        "Nyx 原始数据为 little-endian float32 的 `.dat` 二进制体数据，不能直接作为图片查看。本项目根据文件大小推断三维网格尺寸，按照 z-y-x 存储顺序恢复为 x-y-z 体数据，并使用中心切片检查读取方向和轴顺序，为后续可视化分析建立可靠数据基础。"
    ),
    (
        "针对问题二：宇宙密度时序统计特征分析",
        "为刻画完整演化过程，本项目遍历全部时间步，计算 mean、std、P99、P99/mean、P99-P01 等统计指标，并构建 log-density histogram 与 time-density heatmap。结果显示，早期分布较集中，后期分布逐渐变宽，高密度右尾增强。"
    ),
    (
        "针对问题三：体数据渲染、传递函数与光照效果设计",
        "针对三维体数据难以通过单一切片表达的问题，项目采用 log-density、百分位归一化、自定义 alpha compositing、LUT 传递函数和梯度增强生成关键帧体绘制结果。不同传递函数分别突出空洞、中密度丝状结构和高密度节点。"
    ),
    (
        "针对问题四：P99 高密度筛选与宇宙网节点验证",
        "项目使用原始 density 的 P99 作为 top 1% 高密度阈值，生成高密度 mask，并通过 X/Y/Z 最大强度投影与多阈值等值面观察空间分布。P99 以上体素集中在丝状交汇和节点区域，说明直方图右尾具有明确空间结构意义。"
    ),
    (
        "针对问题五：相空间交互式刷选与 Web 最终展示系统",
        "在静态图基础上，项目进一步构建 linked brushing 与 Nyx Density Explorer Web 页面。用户可通过 time slider、histogram brush、Top 1% 快捷筛选、MIP/Mask/Slice 模式切换，将统计分布中的密度区间实时映射回空间视图。"
    ),
]

CONCLUSION = (
    "综合结论",
    "本作品形成了“数据恢复—统计分析—体绘制—高密度验证—交互展示”的完整流程。可视化结果表明，Nyx 密度场在课程实验层面呈现由相对均匀到高低密度分化的趋势，高密度尾部和空间节点结构相互印证。相关结论基于给定数据和可视化结果，不作为正式天体物理结论。"
)


def rel(path: Path) -> str:
    return "../" + str(path.relative_to(ROOT).as_posix())


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borderless(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tc_pr.append(borders)


def set_run_font(run, name="宋体", size=9, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def compact_paragraph(paragraph, after=1, before=0, line=1.0):
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.line_spacing = line


def add_problem(cell, title: str, body: str):
    p = cell.add_paragraph()
    compact_paragraph(p, after=1)
    r = p.add_run(title)
    set_run_font(r, name="黑体", size=8.5, bold=True, color="1F4E79")
    p = cell.add_paragraph()
    compact_paragraph(p, after=3, line=1.0)
    r = p.add_run(body)
    set_run_font(r, size=7.7)


def add_image_block(cell, caption: str, path: Path, max_width_cm=7.0):
    p = cell.add_paragraph()
    compact_paragraph(p, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if path.exists():
        run = p.add_run()
        with Image.open(path) as img:
            width, height = img.size
        # Keep very wide plots readable and tall screenshots contained.
        if height > width * 0.9:
            width_cm = min(max_width_cm, 6.4)
        else:
            width_cm = max_width_cm
        run.add_picture(str(path), width=Cm(width_cm))
    else:
        run = p.add_run(f"[此处插入：{path}]")
        set_run_font(run, size=7.5, color="9B1C1C")
    cap = cell.add_paragraph()
    compact_paragraph(cap, after=2)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, size=7.3, color="4B5563")


def build_markdown() -> str:
    lines = [
        "# 摘要",
        "",
        "**赛题 II：科学可视化挑战赛**  ",
        "**基于体绘制与相空间刷选的 Nyx 宇宙密度演化可视分析**",
        "",
        "姓名：  ",
        "学号：  ",
        "班级：  ",
        "",
    ]
    for title, body in PROBLEMS:
        lines.append(f"## {title}")
        lines.append(body)
        lines.append("")
    lines.append(f"## {CONCLUSION[0]}")
    lines.append(CONCLUSION[1])
    lines.append("")
    for caption, path in IMAGES:
        if path.exists():
            lines.append(f"![{caption}]({rel(path)})")
            lines.append(caption)
        else:
            lines.append(f"[此处插入：{path}]")
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.05)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.05)
    section.right_margin = Cm(1.05)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(8)

    title = doc.add_paragraph()
    compact_paragraph(title, after=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("摘要")
    set_run_font(r, name="黑体", size=18, bold=True)

    subtitle = doc.add_paragraph()
    compact_paragraph(subtitle, after=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("赛题 II：科学可视化挑战赛\n基于体绘制与相空间刷选的 Nyx 宇宙密度演化可视分析")
    set_run_font(r, name="宋体", size=9.5, bold=True, color="1F2937")

    info = doc.add_paragraph()
    compact_paragraph(info, after=3)
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info.add_run("姓名：        学号：        班级：")
    set_run_font(r, size=8)

    main = doc.add_table(rows=1, cols=2)
    main.autofit = False
    left, right = main.rows[0].cells
    set_cell_width(left, 11.0)
    set_cell_width(right, 7.6)
    set_cell_borderless(left)
    set_cell_borderless(right)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    for title_text, body in PROBLEMS:
        add_problem(left, title_text, body)
    add_problem(left, CONCLUSION[0], CONCLUSION[1])

    for caption, path in IMAGES:
        add_image_block(right, caption, path, max_width_cm=7.1)

    note = doc.add_paragraph()
    compact_paragraph(note, before=1, after=0)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("说明：本页为图文摘要页，图像用于对应赛题问题的可视化支撑；详细方法与代码见技术报告正文及附录。")
    set_run_font(r, size=7.4, color="6B7280")

    doc.save(OUT_DOCX)


def main():
    REPORT_DIR.mkdir(exist_ok=True)
    OUT_MD.write_text(build_markdown(), encoding="utf-8")
    build_docx()
    print(OUT_MD)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
