from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "report"


PRESET = {
    "font_cn": "Microsoft YaHei",
    "font_en": "Calibri",
    "code_font": "Consolas",
    "blue": RGBColor(46, 116, 181),
    "dark_blue": RGBColor(31, 77, 120),
    "muted": RGBColor(90, 90, 90),
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for i, width in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = width
                tc_pr = row.cells[i]._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
                tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, bold=None, color=None, font=None):
    font_name = font or PRESET["font_cn"]
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), PRESET["font_en"] if font is None else font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), PRESET["font_en"] if font is None else font_name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def add_field(paragraph, instr):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_doc(title):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = PRESET["font_cn"]
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), PRESET["font_cn"])
    normal._element.rPr.rFonts.set(qn("w:ascii"), PRESET["font_en"])
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), PRESET["font_en"])
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, PRESET["blue"], 16, 8),
        ("Heading 2", 13, PRESET["blue"], 12, 6),
        ("Heading 3", 12, PRESET["dark_blue"], 8, 4),
    ]:
        st = styles[name]
        st.font.name = PRESET["font_cn"]
        st._element.rPr.rFonts.set(qn("w:eastAsia"), PRESET["font_cn"])
        st._element.rPr.rFonts.set(qn("w:ascii"), PRESET["font_en"])
        st._element.rPr.rFonts.set(qn("w:hAnsi"), PRESET["font_en"])
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("第 ")
    add_field(footer, "PAGE")
    footer.add_run(" 页")
    for r in footer.runs:
        set_run_font(r, size=9, color=PRESET["muted"])

    header = sec.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in header.runs:
        set_run_font(r, size=9, color=PRESET["muted"])
    return doc


def add_cover(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run(title)
    set_run_font(r, size=22, bold=True, color=PRESET["dark_blue"])
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        set_run_font(r, size=13, color=PRESET["muted"])
    doc.add_paragraph()


def add_meta_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_width(table, [Inches(1.65), Inches(4.65)])
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.text = k
        c1.text = v
        set_cell_shading(c0, "F2F4F7")
        for cell in (c0, c1):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run_font(r, size=10.5, bold=(cell is c0))
    doc.add_paragraph()


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    run = p.add_run(code.rstrip())
    set_run_font(run, size=8.5, font=PRESET["code_font"])
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)


def add_markdown_table(doc, lines):
    rows = []
    for line in lines:
        if re.match(r"^\|\s*[-:]+", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    widths = [Inches(6.3 / cols)] * cols
    set_table_width(table, widths)
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.rows[i].cells[j]
            cell.text = row[j] if j < len(row) else ""
            if i == 0:
                set_cell_shading(cell, "F2F4F7")
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run_font(r, size=9.5, bold=(i == 0))
    doc.add_paragraph()


def add_image_if_exists(doc, image_path):
    path = ROOT / image_path
    if not path.exists():
        p = doc.add_paragraph(f"[图片未找到：{image_path}]")
        for r in p.runs:
            set_run_font(r, size=10, color=RGBColor(155, 28, 28))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        run = p.add_run()
        run.add_picture(str(path), width=Inches(5.9))
    except Exception as exc:
        p.add_run(f"[图片插入失败：{image_path}，{exc}]")


def add_para(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    if text.strip() == "":
        return p
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def convert_markdown(doc, md_text, image_mode=False, skip_first_title=True):
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lines = []
    table_lines = []
    first_title_skipped = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, "\n".join(code_lines))
                in_code = False
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines.append(stripped)
            i += 1
            if i >= len(lines) or not (lines[i].strip().startswith("|") and lines[i].strip().endswith("|")):
                add_markdown_table(doc, table_lines)
                table_lines = []
            continue

        if i + 1 < len(lines) and re.match(r"^[=]{3,}\s*$", lines[i + 1].strip()):
            p = doc.add_paragraph(stripped, style="Heading 1")
            for r in p.runs:
                set_run_font(r, bold=True)
            i += 2
            continue

        if i + 1 < len(lines) and re.match(r"^[-]{3,}\s*$", lines[i + 1].strip()):
            p = doc.add_paragraph(stripped, style="Heading 2")
            for r in p.runs:
                set_run_font(r, bold=True)
            i += 2
            continue

        if image_mode and stripped.startswith("建议插入图片："):
            paths = stripped.split("：", 1)[1]
            for part in re.split(r"[、,，]", paths):
                image_path = part.strip()
                if image_path:
                    add_image_if_exists(doc, image_path)
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            if skip_first_title and level == 1 and not first_title_skipped:
                first_title_skipped = True
                i += 1
                continue
            style = "Heading 1" if level <= 1 else "Heading 2" if level == 2 else "Heading 3"
            p = doc.add_paragraph(title, style=style)
            for r in p.runs:
                set_run_font(r, bold=True)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[2:].strip())
            set_run_font(run, size=10.8)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(re.sub(r"^\d+\.\s+", "", stripped))
            set_run_font(run, size=10.8)
            i += 1
            continue

        if stripped == "":
            i += 1
            continue

        # 元数据行在封面后不再重复放大，只作为正文小段落保留。
        p = add_para(doc, stripped)
        if stripped.startswith("图 "):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run_font(r, size=10, bold=True, color=PRESET["dark_blue"])
        i += 1


def build_report():
    md = (REPORT / "大作业技术报告草稿.md").read_text(encoding="utf-8")
    title = "基于体绘制与相空间刷选的 Nyx 宇宙密度演化可视分析"
    doc = setup_doc(title)
    add_cover(doc, title, "数据可视化课程大作业技术报告")
    add_meta_table(doc, [
        ("课程名称", "数据可视化"),
        ("作品类型", "科学可视化 / 体数据可视化 / 时序数据可视分析"),
        ("数据来源", "Nyx 宇宙学模拟密度数据"),
        ("实现工具", "MATLAB、Python、FFmpeg"),
    ])
    doc.add_page_break()
    convert_markdown(doc, md, image_mode=True, skip_first_title=True)
    out = REPORT / "Nyx宇宙密度演化可视分析_技术报告.docx"
    doc.save(out)
    return out


def build_simple_doc(md_name, out_name, title, subtitle):
    md = (REPORT / md_name).read_text(encoding="utf-8") if (REPORT / md_name).exists() else (ROOT / md_name).read_text(encoding="utf-8")
    doc = setup_doc(title)
    add_cover(doc, title, subtitle)
    doc.add_page_break()
    convert_markdown(doc, md, image_mode=False, skip_first_title=True)
    out = REPORT / out_name
    doc.save(out)
    return out


def main():
    outputs = [
        build_report(),
        build_simple_doc("Answer_Sheet_简版.md", "Nyx宇宙密度演化可视分析_Answer_Sheet.docx", "Answer Sheet", "Nyx 宇宙密度演化可视分析"),
        build_simple_doc("README.txt", "Nyx宇宙密度演化可视分析_README.docx", "README", "项目运行与文件说明"),
        build_simple_doc("代码说明.md", "Nyx宇宙密度演化可视分析_代码说明.docx", "代码说明", "MATLAB 项目结构与模块职责"),
    ]
    for out in outputs:
        print(out)


if __name__ == "__main__":
    main()
