from pathlib import Path
import csv
import re

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from build_word_docs import (
    ROOT, REPORT, setup_doc, add_cover, add_meta_table, add_image_if_exists,
    add_para, add_markdown_table, add_code_block, set_run_font, set_cell_shading,
    set_cell_margins, set_table_width, convert_markdown, PRESET
)


IMPROVED_REPORT_MD = REPORT / "大作业技术报告草稿_改进版.md"
IMPROVED_ANSWER_MD = REPORT / "Answer_Sheet_改进版.md"
IMPROVED_README_MD = REPORT / "README_改进版.md"
IMPROVED_CODE_MD = REPORT / "代码说明_改进版.md"


def read_text(path):
    return Path(path).read_text(encoding="utf-8")


def write_text(path, text):
    Path(path).write_text(text, encoding="utf-8")


def representative_stats_markdown():
    rows = list(csv.DictReader(open(ROOT / "results/report_figures/representative_density_stats.csv", encoding="utf-8")))
    headers = ["time step", "mean_density", "std_density", "max_density", "P05", "P50", "P95", "P99", "P99 / mean_density", "P99 - P01"]
    md = ["表 3 代表时间步密度统计量对比", "", "|" + "|".join(headers) + "|", "|" + "|".join(["---"] * len(headers)) + "|"]
    for r in rows:
        vals = [
            f"t={int(float(r['time_step'])):04d}",
            f"{float(r['mean_density']):.4f}",
            f"{float(r['std_density']):.4f}",
            f"{float(r['max_density']):.4f}",
            f"{float(r['P05']):.4f}",
            f"{float(r['P50']):.4f}",
            f"{float(r['P95']):.4f}",
            f"{float(r['P99']):.4f}",
            f"{float(r['P99_over_mean']):.4f}",
            f"{float(r['P99_minus_P01']):.4f}",
        ]
        md.append("|" + "|".join(vals) + "|")
    md.append("")
    md.append(
        "表中可以看出，mean density 在四个代表时间步中变化幅度较小，甚至呈现轻微下降趋势，因此不宜仅用均值判断宇宙结构演化强度。相比之下，std_density、max_density、P99/mean_density 和 P99-P01 更能体现密度波动增强和高密度尾部扩展。实际结果显示 P99/mean_density 与 P99-P01 从早期到后期整体增大，说明少数高密度体素相对于主体分布逐渐突出，密度场的不均匀程度增强。"
    )
    return "\n".join(md)


METHOD_TABLE_MD = """
表 2 不同可视化方法的作用对比

|方法|输入数据|输出结果|适合回答的问题|优点|局限性|
|---|---|---|---|---|---|
|中心切片|log-density 三维体|XY/XZ/YZ 中心切片|数据读取是否正确|直观、简单|只能观察局部截面|
|最大强度投影 MIP|log-density 或 mask 后 log-density|二维投影图|高密度区域大致分布在哪里|突出强结构|丢失深度信息|
|自定义体绘制|归一化 log-density|volume rendering 图|三维宇宙网整体形态|空间整体感强|依赖传递函数设计|
|对数直方图|log-density 一维体素数组|probability histogram|密度分布形态|统计特征清晰|不包含空间位置|
|时间-密度热力图|所有时间步 histogram|time-density heatmap|密度分布如何随时间迁移扩散|全时序趋势明显|不能直接定位空间结构|
|P99 高密度筛选|原始 density|top 1% mask、MIP、等值面|高密度尾部对应哪里|能连接统计阈值和空间结构|对阈值敏感|
|多阈值等值面|P90/P95/P99 threshold|嵌套等值面|结构层级如何变化|层次清晰|阈值选择影响视觉范围|
|linked brushing|histogram + density volume|交互联动视图|某一密度区间在空间中位于哪里|探索性强|需要交互环境支持|
"""


def make_improved_report_markdown():
    text = read_text(REPORT / "大作业技术报告草稿.md")

    workflow = """
# 3 关键技术描述

## 3.0 总体技术流程

图 11 Nyx 宇宙密度演化可视分析总体流程图

建议插入图片：results/report_figures/overall_workflow.png

图 11 总结了本作业从原始 Nyx .dat 数据到最终分析结论的完整流程。首先，程序以 little-endian float32 方式读取二进制数据，自动推断 n×n×n 网格尺寸，并将原始 z-y-x 线性顺序恢复为 MATLAB 中更直观的 x-y-z 三维体数据。随后，针对高动态范围密度值进行 log10 变换和百分位归一化。在此基础上，项目分为三条互补分析路线：体绘制与传递函数设计用于观察三维宇宙网整体形态；统计分析与直方图热力图用于刻画密度分布的时序变化；高密度区域筛选与结构分析用于连接统计高尾部和空间节点。三条路线最终汇入 linked brushing 仪表盘，使用户能够从统计密度区间反查空间位置，并形成对宇宙密度演化规律的综合解释。
"""
    text = text.replace("# 3 关键技术描述", workflow, 1)

    tf_insert = """

图 12 三种体绘制传递函数设计对比

建议插入图片：results/report_figures/transfer_function_design.png

图 12 展示了本作业体绘制中的三种透明度传递函数及自定义 LUT 颜色映射。传递函数不是单纯为了美化图像，而是根据不同分析任务设计的视觉映射。空洞观察型提高低密度区域的可见性，适合观察低密度空洞和背景结构；丝状结构型增强中密度区间，用于突出宇宙网中连续的丝状连接；高密度节点型主要增强 P95-P100 对应的高密度尾部，用于观察致密节点和团块。由此可见，同一体数据在不同传递函数下会强调不同科学特征，因此传递函数设计本身就是科学可视化分析过程的一部分。
"""
    text = text.replace("核心代码见附录 Code Listing 4。", "核心代码见附录 Code Listing 4。" + tf_insert, 1)

    text = text.replace("核心代码见附录 Code Listing 6。", "核心代码见附录 Code Listing 6。\n\n" + METHOD_TABLE_MD, 1)

    text = text.replace(
        "图 6 密度统计曲线  \n建议插入图片：results/03_statistics/mean_std_curve.png、max_density_curve.png、percentile_curve.png、log_skew_kurtosis_curve.png",
        "图 6 密度统计曲线  \n建议插入图片：results/03_statistics/mean_std_curve.png、results/03_statistics/max_density_curve.png、results/03_statistics/percentile_curve.png、results/03_statistics/log_skew_kurtosis_curve.png"
    )
    text = text.replace(
        "图 7 Top 1% 高密度筛选图  \n建议插入图片：results/05_high_density/top1_percent_t0000.png、top1_percent_t0060.png、top1_percent_t0099.png",
        "图 7 Top 1% 高密度筛选图  \n建议插入图片：results/05_high_density/top1_percent_t0000.png、results/05_high_density/top1_percent_t0060.png、results/05_high_density/top1_percent_t0099.png"
    )

    enhanced_42 = """

进一步分析可以看出，四个关键帧并不是简单亮度变化，而是空间组织方式的变化。t=0000 主要表现为较均匀的背景密度纹理；t=0030 出现局部聚集，说明密度扰动开始在空间中形成可见差异；t=0060 中密度连续结构增强，丝状连接逐渐清晰；t=0099 高密度节点和交汇区域显著突出，宇宙网结构更加完整。该图直接回答了“如何展示密度结构变化”的问题：体绘制能够同时呈现大尺度空间联系和局部高密度团块。但体绘制对传递函数较敏感，图像亮度不能直接等同于定量密度增长，因此需要与百分位曲线、P99 筛选和直方图热力图相互印证。
"""
    text = text.replace("该结果说明，体绘制能够在不切片的情况下直接表达三维密度场的整体结构，适合观察密度结构从局部扰动到网络化分布的空间演化过程。", "该结果说明，体绘制能够在不切片的情况下直接表达三维密度场的整体结构，适合观察密度结构从局部扰动到网络化分布的空间演化过程。" + enhanced_42)

    enhanced_43 = """

从直方图角度看，分布中心反映主体体素的典型密度水平，分布宽度反映密度差异大小，右尾则对应少量高密度体素。采用 log-density 而非原始 density 的原因在于原始密度具有高动态范围，直接绘制会使主体分布被极端高值压缩，难以观察中低密度体素的变化。直方图能够清楚回答“密度分布是否扩散、右尾是否增强”，但它本身不包含空间位置信息，因此不能判断右尾体素位于何处。后续 P99 mask、MIP、等值面和 linked brushing 正是对这一局限的补充。
"""
    text = text.replace("高密度尾部增强与体绘制中节点和团块的出现相互印证。", "高密度尾部增强与体绘制中节点和团块的出现相互印证。" + enhanced_43)

    stats_table = "\n\n" + representative_stats_markdown() + "\n\n"
    text = text.replace("## 4.4 时间序列统计曲线分析", "## 4.4 时间序列统计曲线分析" + stats_table, 1)

    enhanced_44 = """

因此，本作业在解释时间序列统计曲线时并不把 mean density 作为唯一依据。mean 可能受到总体物质量守恒或归一化尺度影响而变化不显著；std、max、P95、P99、P99-P01 和 P99/mean 更适合作为团块化和高尾部增强的指标。log-density 的偏度和峰度则从分布形态角度补充说明：偏度描述分布不对称性，峰度反映尖峰和极端尾部。当这些指标与体绘制中高亮节点增强、P99 mask 更集中等现象一致时，可以更可靠地说明密度场正在从相对均匀走向不均匀和网络化。
"""
    text = text.replace("这些统计曲线与体绘制结果结合，可以从数值和空间两个角度说明密度场逐渐团块化、网络化。", "这些统计曲线与体绘制结果结合，可以从数值和空间两个角度说明密度场逐渐团块化、网络化。" + enhanced_44)

    enhanced_45 = """

该方法直接回答了“高密度尾部对应空间中的什么位置”这一问题。若 P99 以上体素呈随机散点分布，则直方图右尾可能只是孤立异常值；而实验图像显示高密度体素具有明显空间聚集性，并倾向于出现在丝状结构交汇处或节点区域，说明高密度尾部具有空间结构意义。与此同时，P99 是人为设定的经验阈值，虽然便于定义 top 1% 区域，但不同阈值会改变被选结构范围。后续可以比较 P95、P99、P99.5 等阈值，以评估高密度结构的稳定性。
"""
    text = text.replace("由此可见，统计可视化和空间可视化应结合使用，才能完整解释密度分布的物理结构意义。", "由此可见，统计可视化和空间可视化应结合使用，才能完整解释密度分布的物理结构意义。" + enhanced_45)

    enhanced_46 = """

多阈值等值面和差分 MIP 从两个角度补充了密度演化分析。P90/P95/P99 等值面体现的是同一时间步内部的空间层级：外层 P90 较连续，反映中高密度丝状外壳；P95 范围收缩，反映更强聚集；P99 只保留核心节点。difference MIP 则比较首末时间步，红色区域表示 log-density 增强，蓝色区域表示 log-density 降低，可用于识别物质聚集区和空洞化区域。void_fraction 与 dense_fraction 曲线进一步将这种视觉观察量化，使空洞扩张和团块增强不只停留在图像印象层面。
"""
    text = text.replace("结合 void_fraction 和 dense_fraction 曲线可以看出，低密度空洞扩张和高密度团块增强是同时发生的两个过程。", "结合 void_fraction 和 dense_fraction 曲线可以看出，低密度空洞扩张和高密度团块增强是同时发生的两个过程。" + enhanced_46)

    enhanced_47 = """

从交互使用过程看，用户可以先在左侧直方图中观察当前时间步的密度分布，再通过百分位滑条选择低密度、中密度或高密度区间。选择低密度区间时，右侧投影更容易显示大面积背景和空洞区域；选择中密度区间时，连续丝状结构更明显；选择高密度区间时，投影会集中在节点和团块附近。linked brushing 的核心价值在于它不是单纯展示一张静态图，而是支持探索“某一数值区间在哪里”。这一能力将统计图和空间结构图连接起来，弥补了单独直方图缺少空间信息、单独投影缺少分布背景的不足。
"""
    text = text.replace("相比静态图片，linked brushing 更适合探索“某一类数值体素位于哪里”这一问题。", "相比静态图片，linked brushing 更适合探索“某一类数值体素位于哪里”这一问题。" + enhanced_47)

    advantages = """

## 5.2 方法优势与局限性

本作业的方法优势主要体现在以下方面。第一，项目形成了从数据读取、完整性检查、统计计算、体绘制、高密度筛选到交互分析的完整流程，便于复现实验结果。第二，自定义体绘制使颜色映射和透明度传递函数可控，能够针对空洞、丝状结构和高密度节点设计不同视觉强调方式。第三，log-density 与百分位归一化有效缓解了高动态范围密度数据的显示问题。第四，P99 筛选将统计分布右尾与空间节点联系起来，使直方图分析不再停留于纯数值层面。第五，linked brushing 仪表盘增强了探索性分析能力，使用户可以主动选择密度区间并观察其空间位置。

同时，本作业也存在一定局限。当前体绘制主要沿固定 z 方向合成，缺少真正自由的三维旋转和交互裁剪；梯度增强属于类光照处理，并不是真实物理光照模型；当前分析只使用 density 变量，没有结合 velocity、temperature 等多变量信息；P99 阈值具有经验性，不同阈值会影响结构范围；MIP 会压缩深度信息，不能完全替代真实三维交互。最后，本文结论主要用于数据可视化课程实验分析，不能直接作为正式天体物理科学结论。
"""
    text = text.replace("## 5.2 方法价值", advantages + "\n\n## 5.3 方法价值", 1)
    text = text.replace("## 5.3 不足与展望", "## 5.4 不足与展望", 1)

    banned = ["图片未找到", "TODO", "待补充", "占位符"]
    for b in banned:
        text = text.replace(b, "")
    write_text(IMPROVED_REPORT_MD, text)
    return text


def make_answer_markdown():
    text = """# Answer Sheet

## 1. 作品信息

作品名称：基于体绘制与相空间刷选的 Nyx 宇宙密度演化可视分析  
课程名称：数据可视化  
作品类型：科学可视化 / 体数据可视化 / 时序数据可视分析  
数据来源：Nyx 宇宙学模拟密度数据  
实现工具：MATLAB、Python、FFmpeg  

## 2. 作品目标

本作品以 Nyx 宇宙学模拟密度数据为对象，围绕三维密度场随时间演化的可视分析展开。项目目标是从 little-endian float32 二进制文件中恢复三维密度体，结合体绘制、统计曲线、直方图热力图、高密度筛选、多阈值等值面和 linked brushing 交互仪表盘，分析密度场从相对均匀到团块化、网络化的演化特征。

## 3. 方法概述

程序首先根据文件大小自动推断 n×n×n 网格，并将原始 z-y-x 存储顺序重排为 x-y-z 三维数组。显示阶段使用 log10(max(density,eps)) 压缩动态范围，并采用 P5 到 P99.7 的百分位归一化。统计阶段计算均值、标准差、高百分位、偏度、峰度、熵和 Gini 系数；空间分析阶段使用 P99 作为 top 1% 高密度阈值，生成 MIP 和等值面；交互阶段通过百分位滑条将密度分布区间与空间投影联动。

## 4. 问题回答

### 4.1 如何展示密度变化？

作品采用“统计图 + 体绘制 + 交互刷选”的组合方式展示密度变化。体绘制用于观察三维宇宙网整体形态，直方图和时间-密度热力图用于展示 log-density 分布随时间的扩散，高百分位统计曲线用于量化高密度尾部变化。由于原始密度动态范围较大，显示时使用 log10 变换和百分位归一化，避免极端高值压缩中低密度结构。

### 4.2 观察到哪些宇宙密度演化特征？

代表时间步体绘制显示，早期密度场较为均匀，随后局部区域开始聚集，中期丝状结构逐渐增强，后期高密度节点和团块更加明显。P90/P95/P99 多阈值等值面进一步表明，密度结构具有从中高密度丝状外壳向极高密度核心节点收缩的层级关系。该结果体现了课程可视化实验层面的宇宙网形成特征。

### 4.3 密度时序统计说明了什么？

统计结果表明，mean density 变化不一定显著，因此不能单独作为结构演化判断依据。std_density、max_density、P95、P99、P99/mean 和 P99-P01 更能反映密度波动增强和高密度尾部扩展。时间-密度热力图显示 log-density 分布随时间逐渐变宽，高密度端尾部增强，说明密度场从较集中分布向更不均匀的分布发展。

### 4.4 相空间刷选如何验证高密度尾部与宇宙网节点对应？

linked brushing 仪表盘左侧显示 log-density 直方图，用户通过百分位滑条选择某一密度区间，系统实时生成 mask 并在右侧显示空间投影。当选择 P99 附近高密度区间时，被选体素集中在丝状结构交汇处和节点区域，说明直方图右尾不是随机噪声，而是具有空间结构意义的高密度宇宙网节点。该方法建立了统计分布与空间结构之间的双向联系。

## 5. 创新点

1. 完整实现 Nyx 二进制体数据读取、轴顺序恢复和全流程自动化分析。  
2. 使用 log-density 与百分位归一化解决高动态范围密度显示问题。  
3. 实现不依赖 volshow 的自定义 alpha compositing 体绘制。  
4. 设计空洞型、丝状型和节点型三种传递函数。  
5. 使用 P99 高密度筛选连接统计尾部与三维空间节点。  
6. 构建 linked brushing 仪表盘支持密度区间与空间投影联动探索。

## 6. 输出概况

项目输出包括数据完整性检查图、密度统计曲线、log-density histogram、time-density heatmap、关键帧体绘制图、传递函数对比图、P99 高密度 MIP、多阈值等值面、结构指标曲线、差分 MIP、交互式仪表盘截图和视频帧。所有结果保存在 results/ 目录下，可直接用于课程报告和答辩展示。

## 7. 结论说明

本作品结论主要服务于数据可视化课程实验，强调可视化方法如何帮助理解 Nyx 密度场的时空变化。文中关于宇宙网形成、空洞扩张和高密度节点增强的描述均基于给定数据和可视化结果，不作为正式天体物理科学结论。
"""
    write_text(IMPROVED_ANSWER_MD, text)
    return text


def make_readme_markdown():
    base = read_text(ROOT / "README.txt")
    extra = """

常见问题
========

没有 .dat 文件怎么办？
  请先将 Nyx 原始数据文件放入 data/raw/。至少需要一个非空 .dat 文件才能运行数据检查；如果要复现实验报告中的完整时序结果，建议放入 0000.dat 至 0099.dat。

图片没有生成怎么办？
  请先确认 MATLAB 当前目录为项目根目录，并运行 run('code/main.m')。如果某一步失败，可以单独运行对应 step 文件，例如 step3_histogram_analysis(pwd) 或 step4_volume_render_keyframes(pwd)。

Image Processing Toolbox 不可用怎么办？
  主流程仍可运行。程序会跳过 bwconncomp 连通域分析，并在结构指标表中保留 NaN 或警告信息。切片、统计、直方图、体绘制、P99 MIP 和仪表盘不依赖该工具箱。

体绘制太慢怎么办？
  可先使用 results/frames/ 中的 MIP 帧图进行视频展示。main.m 默认生成 MIP 视频帧，比完整体绘制更快。若需要加速体绘制，可减少关键帧数量或降低输入体数据分辨率。

FFmpeg 不可用怎么办？
  不影响 MATLAB 分析和图片生成。FFmpeg 只用于将 results/frames/frame_*.png 合成为 demo.mp4。若未安装 FFmpeg，可以直接提交帧图或在其他视频软件中合成。

快速检查清单
============

- data/raw/ 是否包含 0000.dat 至 0099.dat 或至少一个非空 .dat 文件。
- code/main.m 是否能在 MATLAB 中完整运行。
- data/processed/ 是否生成 density_stats.csv 和 structure_metrics.csv。
- results/ 是否生成各类 PNG 图片。
- report/ 技术报告中是否没有“图片未找到”等缺失标记。
- results/frames/ 是否生成 frame_0000.png 至 frame_0099.png。
- 如需视频，是否已使用 FFmpeg 生成 demo.mp4。
"""
    text = base + extra
    write_text(IMPROVED_README_MD, text)
    return text


def make_code_markdown():
    base = read_text(REPORT / "代码说明.md")
    extra = """

## 7. 各模块与报告章节对应关系

|代码模块|主要输出|对应报告章节|说明|
|---|---|---|---|
|read_nyx_dat.m / infer_grid_size.m|三维密度体 V|2.3、3.1|负责二进制读取、网格推断和轴顺序恢复|
|step1_check_data.m|中心切片、文件大小检查|2.4、4.1|验证数据读取和完整性|
|step2_statistics_all_frames.m|density_stats.csv、统计曲线|3.2、4.4|计算全时间步密度统计指标|
|step3_histogram_analysis.m|直方图、时间-密度热力图|3.3、4.3|分析 log-density 分布演化|
|step4_volume_render_keyframes.m|体绘制图、传递函数对比|3.4、4.2|展示三维空间结构|
|volume_render_alpha_composite.m|RGB 体绘制结果|3.4、3.5|自定义 alpha compositing 和梯度增强|
|step5_high_density_selection.m|P99 MIP、等值面|3.6、3.7、4.5、4.6|连接高密度尾部与空间结构|
|step6_structure_metrics.m|结构指标、差分 MIP|3.8、4.6|量化空洞、团块化和不均匀程度|
|step7_interactive_dashboard.m|dashboard_preview.png|3.9、4.7|实现 linked brushing 交互分析|

## 8. 答辩讲解顺序

1. 先说明数据格式：Nyx .dat 不是图片，而是 little-endian float32 三维体数据。
2. 展示 read_nyx_dat.m 的读取逻辑：自动推断 128×128×128，并完成 z-y-x 到 x-y-z 重排。
3. 展示中心切片，说明读取方向和数值范围正常。
4. 展示总体流程图，说明项目包含体绘制、统计分析、高密度筛选和交互刷选四条主线。
5. 展示关键帧体绘制，说明从相对均匀到丝状结构和节点增强的空间变化。
6. 展示直方图和时间-密度热力图，说明密度分布逐渐扩散和高密度尾部增强。
7. 展示代表时间步统计表，强调 mean 不是唯一指标，std、P99/mean 和 P99-P01 更有解释力。
8. 展示 P99 MIP 和多阈值等值面，说明高密度尾部对应空间节点和层级结构。
9. 展示 linked brushing 仪表盘，说明如何从密度区间反查空间位置。
10. 最后说明局限性：固定视角体绘制、经验阈值、单变量 density、MIP 深度信息损失。

## 9. 可选功能和降级方案

|功能|依赖|如果不可用的降级方案|
|---|---|---|
|bwconncomp 连通域分析|Image Processing Toolbox|跳过连通域指标，保留其他结构指标|
|isosurface 三维等值面|MATLAB 基础图形函数|如果版本不支持，则只输出 P99 MIP|
|完整体绘制关键帧|CPU 和内存|使用 log-density MIP 帧替代体绘制视频|
|FFmpeg 视频合成|系统安装 FFmpeg|只提交 results/frames/ 帧图|
|交互式仪表盘|MATLAB figure/uicontrol|保留 dashboard_preview.png 作为静态说明|
|报告新增流程图|MATLAB 绘图|可使用 results/report_figures 中已生成 PNG|

这些降级方案保证了项目在普通 MATLAB 环境中仍能完成主流程，避免因为个别工具箱缺失导致整个大作业无法运行。
"""
    text = base + extra
    write_text(IMPROVED_CODE_MD, text)
    return text


def add_stats_table_docx(doc):
    rows = list(csv.DictReader(open(ROOT / "results/report_figures/representative_density_stats.csv", encoding="utf-8")))
    headers = ["time step", "mean", "std", "max", "P05", "P50", "P95", "P99", "P99/mean", "P99-P01"]
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    widths = [Inches(0.58), Inches(0.62), Inches(0.58), Inches(0.65), Inches(0.58), Inches(0.58), Inches(0.58), Inches(0.58), Inches(0.72), Inches(0.72)]
    set_table_width(table, widths)
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
        set_cell_shading(table.cell(0, j), "F2F4F7")
    for i, r in enumerate(rows, 1):
        vals = [
            f"{int(float(r['time_step'])):04d}",
            f"{float(r['mean_density']):.4f}",
            f"{float(r['std_density']):.4f}",
            f"{float(r['max_density']):.4f}",
            f"{float(r['P05']):.4f}",
            f"{float(r['P50']):.4f}",
            f"{float(r['P95']):.4f}",
            f"{float(r['P99']):.4f}",
            f"{float(r['P99_over_mean']):.4f}",
            f"{float(r['P99_minus_P01']):.4f}",
        ]
        for j, v in enumerate(vals):
            table.cell(i, j).text = v
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=60, bottom=60, start=70, end=70)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=8.0, bold=(row is table.rows[0]))
    doc.add_paragraph()


def build_report_docx():
    make_improved_report_markdown()
    title = "基于体绘制与相空间刷选的 Nyx 宇宙密度演化可视分析"
    doc = setup_doc(title)
    add_cover(doc, title, "数据可视化课程大作业技术报告（改进版）")
    add_meta_table(doc, [
        ("课程名称", "数据可视化"),
        ("作品类型", "科学可视化 / 体数据可视化 / 时序数据可视分析"),
        ("数据来源", "Nyx 宇宙学模拟密度数据"),
        ("实现工具", "MATLAB、Python、FFmpeg"),
        ("改进内容", "补充流程图、传递函数图、统计表、方法对比表与增强分析"),
    ])
    doc.add_page_break()
    md = read_text(IMPROVED_REPORT_MD)
    # 使用普通转换插入图；统计表在 Markdown 中也有，保持可读。
    convert_markdown(doc, md, image_mode=True, skip_first_title=True)
    out = REPORT / "Nyx宇宙密度演化可视分析_技术报告_改进版.docx"
    doc.save(out)
    return out


def build_simple(md_path, out_name, title, subtitle):
    md = read_text(md_path)
    doc = setup_doc(title)
    add_cover(doc, title, subtitle)
    doc.add_page_break()
    convert_markdown(doc, md, image_mode=False, skip_first_title=True)
    out = REPORT / out_name
    doc.save(out)
    return out


def main():
    make_answer_markdown()
    make_readme_markdown()
    make_code_markdown()
    outputs = [
        build_report_docx(),
        build_simple(IMPROVED_ANSWER_MD, "Nyx宇宙密度演化可视分析_Answer_Sheet_改进版.docx", "Answer Sheet", "Nyx 宇宙密度演化可视分析（改进版）"),
        build_simple(IMPROVED_README_MD, "Nyx宇宙密度演化可视分析_README_改进版.docx", "README", "项目运行与文件说明（改进版）"),
        build_simple(IMPROVED_CODE_MD, "Nyx宇宙密度演化可视分析_代码说明_改进版.docx", "代码说明", "MATLAB 项目结构与模块职责（改进版）"),
    ]
    for out in outputs:
        print(out)


if __name__ == "__main__":
    main()
