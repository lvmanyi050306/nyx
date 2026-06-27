from pathlib import Path
import re
import csv

from PIL import Image, ImageOps, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "report"
FIGDIR = ROOT / "results" / "report_figures"

FONT_CN = "SimSun"
FONT_HEI = "SimHei"
FONT_EN = "Times New Roman"
FONT_CODE = "Consolas"
BODY_PT = 10.5


def rel(path):
    return str(Path(path).as_posix())


def ensure_report_figures():
    FIGDIR.mkdir(parents=True, exist_ok=True)
    make_collage(
        [
            ROOT / "results/03_statistics/mean_std_curve.png",
            ROOT / "results/03_statistics/max_density_curve.png",
            ROOT / "results/03_statistics/percentile_curve.png",
            ROOT / "results/03_statistics/log_skew_kurtosis_curve.png",
        ],
        ["Mean ± Std", "Max density", "Percentiles", "Log skewness / kurtosis"],
        FIGDIR / "combined_statistics_curves.png",
        cols=2,
    )
    make_collage(
        [
            ROOT / "results/05_high_density/top1_percent_t0000.png",
            ROOT / "results/05_high_density/top1_percent_t0060.png",
            ROOT / "results/05_high_density/top1_percent_t0099.png",
        ],
        ["t=0000", "t=0060", "t=0099"],
        FIGDIR / "combined_top1_percent_mips.png",
        cols=1,
    )


def make_collage(paths, labels, out, cols=2):
    images = []
    for p in paths:
        im = Image.open(p).convert("RGB")
        im = ImageOps.contain(im, (980, 620), method=Image.LANCZOS)
        canvas = Image.new("RGB", (1000, 690), "white")
        canvas.paste(im, ((1000 - im.width) // 2, 42))
        d = ImageDraw.Draw(canvas)
        try:
            font = ImageFont.truetype("arial.ttf", 28)
        except Exception:
            font = ImageFont.load_default()
        label = labels[len(images)]
        d.text((24, 10), label, fill=(25, 40, 70), font=font)
        images.append(canvas)
    rows = (len(images) + cols - 1) // cols
    w, h = images[0].size
    collage = Image.new("RGB", (cols * w, rows * h), "white")
    for i, im in enumerate(images):
        collage.paste(im, ((i % cols) * w, (i // cols) * h))
    collage.save(out, quality=95)


def set_run_font(run, size=BODY_PT, bold=None, color=None, cn_font=FONT_CN, en_font=FONT_EN):
    run.font.name = en_font
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)
    run._element.rPr.rFonts.set(qn("w:ascii"), en_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), en_font)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_base(p, first_indent=True, keep_with_next=False):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    if first_indent:
        pf.first_line_indent = Pt(21)
    if keep_with_next:
        pf.keep_with_next = True
    for r in p.runs:
        set_run_font(r)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill="F2F4F7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders_three_line(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "bottom", "insideH"]:
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "10" if edge in ["top", "bottom"] else "4")
        elem.set(qn("w:color"), "666666")
    for edge in ["left", "right", "insideV"]:
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "nil")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for i, width in enumerate(widths):
            if i >= len(row.cells):
                continue
            cell = row.cells[i]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_field_run(paragraph, instr):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    return run


def add_page_number_footer(section, include_header=True):
    section.different_first_page_header_footer = True
    if include_header:
        hp = section.header.paragraphs[0]
        hp.text = "Nyx 宇宙密度演化可视分析"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in hp.runs:
            set_run_font(r, size=9, color=RGBColor(120, 120, 120), cn_font=FONT_CN)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for txt in ["第 "]:
        r = fp.add_run(txt)
        set_run_font(r, size=9, color=RGBColor(100, 100, 100))
    add_field_run(fp, "PAGE")
    r = fp.add_run(" 页 / 共 ")
    set_run_font(r, size=9, color=RGBColor(100, 100, 100))
    add_field_run(fp, "NUMPAGES")
    r = fp.add_run(" 页")
    set_run_font(r, size=9, color=RGBColor(100, 100, 100))


def setup_doc(include_header=True):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    add_page_number_footer(sec, include_header=include_header)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_EN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    normal.font.size = Pt(BODY_PT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.first_line_indent = Pt(21)

    heading_specs = [
        ("Heading 1", 15, 12, 6),
        ("Heading 2", 14, 8, 4),
        ("Heading 3", 12, 6, 3),
    ]
    for name, size, before, after in heading_specs:
        st = styles[name]
        st.font.name = FONT_HEI
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEI)
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(31, 77, 120) if name != "Heading 1" else RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    return doc


def add_cover_report(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(100)
    r = p.add_run("基于体绘制与相空间刷选的 Nyx 宇宙密度演化可视分析")
    set_run_font(r, size=22, bold=True, cn_font=FONT_HEI, en_font=FONT_EN)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(60)
    r = p.add_run("数据可视化课程大作业技术报告")
    set_run_font(r, size=16, bold=True, cn_font=FONT_HEI, en_font=FONT_EN)

    rows = [
        ("课程名称", "数据可视化"),
        ("作品类型", "科学可视化 / 体数据可视化 / 时序数据可视分析"),
        ("数据来源", "Nyx 宇宙学模拟密度数据"),
        ("实现工具", "MATLAB、Python、FFmpeg"),
        ("姓名", ""),
        ("学号", ""),
        ("班级", ""),
        ("完成日期", ""),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, [Inches(1.6), Inches(4.5)])
    set_table_borders_three_line(table)
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k + "："
        table.cell(i, 1).text = v
        for j in range(2):
            c = table.cell(i, j)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(c, top=90, bottom=90, start=120, end=120)
            for pp in c.paragraphs:
                pp.paragraph_format.first_line_indent = Pt(0)
                for rr in pp.runs:
                    set_run_font(rr, size=11, bold=(j == 0), cn_font=FONT_CN)
    doc.add_page_break()


def add_simple_cover(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run(title)
    set_run_font(r, size=22, bold=True, cn_font=FONT_HEI)
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        r = p.add_run(subtitle)
        set_run_font(r, size=12, color=RGBColor(90, 90, 90), cn_font=FONT_HEI)
    doc.add_page_break()


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目录")
    set_run_font(r, size=16, bold=True, cn_font=FONT_HEI)
    p.paragraph_format.space_after = Pt(20)
    toc = doc.add_paragraph()
    toc.paragraph_format.first_line_indent = Pt(0)
    run = toc.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "右键更新目录"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)
    set_run_font(run, size=10.5)
    doc.add_page_break()


def add_heading(doc, text, level):
    p = doc.add_paragraph(text, style=f"Heading {min(level,3)}")
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.keep_with_next = True
    for r in p.runs:
        set_run_font(r, size={1: 15, 2: 14, 3: 12}.get(level, 12), bold=True, cn_font=FONT_HEI)
    return p


def add_body_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if style in ("List Bullet", "List Number"):
        p.paragraph_format.first_line_indent = None
    else:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = False
    r = p.add_run(text)
    set_run_font(r, size=9, bold=False, cn_font=FONT_CN)
    return p


def add_image(doc, image_path, caption, wide=False):
    path = ROOT / image_path
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    width = Inches(6.0 if wide else 5.3)
    run.add_picture(str(path), width=width)
    add_caption(doc, caption)


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Cm(0.25)
    p.paragraph_format.right_indent = Cm(0.25)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(code.rstrip())
    set_run_font(r, size=9, cn_font=FONT_CODE, en_font=FONT_CODE)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    p_pr.append(shd)
    pbdr = OxmlElement("w:pBdr")
    for edge in ["top", "left", "bottom", "right"]:
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "D0D0D0")
        pbdr.append(e)
    p_pr.append(pbdr)


def parse_markdown_table(lines):
    rows = []
    for line in lines:
        if re.match(r"^\|\s*[-:]+", line):
            continue
        rows.append([c.strip() for c in line.strip().strip("|").split("|")])
    return rows


def add_table_from_rows(doc, rows, title=None, wide=False):
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(title)
        set_run_font(r, size=9, cn_font=FONT_CN)
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    set_table_borders_three_line(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total = 6.2
    if cols <= 2:
        widths = [Inches(1.7), Inches(4.5)]
    elif cols >= 8:
        widths = [Inches(total / cols)] * cols
    else:
        widths = [Inches(total / cols)] * cols
    set_table_width(table, widths)
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.text = row[j] if j < len(row) else ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=70, bottom=70, start=80, end=80)
            if i == 0:
                set_cell_shading(cell, "EDEDED")
            for pp in cell.paragraphs:
                pp.paragraph_format.first_line_indent = Pt(0)
                pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT if (i > 0 and re.fullmatch(r"[-+]?\\d+(\\.\\d+)?", pp.text.strip())) else WD_ALIGN_PARAGRAPH.LEFT
                for rr in pp.runs:
                    set_run_font(rr, size=8.5 if cols >= 6 else 9.5, bold=(i == 0))
    doc.add_paragraph()


def preprocess_report_markdown(md):
    # 删除手工目录块，只保留自动目录。
    md = re.sub(r"## 目录.*?(?=## 摘要)", "", md, flags=re.S)
    # 删除封面元信息，封面单独重建。
    md = re.sub(r"课程名称：.*?说明：.*?\n", "", md, flags=re.S)
    md = md.replace("实验环境如下：\n\n| 项目 | 说明 |", "实验环境如下：\n\n表 1 实验环境与工具配置\n\n| 项目 | 说明 |")
    md = md.replace("表 1 代表时间步密度统计量对比", "表 3 代表时间步密度统计量对比")
    md = md.replace("表 2 不同可视化方法的作用对比", "表 2 不同可视化方法的作用对比")
    # 替换多图组为组合图。
    md = md.replace(
        "建议插入图片：results/03_statistics/mean_std_curve.png、results/03_statistics/max_density_curve.png、results/03_statistics/percentile_curve.png、results/03_statistics/log_skew_kurtosis_curve.png",
        "建议插入图片：results/report_figures/combined_statistics_curves.png",
    )
    md = md.replace(
        "建议插入图片：results/05_high_density/top1_percent_t0000.png、results/05_high_density/top1_percent_t0060.png、results/05_high_density/top1_percent_t0099.png",
        "建议插入图片：results/report_figures/combined_top1_percent_mips.png",
    )
    md = md.replace("图 11 Nyx 宇宙密度演化可视分析总体流程图", "图 1 Nyx 宇宙密度演化可视分析总体流程图")
    md = md.replace("图 12 三种体绘制传递函数设计对比", "图 2 三种体绘制传递函数设计对比")
    for bad in ["图片未找到", "TODO", "待补充", "占位符"]:
        md = md.replace(bad, "")
    return md


def scan_figure_captions(md):
    fig_entries = []
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        m = re.match(r"^图\s*\d+\s+(.+)$", s)
        if m:
            title = m.group(1).strip()
            j = i + 1
            while j < len(lines) and lines[j].strip() == "":
                j += 1
            if j < len(lines) and lines[j].strip().startswith(("建议插入图片：", "建议插入图片:")):
                fig_entries.append((s, title))
        i += 1
    mapping = {}
    for idx, (raw, title) in enumerate(fig_entries, 1):
        mapping[raw] = f"图 {idx} {title}"
    return mapping


def convert_markdown_to_doc(doc, md, is_report=False):
    md = preprocess_report_markdown(md) if is_report else md
    fig_map = scan_figure_captions(md)
    lines = md.splitlines()
    i = 0
    in_code = False
    code_lines = []
    table_lines = []
    pending_table_title = None
    pending_figure_raw = None
    pending_figure_caption = None

    while i < len(lines):
        raw = lines[i]
        s = raw.strip()

        if s.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, "\n".join(code_lines))
                in_code = False
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not s:
            i += 1
            continue

        if s == "Nyx_Visualization_Project/":
            block = [raw]
            i += 1
            while i < len(lines):
                nxt = lines[i]
                if nxt.startswith("  ") or nxt.strip().startswith(("├", "└", "│")):
                    block.append(nxt)
                    i += 1
                else:
                    break
            add_code_block(doc, "\n".join(block))
            continue

        # Setext headings in README.
        if i + 1 < len(lines) and re.match(r"^[=]{3,}\s*$", lines[i + 1].strip()):
            add_heading(doc, s, 1)
            i += 2
            continue
        if i + 1 < len(lines) and re.match(r"^[-]{3,}\s*$", lines[i + 1].strip()):
            add_heading(doc, s, 2)
            i += 2
            continue

        hm = re.match(r"^(#{1,6})\s+(.*)$", s)
        if hm:
            level = min(len(hm.group(1)), 3)
            title = hm.group(2).strip()
            # 顶部重复主标题跳过。
            if level == 1 and "基于体绘制与相空间刷选" in title:
                i += 1
                continue
            add_heading(doc, title, level)
            i += 1
            continue

        if re.match(r"^表\s*\d+\s+", s):
            pending_table_title = s
            i += 1
            continue

        if re.match(r"^图\s*\d+\s+", s):
            pending_figure_raw = s
            pending_figure_caption = fig_map.get(s, s)
            # 如果下一行不是建议插图，那就作为普通居中图题处理。
            i += 1
            continue

        if s.startswith(("建议插入图片：", "建议插入图片:")) and pending_figure_caption:
            sep = "：" if "：" in s else ":"
            paths = [p.strip() for p in re.split(r"[、,，]", s.split(sep, 1)[1]) if p.strip()]
            for pth in paths:
                wide = any(key in pth for key in ["workflow", "heatmap", "compare", "combined", "transfer_function_design", "dashboard", "nested"])
                add_image(doc, pth, pending_figure_caption, wide=wide)
            pending_figure_raw = None
            pending_figure_caption = None
            i += 1
            continue

        if s.startswith("|") and s.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = parse_markdown_table(table_lines)
            add_table_from_rows(doc, rows, title=pending_table_title)
            pending_table_title = None
            continue

        if s.startswith("- "):
            add_body_paragraph(doc, s[2:].strip(), style="List Bullet")
            i += 1
            continue
        if re.match(r"^\d+\.\s+", s):
            add_body_paragraph(doc, re.sub(r"^\d+\.\s+", "", s), style="List Number")
            i += 1
            continue

        # README checklist.
        if s.startswith("□"):
            add_body_paragraph(doc, s, style="List Bullet")
            i += 1
            continue

        # Code-like one-liners.
        if re.match(r"^(run\(|ffmpeg|step\d|[A-Za-z]:\\|Nyx_Visualization_Project/)", s):
            add_code_block(doc, s)
            i += 1
            continue

        p = add_body_paragraph(doc, s)
        # Code listing titles.
        if s.startswith("Code Listing"):
            for r in p.runs:
                r.bold = True
            p.paragraph_format.keep_with_next = True
        i += 1


def build_technical_report():
    ensure_report_figures()
    md = (REPORT / "大作业技术报告草稿_改进版.md").read_text(encoding="utf-8")
    doc = setup_doc(include_header=True)
    add_cover_report(doc)
    add_toc(doc)
    convert_markdown_to_doc(doc, md, is_report=True)
    out = REPORT / "Nyx宇宙密度演化可视分析_技术报告_排版优化版.docx"
    doc.save(out)
    return out


def build_simple_doc(md_file, out_name, title, subtitle=None, include_header=False):
    md = (REPORT / md_file).read_text(encoding="utf-8")
    doc = setup_doc(include_header=include_header)
    add_simple_cover(doc, title, subtitle)
    convert_markdown_to_doc(doc, md, is_report=False)
    out = REPORT / out_name
    doc.save(out)
    return out


def build_all():
    outs = [
        build_technical_report(),
        build_simple_doc("Answer_Sheet_改进版.md", "Nyx宇宙密度演化可视分析_Answer_Sheet_排版优化版.docx", "Nyx 宇宙密度演化可视分析 Answer Sheet"),
        build_simple_doc("README_改进版.md", "Nyx宇宙密度演化可视分析_README_排版优化版.docx", "Nyx 宇宙密度演化可视分析 README", "项目运行手册"),
        build_simple_doc("代码说明_改进版.md", "Nyx宇宙密度演化可视分析_代码说明_排版优化版.docx", "Nyx 宇宙密度演化可视分析代码说明", "模块说明书"),
    ]
    for out in outs:
        print(out)


if __name__ == "__main__":
    build_all()
